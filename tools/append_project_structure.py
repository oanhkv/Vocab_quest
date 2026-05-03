"""
Append PHẦN XVI — CẤU TRÚC DỰ ÁN vào VocabQuest_TongHop_v2.docx
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TongHop_v2.docx'
d = Document(path)


def add_table(rows):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            if r_idx == 0:
                run.bold = True
    return t


def add_code_block(code):
    p = d.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    return p


def H1(t): d.add_paragraph(t, style='Heading 1')
def H2(t): d.add_paragraph(t, style='Heading 2')
def P(t):  d.add_paragraph(t, style='Normal')


# ===== TITLE =====
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN XVI — CẤU TRÚC DỰ ÁN', style='Title')

H1('1. Cấu trúc thư mục gốc')
P('Project sử dụng cấu trúc chuẩn Flutter, gồm thư mục lib/ chứa toàn bộ mã Dart, '
  'các thư mục build cho từng nền tảng (android/ios/web/windows/macos/linux), assets/ '
  'cho tài nguyên tĩnh, và test/ cho unit/widget test.')

tree = '''vocab_quest/
├── android/                ← native Android (Gradle, AndroidManifest, google-services.json)
├── ios/                    ← native iOS (Info.plist, Runner.xcodeproj)
├── web/                    ← bootstrap HTML/manifest cho Flutter Web
├── windows/  macos/  linux/← Flutter desktop bootstrap (auto-generated plugin registrant)
├── assets/
│   ├── audio/              ← âm thanh sound effect (click, correct, wrong)
│   ├── data/               ← JSON từ vựng (vocab_beginner.json, ...)
│   └── images/             ← icon, logo, ảnh minh hoạ
├── lib/                    ← TOÀN BỘ MÃ DART
│   ├── main.dart           ← entry point
│   ├── firebase_options.dart
│   ├── config/             ← cấu hình hằng số, theme, design tokens
│   ├── models/             ← data classes (POJO Dart)
│   ├── providers/          ← state management (ChangeNotifier)
│   ├── services/           ← gọi Firebase, audio, notification, JSON
│   ├── screens/            ← các màn hình UI
│   │   ├── auth/           ← Login, Register
│   │   ├── home/           ← Home dashboard
│   │   ├── games/          ← 4 mini-game + flow chọn pack/level/result
│   │   ├── leaderboard/    ← bảng xếp hạng
│   │   ├── profile/        ← hồ sơ + edit
│   │   ├── settings/       ← cài đặt
│   │   ├── shop/           ← cửa hàng pack
│   │   ├── favorites/      ← yêu thích
│   │   ├── history/        ← lịch sử chơi
│   │   └── splash_screen.dart
│   ├── widgets/            ← widget tái sử dụng
│   └── utils/              ← helper: validator, localization, streak calc, ...
├── test/
│   └── widget_test.dart    ← test mặc định (chưa mở rộng)
├── tools/                  ← script Python cho tài liệu (không liên quan app runtime)
├── pubspec.yaml            ← khai báo dependency Flutter
├── pubspec.lock            ← lock version
├── analysis_options.yaml   ← cấu hình lint
├── firebase.json           ← cấu hình Firebase CLI
├── firestore.rules         ← Security Rules Firestore
├── firestore.indexes.json  ← khai báo composite index
├── storage.rules           ← Security Rules cho Firebase Storage
├── README.md
└── .gitignore
'''
add_code_block(tree)

# ============================================================
# 2. Root config files
# ============================================================
H1('2. File cấu hình root')
add_table([
    ['File', 'Vai trò'],
    ['pubspec.yaml',
     'Khai báo metadata app + dependency Flutter (firebase_*, provider, '
     'shared_preferences, flutter_tts, lucide_icons, lottie, ...).'],
    ['analysis_options.yaml', 'Cấu hình lint rule cho Dart analyzer.'],
    ['firebase.json',
     'Liên kết project với Firebase CLI (rules, indexes, hosting).'],
    ['firestore.rules',
     'Security Rules cho Cloud Firestore — quy định ai được đọc/ghi document nào.'],
    ['firestore.indexes.json',
     'Khai báo composite index (vd: userId + playedAt cho lịch sử).'],
    ['storage.rules',
     'Security Rules cho Firebase Storage (avatar lớn).'],
    ['.firebaserc',
     'Trỏ Firebase CLI tới project ID cụ thể.'],
    ['README.md', 'Tài liệu mô tả tổng quan project.'],
])

# ============================================================
# 3. lib/main.dart + firebase_options
# ============================================================
H1('3. Entry point — lib/main.dart & firebase_options.dart')
add_table([
    ['File', 'Chức năng chính'],
    ['lib/main.dart',
     'Hàm main() — gọi WidgetsFlutterBinding.ensureInitialized(), khởi tạo Firebase, '
     'khởi tạo NotificationService, AudioService, đăng ký các Provider, chạy '
     'MaterialApp với theme và locale động.'],
    ['lib/firebase_options.dart',
     'Auto-generate bởi flutterfire CLI — chứa apiKey, appId, projectId cho từng '
     'nền tảng (Android/iOS/Web/...).'],
])

# ============================================================
# 4. lib/config/
# ============================================================
H1('4. lib/config/ — Hằng số & Theme')
add_table([
    ['File', 'Chức năng chính'],
    ['constants.dart',
     'AppConstants: tên collection (usersCollection, gameResultsCollection), '
     'mã game (gameQuiz, gameMatching, ...), key SharedPreferences. '
     'LevelSystem: ngưỡng XP cho 10 level user. Hằng số chấm sao, ngưỡng accuracy.'],
    ['design_tokens.dart',
     'AppSpacing (sm/md/lg), AppRadius (pill, sm, md), AppShadow (soft/strong) — '
     'token thiết kế dùng đồng bộ trong toàn app.'],
    ['theme.dart',
     'AppColors (primary, gradientPurple, gradientBlue, ...), buildLightTheme(), '
     'buildDarkTheme() — định nghĩa Material ThemeData.'],
    ['rewards.dart',
     'LevelReward.byLevel(level) trả về phần thưởng coin/XP khi pass level lần đầu '
     '(L1=30/20, L2=60/40, L3=150/100). Định nghĩa milestone streak.'],
])

# ============================================================
# 5. lib/models/
# ============================================================
H1('5. lib/models/ — Data Models (POJO)')
add_table([
    ['File', 'Class', 'Chức năng chính'],
    ['user_model.dart', 'UserModel',
     'Đại diện 1 user: uid, email, displayName, totalScore/Coins/XP/level, streak, '
     'longestStreak, ownedPacks, progress map. Có fromFirestore(), toMap(), copyWith().'],
    ['game_result_model.dart', 'GameResultModel',
     '1 lần chơi xong: userId, gameType, level, score, correctAnswers, '
     'timeTakenSeconds, coinsEarned, xpEarned, playedAt. Computed: accuracy %, '
     'stars (0–3 dựa accuracy).'],
    ['vocab_model.dart', 'VocabModel',
     'Từ vựng đơn lẻ: id, word (English), meaning (tiếng Việt), level, topic, '
     'audioUrl (tuỳ chọn).'],
    ['vocab_pack_model.dart', 'VocabPackModel',
     'Gói từ vựng: id, name, description, level, price (coin), iconUrl, totalWords.'],
    ['level_model.dart', 'LevelModel',
     'Đại diện 1 level (1/2/3) trong pack: stars đã đạt, isUnlocked, isCompleted.'],
    ['level_reward_model.dart', 'LevelReward',
     'Phần thưởng pass level lần đầu: coinReward, xpReward, packCompleted (bool).'],
])

# ============================================================
# 6. lib/providers/
# ============================================================
H1('6. lib/providers/ — State Management')
add_table([
    ['File', 'Class', 'Chức năng chính'],
    ['user_provider.dart', 'UserProvider',
     'Lưu trữ UserModel hiện tại. Hàm chính: load(uid), updateLocalUser(addScore/Coins/XP), '
     'recordLevelComplete(gameType, packId, level) → trả LevelReward, applyStreakOutcome(), '
     'updateAvatar(file). Notify listeners để các Consumer rebuild.'],
    ['game_provider.dart', 'GameProvider',
     'Quản lý 1 phiên chơi game: startGame() set _startTime, addScore(d), addCorrect(), '
     'addWrong(), finishGame(...) → tính reward (coin/XP/bonus accuracy/speed/multiplier '
     'theo level), gọi FirestoreService.saveGameResult, trả GameOutcome.'],
    ['favorites_provider.dart', 'FavoritesProvider',
     'Set<String> các id pack/game đã favorite. Lưu local SharedPreferences. '
     'toggle(id), contains(id), clear().'],
    ['settings_provider.dart', 'SettingsProvider',
     'theme (ThemeMode), language (vi/en), soundEnabled, notificationEnabled, '
     'reminderHour/Minute. Lưu/load qua LocalStorage.'],
])

# ============================================================
# 7. lib/services/
# ============================================================
H1('7. lib/services/ — Service Layer (gọi Firebase, OS, asset)')
add_table([
    ['File', 'Class', 'Chức năng chính'],
    ['auth_service.dart', 'AuthService',
     'Wrapper quanh FirebaseAuth: signUp(email, password) — tạo tài khoản + user '
     'document trên Firestore với defaults (100 coin, 3 pack free); signIn(); '
     'signOut(); resetPassword(email); ensureUserDoc() — tự tạo doc nếu thiếu '
     '(Pigeon bug recovery).'],
    ['firestore_service.dart', 'FirestoreService',
     'Toàn bộ giao tiếp Cloud Firestore: getUser/streamUser, updateUser, '
     'updateDisplayName, updateAvatarUrl, updateLevelProgress (transaction), '
     'purchasePack (transaction atomic), saveGameResult (transaction lớn — lưu '
     'result + cộng coin/XP/score + streak + milestone), streamLeaderboard, '
     'streamUserHistory, seedLeaderboardDemoUsers.'],
    ['storage_service.dart', 'StorageService',
     'Wrapper Firebase Storage cho upload/download avatar lớn (hiện app dùng '
     'base64 trong Firestore, service dự phòng cho migration sau).'],
    ['audio_service.dart', 'AudioService',
     'Singleton phát sound effect: playClick, playCorrect, playWrong, '
     'playLevelUp. Đọc setting soundEnabled từ SettingsProvider.'],
    ['notification_service.dart', 'NotificationService',
     'flutter_local_notifications + timezone (Asia/HCM). scheduleDailyReminder(hour, '
     'minute), cancelAll(), requestPermission(). Auto-reschedule khi đổi múi giờ.'],
    ['json_service.dart', 'JsonService',
     'loadVocab(level) — đọc assets/data/vocab_<level>.json, parse thành List<VocabModel>. '
     'Cache trong RAM giữa các lần gọi.'],
    ['pack_service.dart', 'PackService',
     'Quản lý danh sách VocabPackModel: load metadata pack, mapping pack ↔ vocab.'],
    ['local_storage.dart', 'LocalStorage',
     'Wrapper SharedPreferences: setString/getString/setBool/getBool... + '
     'key constants. Dùng cho settings và favorites.'],
])

# ============================================================
# 8. lib/screens/
# ============================================================
H1('8. lib/screens/ — Màn hình UI')

H2('8.1. screens/splash_screen.dart')
add_table([
    ['File', 'Chức năng'],
    ['splash_screen.dart',
     'Splash với logo animation. Kiểm tra FirebaseAuth.currentUser — nếu có '
     'session thì điều hướng Home, không thì Login.'],
])

H2('8.2. screens/auth/ — Xác thực')
add_table([
    ['File', 'Chức năng'],
    ['login_screen.dart',
     'Form email + mật khẩu, nút "Đăng nhập", link "Quên mật khẩu" và "Đăng ký". '
     'Validator client-side trước khi gọi AuthService.signIn.'],
    ['register_screen.dart',
     'Form 4 trường (tên, email, mật khẩu, xác nhận) + checkbox điều khoản. '
     'Validator độ dài tên 2–30, email regex, mật khẩu ≥ 6 ký tự, confirm khớp. '
     'Gọi AuthService.signUp.'],
])

H2('8.3. screens/home/')
add_table([
    ['File', 'Chức năng'],
    ['home_screen.dart',
     'Dashboard chính: header (avatar/tên/level/streak/coin), banner Daily '
     'Challenge, grid tile dẫn tới Game Menu / Leaderboard / Shop / Profile / '
     'Favorites / History / Settings. Pull-to-refresh để re-fetch user.'],
])

H2('8.4. screens/games/ — 4 mini-game + flow')
add_table([
    ['File', 'Chức năng'],
    ['pack_selection_screen.dart',
     'Hiển thị danh sách VocabPack cho gameType vừa chọn. Phân biệt owned/locked, '
     'có nút Mua nếu chưa sở hữu (extend → Shop).'],
    ['level_map_screen.dart',
     'Chọn level 1/2/3 trong pack. Level chỉ unlock khi level trước đạt ≥ 2 sao '
     '(hoặc 3 sao cho Memory). Hiển thị star indicator đã đạt.'],
    ['game_menu_screen.dart',
     'Chọn 1 trong 4 mini-game (Quiz, Matching, Memory, Word Puzzle) và độ khó '
     '(beginner/intermediate/advanced).'],
    ['quiz_game.dart',
     'Trắc nghiệm 10 câu, 15s/câu. Có nút TTS phát âm từ tiếng Anh. Điểm = 10 + '
     'thời gian còn lại; sai = 0. Highlight xanh/đỏ feedback.'],
    ['matching_game.dart',
     'Nối từ ↔ nghĩa, 4 cặp/round × 3 round, 60s tổng. Điểm 10/cặp đúng lần đầu, '
     'sai báo đỏ → reset không cộng điểm.'],
    ['memory_game.dart',
     'Lật thẻ Pexeso. pairs/round 4-6-8 theo level, 90s. Match khi cùng vocabId + '
     '1 word + 1 meaning. Điểm scale max 100; sao theo timeLeft (≥50s = 3⭐).'],
    ['word_puzzle_game.dart',
     'Sắp xếp chữ cái thành từ. 5 từ/ván, không timer. Hint -10 coin tiết lộ 1 chữ. '
     '20 điểm/từ đúng.'],
    ['game_result_screen.dart',
     'Tổng kết ván chơi: animation sao, điểm, accuracy, coin/XP earn, streak '
     'milestone banner (nếu có), Level Reward (nếu pass level mới). Nút Chơi lại '
     '/ Về Home.'],
])

H2('8.5. screens/profile/')
add_table([
    ['File', 'Chức năng'],
    ['profile_screen.dart',
     'Hiển thị avatar, tên, email, stats theo từng loại game. Icon edit → sang '
     'EditProfile. Nút Đăng xuất.'],
    ['edit_profile_screen.dart',
     'Sửa displayName + đổi avatar (image_picker → resize → base64 ≤ 500KB → '
     'lưu Firestore). Nút Lưu cập nhật cả Provider lẫn Firestore.'],
])

H2('8.6. screens/leaderboard/, history/, favorites/, shop/, settings/')
add_table([
    ['File', 'Chức năng'],
    ['leaderboard/leaderboard_screen.dart',
     'Hai tab Score/XP, top 100 user. Highlight dòng user hiện tại. Stream từ '
     'Firestore (orderBy DESC).'],
    ['history/history_screen.dart',
     '50 game gần nhất của user. Mỗi dòng: gameType + level + điểm + sao + thời '
     'lượng + ngày chơi.'],
    ['favorites/favorites_screen.dart',
     'Danh sách đã favorite từ FavoritesProvider. Empty state nếu rỗng.'],
    ['shop/shop_screen.dart',
     'Danh sách pack có thể mua, giá coin. Nhấn Mua → gọi UserProvider.purchasePack '
     '(transaction). Đã sở hữu hiển thị "Đã mua".'],
    ['settings/settings_screen.dart',
     'Switch dark mode, dropdown ngôn ngữ VI/EN, switch sound, switch notification, '
     'time picker giờ nhắc. Lưu qua SettingsProvider → SharedPreferences.'],
])

# ============================================================
# 9. lib/widgets/
# ============================================================
H1('9. lib/widgets/ — Widget tái sử dụng')
add_table([
    ['File', 'Chức năng'],
    ['custom_button.dart',
     'Nút bo tròn gradient với loading state, dùng đồng nhất cho các CTA chính.'],
    ['user_avatar.dart',
     'Hiển thị avatar từ data URI base64 hoặc URL, fallback initials. Hỗ trợ '
     'border level color.'],
    ['vocab_card.dart',
     'Card hiển thị từ vựng (word + meaning) — dùng trong Quiz, Matching, history.'],
    ['loading_widget.dart',
     'Spinner toàn màn + message tuỳ biến — hiển thị khi đang load Firestore/asset.'],
    ['bubble_back_button.dart',
     'Nút Back tròn dạng bubble với icon, dùng đồng nhất các màn hình chi tiết.'],
])

# ============================================================
# 10. lib/utils/
# ============================================================
H1('10. lib/utils/ — Helper')
add_table([
    ['File', 'Chức năng'],
    ['validators.dart',
     'Hàm static validator cho Form: email regex, password tối thiểu 6 ký tự, '
     'displayName 2–30, confirm match.'],
    ['app_localizations.dart',
     'Lớp AppLocalizations: load JSON i18n từ assets, hàm extension '
     'context.t(key) trả text theo locale hiện tại. Hỗ trợ VI/EN.'],
    ['streak_calculator.dart',
     'StreakCalculator.calculate(prevStreak, lastPlayedDate, now) → trả '
     'StreakOutcome { newStreak, milestoneHit, bonusCoins, bonusXP, '
     'streakIncreased }. Quy tắc: same-day giữ nguyên; ngày kế tiếp +1; '
     'cách ≥ 1 ngày reset = 1.'],
    ['helpers.dart',
     'Hàm tiện ích chung: formatNumber, formatDuration, capitalize, '
     'showSnackbar wrapper, ...'],
])

# ============================================================
# 11. assets/
# ============================================================
H1('11. assets/ — Tài nguyên tĩnh')
add_table([
    ['Thư mục', 'Nội dung'],
    ['assets/audio/',
     'File mp3/ogg sound effect: click.mp3, correct.mp3, wrong.mp3, level_up.mp3.'],
    ['assets/data/',
     'JSON từ vựng theo cấp độ: vocab_beginner.json, vocab_intermediate.json, '
     'vocab_advanced.json. Mỗi file là array object {id, word, meaning, ...}.'],
    ['assets/images/',
     'Logo, icon từng game, illustration empty state, badge level, icon pack.'],
])

# ============================================================
# 12. Quan hệ giữa các tầng (PlantUML)
# ============================================================
H1('12. Sơ đồ quan hệ giữa các tầng (gói cấp cao)')
P('Sơ đồ tóm tắt cách các thư mục/tầng phụ thuộc vào nhau. Tầng cao chỉ phụ thuộc tầng '
  'thấp hơn — không có cycle.')

deps_uml = '''@startuml VocabQuest_Layers
skinparam linetype ortho
skinparam shadowing false
skinparam package {
  BackgroundColor #FFF8E7
  BorderColor #444444
}
title Phụ thuộc giữa các thư mục lib/

package "screens/"   as SCR
package "widgets/"   as WID
package "providers/" as PRV
package "services/"  as SVC
package "models/"    as MOD
package "utils/"     as UTL
package "config/"    as CFG

SCR --> WID
SCR --> PRV
SCR --> UTL
SCR --> CFG
WID --> MOD
WID --> CFG
PRV --> SVC
PRV --> MOD
PRV --> UTL
SVC --> MOD
SVC --> UTL
SVC --> CFG
MOD --> UTL
@enduml
'''
add_code_block(deps_uml)

# ============================================================
# 13. Thống kê & quy ước
# ============================================================
H1('13. Thống kê & quy ước đặt tên')
add_table([
    ['Hạng mục', 'Số lượng / Quy ước'],
    ['Tổng số file Dart trong lib/', '52 file'],
    ['Số screen', '20 screen (auth 2, home 1, games 8, profile 2, leaderboard/history/favorites/shop/settings/splash mỗi cái 1)'],
    ['Số provider', '4 (User, Game, Favorites, Settings)'],
    ['Số service', '8 (Auth, Firestore, Storage, Audio, Notification, Json, Pack, LocalStorage)'],
    ['Số model', '6 (User, GameResult, Vocab, VocabPack, Level, LevelReward)'],
    ['Số widget tái sử dụng', '5'],
    ['Quy ước tên file', 'snake_case (vd: game_result_screen.dart)'],
    ['Quy ước tên class', 'PascalCase (vd: GameResultScreen)'],
    ['Quy ước biến private', 'Bắt đầu bằng _ (vd: _score, _isLoading)'],
    ['Quy ước comment', 'Tiếng Việt, /// docstring cho service & game logic; '
     'không comment what, chỉ why khi cần'],
])

# ===== Save =====
try:
    d.save(path)
    print('Saved:', path)
except PermissionError:
    alt = path.replace('_v2.docx', '_v3.docx')
    d.save(alt)
    print('Locked. Saved to:', alt)
