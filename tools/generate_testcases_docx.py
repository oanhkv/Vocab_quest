# -*- coding: utf-8 -*-
"""
Sinh tai lieu Test Case cho ung dung VocabQuest (Flutter).
Xuat ra file .docx tai: C:\\Users\\Kieu Anh\\Desktop\\CD2\\VocabQuest_TestCases.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

OUT_PATH = r"C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TestCases.docx"

# ------------- HELPERS ---------------

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_borders(cell, color="000000", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def add_heading(doc, text, level=1, color=None, align=None):
    h = doc.add_heading(text, level=level)
    if align is not None:
        h.alignment = align
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h

def add_paragraph(doc, text, bold=False, size=11, italic=False, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    if color:
        r.font.color.rgb = color
    return p

def style_table_header(row, bg="305496", text_color=RGBColor(0xFF, 0xFF, 0xFF)):
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = text_color
                run.font.size = Pt(10.5)
                run.font.name = "Times New Roman"
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def fill_cell(cell, text, bold=False, size=10, align=None, color=None, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(str(text) if text is not None else "")
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    if color:
        r.font.color.rgb = color
    if bg:
        set_cell_bg(cell, bg)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])

def severity_color(sev):
    return {
        "Cao": RGBColor(0xC0, 0x00, 0x00),
        "Trung binh": RGBColor(0xBF, 0x90, 0x00),
        "Thap": RGBColor(0x00, 0x6D, 0x2C),
    }.get(sev, RGBColor(0x00, 0x00, 0x00))

def status_bg(status):
    # Mac dinh "Chua thuc thi" -> vang nhat
    return {
        "Pass": "C6EFCE",
        "Fail": "FFC7CE",
        "Chua thuc thi": "FFF2CC",
    }.get(status, "FFFFFF")

# ------------- NOI DUNG TEST CASES ---------------
# Cac cot: ID, Module, Ten test case, Muc tieu, Dieu kien tien quyet,
# Du lieu dau vao, Cac buoc thuc hien, Ket qua mong doi, Loai test,
# Muc do uu tien, Trang thai, Ghi chu

TESTCASES = [
    # ============ 1. DANG KY TAI KHOAN ============
    {
        "id": "TC_REG_001", "module": "Đăng ký",
        "name": "Đăng ký thành công với thông tin hợp lệ",
        "goal": "Xác minh người dùng có thể tạo tài khoản mới khi nhập đầy đủ thông tin hợp lệ.",
        "pre": "Ứng dụng đã được cài đặt; đang ở màn hình đăng nhập; có kết nối Internet; email chưa từng đăng ký.",
        "input": "Tên: 'Nguyễn Văn A'; Email: 'test123@gmail.com'; Mật khẩu: 'abc123'; Xác nhận: 'abc123'; Tích điều khoản.",
        "steps": "1. Mở màn hình đăng nhập\n2. Nhấn 'Đăng ký ngay'\n3. Nhập tên, email, mật khẩu, xác nhận mật khẩu\n4. Tích ô đồng ý điều khoản\n5. Nhấn nút 'Đăng ký'",
        "expect": "Tạo tài khoản thành công trên Firebase Auth; tài liệu user được tạo trong Firestore với 100 coins, 5 hearts, level 1; chuyển sang màn hình Home.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": "Kiểm tra document user trên Firestore có đủ các trường khởi tạo."
    },
    {
        "id": "TC_REG_002", "module": "Đăng ký",
        "name": "Đăng ký thất bại khi email đã tồn tại",
        "goal": "Đảm bảo hệ thống từ chối đăng ký khi email đã được sử dụng.",
        "pre": "Có kết nối Internet; email đã đăng ký trước đó.",
        "input": "Email: 'test123@gmail.com' (đã tồn tại); các trường còn lại hợp lệ.",
        "steps": "1. Mở màn hình đăng ký\n2. Nhập email đã tồn tại\n3. Nhập các trường hợp lệ còn lại\n4. Nhấn 'Đăng ký'",
        "expect": "Hiển thị thông báo lỗi tiếng Việt: 'Email đã được đăng ký'; không tạo tài khoản mới.",
        "type": "Âm tính", "priority": "Cao", "status": "Chưa thực thi",
        "note": "Firebase code: email-already-in-use."
    },
    {
        "id": "TC_REG_003", "module": "Đăng ký",
        "name": "Validate mật khẩu dưới 6 ký tự",
        "goal": "Kiểm tra ràng buộc độ dài tối thiểu của mật khẩu.",
        "pre": "Đang ở màn hình đăng ký.",
        "input": "Mật khẩu: '123' (3 ký tự); xác nhận: '123'.",
        "steps": "1. Nhập đầy đủ các trường\n2. Mật khẩu chỉ 3 ký tự\n3. Nhấn 'Đăng ký'",
        "expect": "Validator chặn submit; hiển thị lỗi 'Mật khẩu tối thiểu 6 ký tự'.",
        "type": "Giá trị biên", "priority": "Cao", "status": "Chưa thực thi",
        "note": "Firebase code dự phòng: weak-password."
    },
    {
        "id": "TC_REG_004", "module": "Đăng ký",
        "name": "Validate mật khẩu và xác nhận không khớp",
        "goal": "Đảm bảo trường 'Xác nhận mật khẩu' phải trùng khớp với mật khẩu.",
        "pre": "Đang ở màn hình đăng ký.",
        "input": "Mật khẩu: 'abc123'; Xác nhận: 'abc124'.",
        "steps": "1. Nhập mật khẩu hợp lệ\n2. Nhập xác nhận khác mật khẩu\n3. Nhấn 'Đăng ký'",
        "expect": "Hiển thị lỗi 'Mật khẩu xác nhận không trùng khớp'; nút Đăng ký không thực thi.",
        "type": "Xác thực", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_REG_005", "module": "Đăng ký",
        "name": "Validate định dạng email không hợp lệ",
        "goal": "Kiểm tra hệ thống chặn email sai định dạng.",
        "pre": "Đang ở màn hình đăng ký.",
        "input": "Email: 'abc.com' (thiếu @), 'abc@', 'abc@domain'.",
        "steps": "1. Lần lượt nhập các email không hợp lệ\n2. Nhấn 'Đăng ký'",
        "expect": "Validator hiển thị 'Email không hợp lệ' cho từng trường hợp; form không được submit.",
        "type": "Xác thực", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": "Regex email chuẩn."
    },
    {
        "id": "TC_REG_006", "module": "Đăng ký",
        "name": "Chưa tích điều khoản không được đăng ký",
        "goal": "Đảm bảo người dùng buộc phải chấp nhận điều khoản.",
        "pre": "Đang ở màn hình đăng ký; thông tin hợp lệ.",
        "input": "Mọi trường hợp lệ; không tích vào checkbox điều khoản.",
        "steps": "1. Nhập đầy đủ thông tin\n2. KHÔNG tích điều khoản\n3. Nhấn 'Đăng ký'",
        "expect": "Nút Đăng ký bị vô hiệu hoặc hiển thị yêu cầu tích điều khoản; không tạo tài khoản.",
        "type": "Xác thực", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_REG_007", "module": "Đăng ký",
        "name": "Tên hiển thị quá ngắn (<2) hoặc quá dài (>30)",
        "goal": "Kiểm tra giới hạn độ dài tên hiển thị.",
        "pre": "Đang ở màn hình đăng ký.",
        "input": "Tên: 'A' (1 ký tự); Tên: 'A' * 31.",
        "steps": "1. Nhập tên 1 ký tự, submit\n2. Quay lại nhập tên 31 ký tự, submit",
        "expect": "Cả hai lần đều hiển thị lỗi 'Tên phải từ 2 đến 30 ký tự'; không lưu được.",
        "type": "Giá trị biên", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_REG_008", "module": "Đăng ký",
        "name": "Mất mạng khi đăng ký",
        "goal": "Kiểm tra hành vi khi không có kết nối Internet.",
        "pre": "Tắt Wi-Fi và dữ liệu di động.",
        "input": "Thông tin đăng ký hợp lệ.",
        "steps": "1. Tắt mạng\n2. Nhập form hợp lệ\n3. Nhấn 'Đăng ký'",
        "expect": "Hiển thị thông báo lỗi mạng; không lock UI; người dùng có thể thử lại sau khi bật mạng.",
        "type": "Âm tính", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 2. DANG NHAP ============
    {
        "id": "TC_LOGIN_001", "module": "Đăng nhập",
        "name": "Đăng nhập thành công",
        "goal": "Xác minh đăng nhập với email và mật khẩu đúng.",
        "pre": "Tài khoản đã tồn tại; có Internet.",
        "input": "Email: 'test123@gmail.com'; Mật khẩu: 'abc123'.",
        "steps": "1. Mở app\n2. Nhập email, mật khẩu\n3. Nhấn 'Đăng nhập'",
        "expect": "Đăng nhập thành công; chuyển sang màn hình Home; UserProvider có dữ liệu user.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LOGIN_002", "module": "Đăng nhập",
        "name": "Đăng nhập với mật khẩu sai",
        "goal": "Kiểm tra phản hồi khi mật khẩu sai.",
        "pre": "Tài khoản đã tồn tại.",
        "input": "Email đúng; mật khẩu sai.",
        "steps": "1. Nhập email đúng\n2. Nhập mật khẩu sai\n3. Nhấn 'Đăng nhập'",
        "expect": "Hiển thị 'Mật khẩu không chính xác'; không chuyển màn hình.",
        "type": "Âm tính", "priority": "Cao", "status": "Chưa thực thi",
        "note": "Firebase code: wrong-password."
    },
    {
        "id": "TC_LOGIN_003", "module": "Đăng nhập",
        "name": "Đăng nhập với email không tồn tại",
        "goal": "Kiểm tra phản hồi khi tài khoản không tồn tại.",
        "pre": "Email chưa đăng ký bao giờ.",
        "input": "Email: 'khongco@mail.com'; Mật khẩu tùy ý.",
        "steps": "1. Nhập email chưa tồn tại\n2. Nhập mật khẩu\n3. Nhấn 'Đăng nhập'",
        "expect": "Hiển thị 'Người dùng không tồn tại'.",
        "type": "Âm tính", "priority": "Cao", "status": "Chưa thực thi",
        "note": "Firebase code: user-not-found."
    },
    {
        "id": "TC_LOGIN_004", "module": "Đăng nhập",
        "name": "Bật/tắt hiển thị mật khẩu",
        "goal": "Kiểm tra icon con mắt hiển thị/ẩn mật khẩu.",
        "pre": "Đang ở màn hình đăng nhập.",
        "input": "Mật khẩu bất kỳ.",
        "steps": "1. Nhập mật khẩu\n2. Nhấn icon con mắt\n3. Nhấn lại icon con mắt",
        "expect": "Lần 1 hiển thị mật khẩu dạng text; lần 2 ẩn thành dấu chấm.",
        "type": "Giao diện", "priority": "Thấp", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LOGIN_005", "module": "Đăng nhập",
        "name": "Nhấn 'Quên mật khẩu?' và gửi email reset",
        "goal": "Kiểm tra luồng reset mật khẩu.",
        "pre": "Tài khoản đã tồn tại.",
        "input": "Email hợp lệ: 'test123@gmail.com'.",
        "steps": "1. Nhấn 'Quên mật khẩu?'\n2. Nhập email\n3. Nhấn gửi",
        "expect": "Hiển thị thông báo 'Đã gửi email khôi phục'; email thực nhận được link reset từ Firebase.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LOGIN_006", "module": "Đăng nhập",
        "name": "Bỏ trống email và mật khẩu",
        "goal": "Kiểm tra ràng buộc bắt buộc.",
        "pre": "Đang ở màn hình đăng nhập.",
        "input": "Cả hai trường để trống.",
        "steps": "1. Không nhập gì\n2. Nhấn 'Đăng nhập'",
        "expect": "Validator báo lỗi 'Vui lòng nhập email' và 'Vui lòng nhập mật khẩu'.",
        "type": "Xác thực", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LOGIN_007", "module": "Đăng nhập",
        "name": "Nhấn nhanh nút Đăng nhập nhiều lần",
        "goal": "Đảm bảo không gửi trùng request khi user nhấn liên tiếp.",
        "pre": "Đã nhập thông tin hợp lệ.",
        "input": "Nhấn nút Đăng nhập 5 lần liên tiếp.",
        "steps": "1. Nhập thông tin\n2. Nhấn 'Đăng nhập' 5 lần liên tục",
        "expect": "Chỉ có 1 request đăng nhập được thực thi; loading indicator hiển thị.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 3. MAN HINH HOME ============
    {
        "id": "TC_HOME_001", "module": "Màn hình Home",
        "name": "Hiển thị avatar, tên, level, streak, coin đúng",
        "goal": "Kiểm tra header hiển thị đầy đủ thông tin người dùng.",
        "pre": "Đã đăng nhập.",
        "input": "User có: displayName='An', level=3, streak=7, coins=250.",
        "steps": "1. Đăng nhập\n2. Quan sát header Home",
        "expect": "Avatar hiển thị; 'An' đúng tên; badge 'Lv3'; chip streak '7'; coin '250'.",
        "type": "Giao diện", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_HOME_002", "module": "Màn hình Home",
        "name": "Thanh tiến độ XP phản ánh đúng tỉ lệ",
        "goal": "Kiểm tra thanh XP tới level kế tiếp chính xác.",
        "pre": "User có totalXP=300, level=3 (ngưỡng 250 -> 500).",
        "input": "XP=300, ngưỡng kế tiếp=500.",
        "steps": "1. Quan sát thanh XP ở Home",
        "expect": "Thanh đầy khoảng 20% = (300-250)/(500-250).",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_HOME_003", "module": "Màn hình Home",
        "name": "Pull-to-refresh cập nhật dữ liệu",
        "goal": "Đảm bảo kéo xuống làm mới dữ liệu user từ Firestore.",
        "pre": "Đang ở Home.",
        "input": "Thay đổi dữ liệu user trực tiếp trên Firestore.",
        "steps": "1. Sửa coins trên Firestore\n2. Kéo xuống tại Home",
        "expect": "Coins trên Home cập nhật theo giá trị mới.",
        "type": "Tích hợp", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_HOME_004", "module": "Màn hình Home",
        "name": "Nhấn tile 'Chơi game' điều hướng đúng",
        "goal": "Kiểm tra điều hướng sang màn hình menu game.",
        "pre": "Đang ở Home.",
        "input": "Nhấn tile 'Chơi game'.",
        "steps": "1. Nhấn tile 'Chơi game'",
        "expect": "Chuyển sang GameMenuScreen, hiển thị 4 loại game.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_HOME_005", "module": "Màn hình Home",
        "name": "Tile 'Xếp hạng' mở leaderboard",
        "goal": "Kiểm tra điều hướng leaderboard.",
        "pre": "Đang ở Home.",
        "input": "Nhấn tile 'Xếp hạng'.",
        "steps": "1. Nhấn tile 'Xếp hạng'",
        "expect": "Chuyển sang màn hình Leaderboard; hiển thị top 100.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_HOME_006", "module": "Màn hình Home",
        "name": "Tile 'Cửa hàng' mở Shop khi nhấn vào coins",
        "goal": "Kiểm tra điều hướng tới Shop qua icon coin.",
        "pre": "Đang ở Home.",
        "input": "Nhấn vào khu vực coin.",
        "steps": "1. Nhấn label coins",
        "expect": "Chuyển sang Shop; không crash.",
        "type": "Chức năng", "priority": "Thấp", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_HOME_007", "module": "Màn hình Home",
        "name": "Banner Daily Challenge cập nhật XP hôm nay",
        "goal": "Kiểm tra banner thử thách ngày.",
        "pre": "Đã chơi hôm nay, earn 20 XP.",
        "input": "xpToday=20, target=50.",
        "steps": "1. Quan sát banner",
        "expect": "Thanh hiển thị 40%; thông báo 'Nhận 30 XP nữa để hoàn thành'.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 4. GAME QUIZ ============
    {
        "id": "TC_QUIZ_001", "module": "Game Quiz",
        "name": "Chơi đủ 10 câu và kết thúc game",
        "goal": "Kiểm tra luồng cơ bản của Quiz.",
        "pre": "Đã chọn pack + level; có từ vựng.",
        "input": "Trả lời đúng 7/10 câu.",
        "steps": "1. Vào Quiz\n2. Trả lời hết 10 câu\n3. Xem kết quả",
        "expect": "Hiển thị GameResult với 7 câu đúng; accuracy=70%; 2 sao; lưu vào game_results.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": "Ngưỡng sao: 90%=3★, 60%=2★, 30%=1★."
    },
    {
        "id": "TC_QUIZ_002", "module": "Game Quiz",
        "name": "Hết 15 giây auto chuyển câu",
        "goal": "Kiểm tra timer tự chuyển câu sau 15 giây.",
        "pre": "Đang chơi Quiz.",
        "input": "Không chọn đáp án cho 1 câu.",
        "steps": "1. Chờ timer đếm ngược 15s\n2. Quan sát",
        "expect": "Câu hỏi tự chuyển; câu không trả lời tính là sai; điểm không cộng.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_QUIZ_003", "module": "Game Quiz",
        "name": "Điểm bonus theo thời gian còn lại",
        "goal": "Kiểm tra công thức 10 + remainingSeconds.",
        "pre": "Đang chơi Quiz.",
        "input": "Trả lời đúng khi còn 12 giây.",
        "steps": "1. Đọc câu hỏi\n2. Nhấn đáp án đúng khi timer còn 12s",
        "expect": "Điểm của câu = 10 + 12 = 22 điểm.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_QUIZ_004", "module": "Game Quiz",
        "name": "Nhấn nút phát âm (TTS)",
        "goal": "Kiểm tra nút speaker phát âm từ vựng.",
        "pre": "Đang trong một câu hỏi Quiz; loa bật.",
        "input": "Nhấn icon loa.",
        "steps": "1. Bấm nút loa ở câu hỏi",
        "expect": "Phát âm đúng từ tiếng Anh đang hiển thị.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_QUIZ_005", "module": "Game Quiz",
        "name": "Bấm back giữa game hiển thị xác nhận",
        "goal": "Kiểm tra PopScope ngăn thoát game đột ngột.",
        "pre": "Đang ở câu thứ 5/10.",
        "input": "Nhấn back device/button.",
        "steps": "1. Nhấn back",
        "expect": "Hiện dialog xác nhận 'Bạn có chắc muốn thoát?'; chọn Huỷ giữ nguyên game.",
        "type": "Giao diện", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_QUIZ_006", "module": "Game Quiz",
        "name": "Đáp án highlight xanh khi đúng, đỏ khi sai",
        "goal": "Kiểm tra feedback trực quan.",
        "pre": "Đang chơi Quiz.",
        "input": "Chọn đáp án đúng; sau đó chọn đáp án sai câu khác.",
        "steps": "1. Chọn đúng\n2. Chọn sai câu khác",
        "expect": "Câu đúng highlight xanh; câu sai highlight đỏ trong ~1.5s trước khi auto-advance.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 5. GAME MATCHING ============
    {
        "id": "TC_MATCH_001", "module": "Game Matching",
        "name": "Hoàn thành đủ 3 round với 12 cặp",
        "goal": "Kiểm tra luồng hoàn thành game Matching.",
        "pre": "Có đủ từ vựng; đang trong Matching.",
        "input": "Ghép đúng tất cả 12 cặp trong 60 giây.",
        "steps": "1. Bắt đầu game\n2. Ghép đúng 4 cặp/round × 3 round",
        "expect": "Điểm tối đa 120; hiện kết quả 3 sao; lưu vào Firestore.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_MATCH_002", "module": "Game Matching",
        "name": "Ghép sai hiển thị màu đỏ rồi reset",
        "goal": "Kiểm tra feedback ghép sai.",
        "pre": "Đang chơi Matching.",
        "input": "Chọn 1 từ trái và 1 nghĩa không tương ứng.",
        "steps": "1. Nhấn từ A (trái)\n2. Nhấn nghĩa B (phải) sai",
        "expect": "Hai thẻ highlight đỏ 600ms rồi trở về trạng thái chưa chọn; không cộng điểm.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_MATCH_003", "module": "Game Matching",
        "name": "Chỉ cộng điểm lần ghép đúng đầu tiên",
        "goal": "Kiểm tra quy tắc chấm điểm 'first-attempt only'.",
        "pre": "Đang trong Matching.",
        "input": "Ghép sai cặp X -> sau đó ghép đúng cặp X.",
        "steps": "1. Ghép sai cặp X\n2. Ghép đúng cặp X",
        "expect": "Cặp X được đánh dấu matched (biến mất) nhưng KHÔNG cộng 10 điểm.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_MATCH_004", "module": "Game Matching",
        "name": "Hết 60 giây game kết thúc",
        "goal": "Kiểm tra timer tổng 60 giây.",
        "pre": "Đang chơi Matching.",
        "input": "Không thao tác cho tới khi hết giờ.",
        "steps": "1. Chờ timer đếm về 0",
        "expect": "Game dừng tại thời điểm 0; hiện kết quả; điểm theo số cặp đã ghép đúng.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_MATCH_005", "module": "Game Matching",
        "name": "Các round không lặp từ",
        "goal": "Kiểm tra deduplication 12 từ độc nhất.",
        "pre": "Pack có >=12 từ.",
        "input": "Chơi đủ 3 round.",
        "steps": "1. Ghi nhận các từ xuất hiện round 1, 2, 3",
        "expect": "Không từ nào xuất hiện trong 2 round khác nhau.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 6. GAME MEMORY ============
    {
        "id": "TC_MEM_001", "module": "Game Memory",
        "name": "Ghép đủ 3 round × 6 cặp",
        "goal": "Kiểm tra luồng hoàn chỉnh Memory.",
        "pre": "Đang ở Memory game.",
        "input": "Ghép đúng 18 cặp trong 90 giây.",
        "steps": "1. Lật thẻ tìm cặp\n2. Hoàn thành 3 round",
        "expect": "Tổng điểm = 18 × 20 = 360; 3 sao; kết quả lưu Firestore.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_MEM_002", "module": "Game Memory",
        "name": "Lật 2 thẻ khác vocabId sẽ lật úp lại sau 800ms",
        "goal": "Kiểm tra cơ chế lật úp khi không khớp.",
        "pre": "Đang chơi Memory.",
        "input": "Lật 2 thẻ không khớp.",
        "steps": "1. Lật thẻ A (EN)\n2. Lật thẻ B (VI) vocabId khác",
        "expect": "Sau 800ms, 2 thẻ lật úp trở lại; không có cặp matched.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_MEM_003", "module": "Game Memory",
        "name": "Không thể lật 3 thẻ cùng lúc",
        "goal": "Kiểm tra ràng buộc 2 thẻ lật đồng thời.",
        "pre": "Đang chơi Memory.",
        "input": "Nhấn liên tiếp 3 thẻ khác nhau rất nhanh.",
        "steps": "1. Tap 3 thẻ liên tiếp",
        "expect": "Chỉ có 2 thẻ đầu được lật; thẻ thứ 3 không có tác dụng cho tới khi cặp được xử lý.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_MEM_004", "module": "Game Memory",
        "name": "Hiệu ứng lật 3D hoạt động mượt",
        "goal": "Kiểm tra animation Matrix4.rotateY.",
        "pre": "Thiết bị có hiệu năng trung bình trở lên.",
        "input": "Lật nhiều thẻ liên tiếp.",
        "steps": "1. Chơi 1 round\n2. Quan sát animation",
        "expect": "Animation mượt, không giật; hiển thị badge EN/VI ở mặt sau đúng.",
        "type": "Hiệu năng", "priority": "Thấp", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 7. WORD PUZZLE ============
    {
        "id": "TC_PUZ_001", "module": "Game Word Puzzle",
        "name": "Hoàn thành 5 từ đúng",
        "goal": "Kiểm tra luồng chính Word Puzzle.",
        "pre": "Có 5 từ vocab.",
        "input": "Ghép đúng 5 từ, mỗi từ trong 10 giây.",
        "steps": "1. Sắp xếp chữ cái đúng\n2. Lặp cho 5 từ",
        "expect": "Điểm = 100; 3 sao; kết quả lưu.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PUZ_002", "module": "Game Word Puzzle",
        "name": "Nút Hint trừ 10 coin và điền 1 chữ đúng",
        "goal": "Kiểm tra logic gợi ý.",
        "pre": "User có >=10 coins; đang giải 1 từ.",
        "input": "Nhấn Hint khi còn 4 ô trống.",
        "steps": "1. Nhấn Hint",
        "expect": "Coins giảm 10; 1 ô trống được điền chữ đúng; Hint không fill lại vị trí đã đúng.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PUZ_003", "module": "Game Word Puzzle",
        "name": "Hint bị vô hiệu khi coin < 10",
        "goal": "Đảm bảo không cho phép dùng hint thiếu coin.",
        "pre": "User có 5 coins.",
        "input": "Nhấn Hint.",
        "steps": "1. Nhấn Hint",
        "expect": "Hiển thị thông báo 'Không đủ xu'; coin không đổi; không có chữ được điền.",
        "type": "Âm tính", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PUZ_004", "module": "Game Word Puzzle",
        "name": "Nút Clear reset các chữ đang chọn",
        "goal": "Kiểm tra nút Clear.",
        "pre": "Đã chọn một vài chữ.",
        "input": "Nhấn Clear.",
        "steps": "1. Chọn 3 chữ\n2. Nhấn Clear",
        "expect": "Tất cả ô trống; chữ trở lại letter bank.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PUZ_005", "module": "Game Word Puzzle",
        "name": "Hết 10 giây tự chuyển từ và hiển thị đáp án",
        "goal": "Kiểm tra auto-advance khi timeout.",
        "pre": "Đang giải 1 từ nhưng chưa hoàn thành.",
        "input": "Chờ timer về 0.",
        "steps": "1. Không làm gì 10 giây",
        "expect": "Từ tự chuyển; hiển thị đáp án đúng; không cộng điểm cho từ đó.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 8. CHON PACK / LEVEL MAP / GAME MENU ============
    {
        "id": "TC_PACK_001", "module": "Pack/Level Map",
        "name": "Hiển thị danh sách pack từ vựng",
        "goal": "Kiểm tra màn hình pack hiển thị các pack có sẵn.",
        "pre": "Đã đăng nhập; có ít nhất 1 pack.",
        "input": "-",
        "steps": "1. Vào Pack Selection",
        "expect": "Hiển thị danh sách pack với ảnh, tên, giá, trạng thái owned/locked.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PACK_002", "module": "Pack/Level Map",
        "name": "Mua pack khi đủ coin",
        "goal": "Kiểm tra purchase pack thành công.",
        "pre": "User có 500 coins; pack giá 300.",
        "input": "Nhấn 'Mua' pack.",
        "steps": "1. Nhấn pack cần mua\n2. Nhấn 'Mua'",
        "expect": "Coin giảm 300; pack thêm vào ownedPacks; user có thể chơi pack này.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": "FirestoreService.purchasePack dùng transaction."
    },
    {
        "id": "TC_PACK_003", "module": "Pack/Level Map",
        "name": "Không mua được pack khi thiếu coin",
        "goal": "Đảm bảo block khi không đủ coin.",
        "pre": "User có 100 coins; pack giá 300.",
        "input": "Nhấn 'Mua'.",
        "steps": "1. Nhấn 'Mua'",
        "expect": "Hiển thị 'Không đủ xu'; coin không đổi.",
        "type": "Âm tính", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PACK_004", "module": "Pack/Level Map",
        "name": "Level map mở khoá dần theo tiến độ",
        "goal": "Kiểm tra level tiếp theo chỉ mở khi hoàn thành level trước với >=2 sao.",
        "pre": "Level 1 đã đạt 2 sao.",
        "input": "Xem Level Map.",
        "steps": "1. Hoàn thành level 1 2★\n2. Xem level 2",
        "expect": "Level 2 mở khoá; các level xa hơn vẫn khoá.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_MENU_001", "module": "Game Menu",
        "name": "Chọn độ khó trước khi chơi",
        "goal": "Kiểm tra chọn beginner/intermediate/advanced.",
        "pre": "Đang ở Game Menu.",
        "input": "Chọn 'intermediate' + game Quiz.",
        "steps": "1. Chọn độ khó\n2. Nhấn Play",
        "expect": "Game load đúng pack intermediate; số từ theo độ khó.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 9. KET QUA GAME ============
    {
        "id": "TC_RES_001", "module": "Kết quả game",
        "name": "Hiển thị điểm, số sao, coin/XP kiếm được",
        "goal": "Kiểm tra trang kết quả hiển thị đủ.",
        "pre": "Vừa hoàn thành một game.",
        "input": "Score=80/100, correct=4/5 (word puzzle).",
        "steps": "1. Kết thúc game\n2. Quan sát màn hình kết quả",
        "expect": "Điểm 80; accuracy 80%; 2 sao; coin/XP tương ứng hiển thị.",
        "type": "Giao diện", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_RES_002", "module": "Kết quả game",
        "name": "Nút 'Chơi lại' bắt đầu game mới cùng độ khó",
        "goal": "Kiểm tra tính năng chơi lại.",
        "pre": "Đang ở Result.",
        "input": "Nhấn 'Chơi lại'.",
        "steps": "1. Nhấn 'Chơi lại'",
        "expect": "Chuyển sang game mới cùng loại, độ khó như lần trước.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_RES_003", "module": "Kết quả game",
        "name": "Hoàn thành level >=2 sao ghi tiến độ",
        "goal": "Đảm bảo progress.{gameType}|{packId} cập nhật.",
        "pre": "Hoàn thành level 3 với 2 sao.",
        "input": "-",
        "steps": "1. Hoàn thành game 2★\n2. Xem Firestore user.progress",
        "expect": "Trường progress['quiz|beginner']=3 (hoặc lớn hơn giá trị cũ).",
        "type": "Tích hợp", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_RES_004", "module": "Kết quả game",
        "name": "Hoàn thành 1 sao KHÔNG ghi level progress",
        "goal": "Đảm bảo chỉ ghi level khi >=2 sao.",
        "pre": "Hoàn thành level 3 với 1 sao.",
        "input": "-",
        "steps": "1. Hoàn thành game 1★\n2. Xem progress",
        "expect": "progress['quiz|beginner'] không đổi (vẫn là giá trị cũ).",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 10. PROFILE ============
    {
        "id": "TC_PROF_001", "module": "Hồ sơ",
        "name": "Hiển thị đầy đủ thông tin hồ sơ",
        "goal": "Kiểm tra profile hiển thị tên, email, avatar, stats.",
        "pre": "Đã đăng nhập.",
        "input": "-",
        "steps": "1. Vào Profile",
        "expect": "Tên, email, avatar, score, coin, streak, level, và số game đã chơi theo loại hiển thị chính xác.",
        "type": "Giao diện", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PROF_002", "module": "Hồ sơ",
        "name": "Nhấn biểu tượng edit vào màn Edit Profile",
        "goal": "Kiểm tra điều hướng sang edit.",
        "pre": "Đang ở Profile.",
        "input": "-",
        "steps": "1. Nhấn icon edit",
        "expect": "Chuyển sang màn EditProfileScreen.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PROF_003", "module": "Chỉnh sửa hồ sơ",
        "name": "Cập nhật tên hiển thị thành công",
        "goal": "Kiểm tra đổi displayName.",
        "pre": "Đang ở Edit Profile.",
        "input": "Tên mới: 'Nguyễn Văn B'.",
        "steps": "1. Sửa tên\n2. Nhấn Lưu",
        "expect": "displayName cập nhật trên Firestore và UI Home/Profile.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PROF_004", "module": "Chỉnh sửa hồ sơ",
        "name": "Upload avatar nhỏ hơn 500KB",
        "goal": "Kiểm tra đổi avatar thành công.",
        "pre": "Đang ở Edit Profile.",
        "input": "Ảnh ~200KB base64.",
        "steps": "1. Chọn ảnh từ gallery\n2. Lưu",
        "expect": "Avatar cập nhật trên tất cả màn; avatarUrl lưu base64 trên Firestore.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PROF_005", "module": "Chỉnh sửa hồ sơ",
        "name": "Từ chối avatar > 500KB",
        "goal": "Đảm bảo giới hạn kích thước ảnh.",
        "pre": "Đang ở Edit Profile.",
        "input": "Ảnh 1MB.",
        "steps": "1. Chọn ảnh 1MB\n2. Lưu",
        "expect": "Hiển thị lỗi 'Ảnh quá lớn'; không upload.",
        "type": "Âm tính", "priority": "Cao", "status": "Chưa thực thi",
        "note": "UserProvider _maxAvatarBytes=500KB."
    },
    {
        "id": "TC_PROF_006", "module": "Chỉnh sửa hồ sơ",
        "name": "Từ chối cấp quyền camera/gallery",
        "goal": "Kiểm tra luồng khi từ chối permission.",
        "pre": "Đang ở image picker.",
        "input": "User từ chối quyền.",
        "steps": "1. Nhấn chọn ảnh\n2. Từ chối quyền",
        "expect": "Hiển thị thông báo cần cấp quyền; không crash.",
        "type": "Âm tính", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 11. FAVORITES ============
    {
        "id": "TC_FAV_001", "module": "Yêu thích",
        "name": "Hiển thị danh sách đã lưu",
        "goal": "Kiểm tra màn Favorites.",
        "pre": "Đã có mục favorite trong UserModel.",
        "input": "-",
        "steps": "1. Vào Favorites từ Home",
        "expect": "Hiển thị danh sách pack/game đã lưu.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_FAV_002", "module": "Yêu thích",
        "name": "Thêm/bỏ yêu thích đồng bộ với provider",
        "goal": "Kiểm tra cập nhật favorites_provider.",
        "pre": "Đang ở màn Pack Selection.",
        "input": "Nhấn icon tim 1 lần (add), nhấn lần 2 (remove).",
        "steps": "1. Bấm tim thêm\n2. Bấm tim bỏ",
        "expect": "Icon đổi trạng thái ngay; Favorites screen phản ánh thay đổi sau khi quay lại.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_FAV_003", "module": "Yêu thích",
        "name": "Empty state khi chưa có favorite",
        "goal": "Kiểm tra hiển thị khi danh sách rỗng.",
        "pre": "Favorites rỗng.",
        "input": "-",
        "steps": "1. Vào Favorites",
        "expect": "Hiển thị thông báo 'Chưa có mục yêu thích'.",
        "type": "Giao diện", "priority": "Thấp", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 12. LEADERBOARD ============
    {
        "id": "TC_LEAD_001", "module": "Bảng xếp hạng",
        "name": "Hiển thị top 100 theo điểm số",
        "goal": "Kiểm tra streamLeaderboard.",
        "pre": "Có >=10 user demo.",
        "input": "-",
        "steps": "1. Vào Leaderboard tab Score",
        "expect": "Hiển thị danh sách tối đa 100 user sắp xếp giảm dần theo totalScore.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LEAD_002", "module": "Bảng xếp hạng",
        "name": "Hiển thị đúng hạng cá nhân",
        "goal": "Kiểm tra getUserRankByScore.",
        "pre": "User đứng hạng 7.",
        "input": "-",
        "steps": "1. Vào Leaderboard",
        "expect": "Hiện 'Hạng của bạn: #7' hoặc tương đương.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LEAD_003", "module": "Bảng xếp hạng",
        "name": "Highlight user hiện tại trong danh sách",
        "goal": "User hiện tại phải có style nổi bật.",
        "pre": "User nằm trong top 100.",
        "input": "-",
        "steps": "1. Scroll tới dòng của mình",
        "expect": "Dòng có màu nền/viền khác biệt.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LEAD_004", "module": "Bảng xếp hạng",
        "name": "Chuyển đổi tab Score ↔ XP",
        "goal": "Kiểm tra chuyển giữa 2 loại xếp hạng.",
        "pre": "-",
        "input": "Tab XP.",
        "steps": "1. Nhấn tab XP",
        "expect": "Danh sách sắp xếp lại theo totalXP giảm dần.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LEAD_005", "module": "Bảng xếp hạng",
        "name": "Mất mạng hiển thị cache/thông báo",
        "goal": "Kiểm tra xử lý offline.",
        "pre": "Tắt mạng.",
        "input": "-",
        "steps": "1. Tắt mạng\n2. Vào Leaderboard",
        "expect": "Hiển thị cache (nếu có) hoặc thông báo 'Không có kết nối'.",
        "type": "Âm tính", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 13. LICH SU ============
    {
        "id": "TC_HIST_001", "module": "Lịch sử",
        "name": "Danh sách game sắp theo thời gian",
        "goal": "Kiểm tra streamUserHistory.",
        "pre": "Đã chơi >=5 game.",
        "input": "-",
        "steps": "1. Vào Lịch sử",
        "expect": "Hiện các game sắp xếp mới nhất trên cùng; tối đa 50 mục.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_HIST_002", "module": "Lịch sử",
        "name": "Hiển thị đúng thời gian, điểm, sao",
        "goal": "Kiểm tra dữ liệu từng dòng.",
        "pre": "Có game result.",
        "input": "-",
        "steps": "1. Xem 1 dòng bất kỳ",
        "expect": "Đúng loại game, độ khó, điểm, số câu đúng/tổng, thời lượng, sao.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_HIST_003", "module": "Lịch sử",
        "name": "Empty state khi chưa có lịch sử",
        "goal": "Kiểm tra màn hình khi chưa chơi game nào.",
        "pre": "User mới đăng ký.",
        "input": "-",
        "steps": "1. Vào Lịch sử",
        "expect": "Hiển thị 'Bạn chưa chơi game nào'.",
        "type": "Giao diện", "priority": "Thấp", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 14. CAI DAT ============
    {
        "id": "TC_SET_001", "module": "Cài đặt",
        "name": "Bật/tắt Dark Mode áp dụng ngay",
        "goal": "Kiểm tra toggle theme.",
        "pre": "Đang ở Settings.",
        "input": "Bật Dark Mode.",
        "steps": "1. Toggle Dark Mode",
        "expect": "Toàn bộ app chuyển màu nền tối ngay lập tức; lưu vào SharedPreferences.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_SET_002", "module": "Cài đặt",
        "name": "Đổi ngôn ngữ Vi <-> En",
        "goal": "Kiểm tra localization.",
        "pre": "App đang chạy.",
        "input": "Chọn English.",
        "steps": "1. Chọn English\n2. Quan sát",
        "expect": "Tất cả text chuyển sang English; lưu key 'language'='en'.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_SET_003", "module": "Cài đặt",
        "name": "Bật/tắt âm thanh",
        "goal": "Kiểm tra soundEnabled.",
        "pre": "App đang phát âm thanh.",
        "input": "Tắt Sound.",
        "steps": "1. Tắt sound\n2. Thử phát âm từ trong Quiz",
        "expect": "Không phát âm; các hiệu ứng âm khác im lặng.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_SET_004", "module": "Cài đặt",
        "name": "Đặt thời gian nhắc nhở",
        "goal": "Kiểm tra time picker.",
        "pre": "Notifications bật.",
        "input": "20:30.",
        "steps": "1. Chọn giờ 20:30\n2. Lưu",
        "expect": "reminderHour=20, reminderMinute=30; lên lịch notification.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_SET_005", "module": "Cài đặt",
        "name": "Cài đặt giữ nguyên sau khi restart app",
        "goal": "Kiểm tra persistence.",
        "pre": "Đã bật Dark mode, En, sound off.",
        "input": "-",
        "steps": "1. Kill app\n2. Mở lại",
        "expect": "Các cài đặt vẫn giữ nguyên như trước khi kill.",
        "type": "Tích hợp", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_SET_006", "module": "Cài đặt",
        "name": "Từ chối quyền notification",
        "goal": "Kiểm tra khi user từ chối quyền thông báo.",
        "pre": "Lần đầu bật notification.",
        "input": "Từ chối permission.",
        "steps": "1. Bật toggle notification\n2. Từ chối quyền hệ thống",
        "expect": "Toggle tự tắt lại; hiện thông báo hướng dẫn vào Settings OS để cấp quyền.",
        "type": "Âm tính", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 15. STREAK ============
    {
        "id": "TC_STK_001", "module": "Streak",
        "name": "Chơi cùng ngày streak không tăng",
        "goal": "Kiểm tra logic streak same-day.",
        "pre": "streak=5, lastPlayedAt=hôm nay 9:00.",
        "input": "Chơi thêm 1 game lúc 15:00 cùng ngày.",
        "steps": "1. Chơi xong 1 game",
        "expect": "streak vẫn =5; longestStreak không đổi.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_STK_002", "module": "Streak",
        "name": "Chơi ngày kế tiếp streak +1",
        "goal": "Kiểm tra tăng streak liên tục.",
        "pre": "streak=5, lastPlayedAt=hôm qua.",
        "input": "Chơi 1 game hôm nay.",
        "steps": "1. Chơi game",
        "expect": "streak=6; longestStreak >= 6.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_STK_003", "module": "Streak",
        "name": "Chơi cách 2 ngày streak reset về 1",
        "goal": "Kiểm tra reset khi bỏ lỡ ngày.",
        "pre": "streak=10, lastPlayedAt=3 ngày trước.",
        "input": "Chơi hôm nay.",
        "steps": "1. Chơi game",
        "expect": "streak=1; longestStreak vẫn =10 (không giảm).",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_STK_004", "module": "Streak",
        "name": "Milestone 7 ngày: +100 coin, +75 XP",
        "goal": "Kiểm tra phần thưởng mốc 7.",
        "pre": "streak=6 hôm qua.",
        "input": "Chơi 1 game hôm nay (→streak=7).",
        "steps": "1. Chơi game\n2. Xem thông báo phần thưởng",
        "expect": "Thông báo 'Chuỗi 7 ngày'; coin +100, XP +75 ngoài phần thưởng game thông thường.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": "Ngưỡng: 3/7/14/30/60/100."
    },
    {
        "id": "TC_STK_005", "module": "Streak",
        "name": "Không nhận milestone 2 lần cùng một ngày",
        "goal": "Đảm bảo streakIncreased=false khi chơi lần 2 cùng ngày.",
        "pre": "Vừa đạt streak=7 hôm nay (đã nhận milestone).",
        "input": "Chơi thêm 1 game.",
        "steps": "1. Chơi game thứ 2 trong ngày",
        "expect": "Không nhận thêm milestone; chỉ cộng điểm game thường.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 16. LEVEL SYSTEM ============
    {
        "id": "TC_LVL_001", "module": "Level System",
        "name": "Tăng level khi đủ XP ngưỡng",
        "goal": "Kiểm tra getLevelFromXP.",
        "pre": "level=2, totalXP=245 (ngưỡng lvl3=250).",
        "input": "Earn 10 XP.",
        "steps": "1. Hoàn thành game +10XP",
        "expect": "level chuyển từ 2 lên 3; badge Home và Profile cập nhật.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_LVL_002", "module": "Level System",
        "name": "Level tối đa = 10",
        "goal": "Đảm bảo level cap.",
        "pre": "totalXP=15000 (>12000).",
        "input": "-",
        "steps": "1. Xem Profile",
        "expect": "level=10; thanh XP đã đạt tối đa.",
        "type": "Giá trị biên", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 17. SHOP ============
    {
        "id": "TC_SHOP_001", "module": "Cửa hàng",
        "name": "Hiển thị danh sách vật phẩm",
        "goal": "Kiểm tra Shop.",
        "pre": "Đã đăng nhập.",
        "input": "-",
        "steps": "1. Vào Shop",
        "expect": "Hiển thị danh sách pack/item với giá.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_SHOP_002", "module": "Cửa hàng",
        "name": "Mua cùng một pack lần 2 bị chặn",
        "goal": "Tránh mua trùng.",
        "pre": "ownedPacks chứa 'pack1'.",
        "input": "Nhấn mua lại pack1.",
        "steps": "1. Tìm pack1\n2. Nhấn 'Mua'",
        "expect": "Hiển thị 'Bạn đã sở hữu pack này'; coin không đổi.",
        "type": "Âm tính", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 18. THONG BAO ============
    {
        "id": "TC_NOTI_001", "module": "Thông báo",
        "name": "Nhận notification đúng giờ đã đặt",
        "goal": "Kiểm tra notification service.",
        "pre": "Đặt giờ nhắc 20:30; notifications bật.",
        "input": "-",
        "steps": "1. Chờ đến 20:30",
        "expect": "Nhận được push notification với nội dung nhắc học từ vựng.",
        "type": "Tích hợp", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": "Test trên thiết bị thật để chính xác."
    },
    {
        "id": "TC_NOTI_002", "module": "Thông báo",
        "name": "Tắt thông báo không nhận push",
        "goal": "Đảm bảo tôn trọng setting.",
        "pre": "notificationEnabled=false.",
        "input": "-",
        "steps": "1. Chờ đến giờ đã set",
        "expect": "Không có notification được gửi.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 19. DANG XUAT ============
    {
        "id": "TC_OUT_001", "module": "Đăng xuất",
        "name": "Đăng xuất hiện dialog xác nhận",
        "goal": "Kiểm tra confirm trước khi logout.",
        "pre": "Đang ở Profile.",
        "input": "Nhấn Logout.",
        "steps": "1. Nhấn 'Đăng xuất'",
        "expect": "Hiển thị dialog 'Bạn có chắc muốn đăng xuất?'.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_OUT_002", "module": "Đăng xuất",
        "name": "Đăng xuất chuyển về màn Login",
        "goal": "Kiểm tra logout.",
        "pre": "Đang đăng nhập.",
        "input": "Xác nhận logout.",
        "steps": "1. Logout\n2. Xác nhận",
        "expect": "Xoá session FirebaseAuth; UserProvider reset; điều hướng về /login.",
        "type": "Chức năng", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 20. TICH HOP FIRESTORE / DU LIEU ============
    {
        "id": "TC_DB_001", "module": "Firestore/Dữ liệu",
        "name": "Transaction lưu game result không lặp thưởng",
        "goal": "Đảm bảo saveGameResult atomic.",
        "pre": "Mô phỏng mất mạng giữa chừng.",
        "input": "Kết thúc 1 game.",
        "steps": "1. Ngắt mạng tại thời điểm commit\n2. Bật lại mạng\n3. Retry",
        "expect": "Hoặc ghi đầy đủ (coin/XP/score + game result) hoặc không ghi gì; không có trạng thái nửa chừng.",
        "type": "Tích hợp", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_DB_002", "module": "Firestore/Dữ liệu",
        "name": "Seed demo users bổ sung top leaderboard",
        "goal": "Kiểm tra seedLeaderboardDemoUsers (dev).",
        "pre": "Chế độ dev.",
        "input": "-",
        "steps": "1. Gọi seed\n2. Xem Leaderboard",
        "expect": "10 user demo xuất hiện với điểm khác nhau.",
        "type": "Tích hợp", "priority": "Thấp", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_DB_003", "module": "Firestore/Dữ liệu",
        "name": "Tự tạo user doc nếu thiếu khi login",
        "goal": "Khắc phục Pigeon bug.",
        "pre": "FirebaseAuth user tồn tại; Firestore users thiếu doc.",
        "input": "Đăng nhập.",
        "steps": "1. Xoá doc users/{uid} trên Firestore\n2. Đăng nhập lại",
        "expect": "AuthService tự tạo lại doc với giá trị mặc định; không crash.",
        "type": "Tích hợp", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 21. LOCAL STORAGE ============
    {
        "id": "TC_LS_001", "module": "Lưu trữ cục bộ",
        "name": "SharedPreferences lưu cờ isFirstTime",
        "goal": "Kiểm tra local storage.",
        "pre": "Cài app lần đầu.",
        "input": "Mở app lần 1, lần 2.",
        "steps": "1. Mở lần 1\n2. Thoát\n3. Mở lần 2",
        "expect": "Lần 1 isFirstTime=true (hoặc hiển thị onboarding); lần 2 bỏ qua onboarding.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },

    # ============ 22. HIEU NANG / KHAC ============
    {
        "id": "TC_PERF_001", "module": "Hiệu năng",
        "name": "Thời gian khởi động app < 3s",
        "goal": "Đảm bảo load nhanh.",
        "pre": "Thiết bị trung cấp; đã đăng nhập.",
        "input": "-",
        "steps": "1. Đóng app hoàn toàn\n2. Mở lại; bấm giờ",
        "expect": "App hiển thị Home trong dưới 3 giây.",
        "type": "Hiệu năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_PERF_002", "module": "Hiệu năng",
        "name": "Xoay màn hình trong game không crash",
        "goal": "Kiểm tra lifecycle.",
        "pre": "Đang chơi 1 game bất kỳ.",
        "input": "Xoay ngang / dọc.",
        "steps": "1. Xoay thiết bị trong game",
        "expect": "App không crash; state game bảo toàn hoặc xử lý gracefully.",
        "type": "Hiệu năng", "priority": "Thấp", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_UI_001", "module": "Giao diện chung",
        "name": "Hiển thị đúng trên thiết bị nhỏ (360x640)",
        "goal": "Kiểm tra responsive.",
        "pre": "Máy ảo 360x640.",
        "input": "-",
        "steps": "1. Duyệt qua tất cả màn chính",
        "expect": "Không có text overflow, không bị khuất nút bấm.",
        "type": "Giao diện", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_UI_002", "module": "Giao diện chung",
        "name": "Hiển thị đúng trên màn hình lớn (tablet)",
        "goal": "Kiểm tra layout tablet.",
        "pre": "Tablet 10 inch.",
        "input": "-",
        "steps": "1. Mở các màn chính",
        "expect": "Bố cục không bị co kéo bất thường; các phần tử căn giữa hợp lý.",
        "type": "Giao diện", "priority": "Thấp", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_I18N_001", "module": "Đa ngôn ngữ",
        "name": "Không có hardcoded text",
        "goal": "Toàn bộ text qua localization.",
        "pre": "Đổi ngôn ngữ En.",
        "input": "-",
        "steps": "1. Duyệt các màn hình chính",
        "expect": "Không còn text tiếng Việt lẫn; tất cả được dịch.",
        "type": "Chức năng", "priority": "Trung bình", "status": "Chưa thực thi",
        "note": ""
    },
    {
        "id": "TC_SEC_001", "module": "Bảo mật",
        "name": "Firestore rule cấm truy cập dữ liệu user khác",
        "goal": "Kiểm tra security rules.",
        "pre": "Có tài khoản A và B.",
        "input": "Tài khoản A cố đọc users/{B}.",
        "steps": "1. Đăng nhập A\n2. Gọi API đọc doc của B",
        "expect": "Firestore trả về permission denied.",
        "type": "Bảo mật", "priority": "Cao", "status": "Chưa thực thi",
        "note": "Kiểm tra firestore.rules."
    },
    {
        "id": "TC_SEC_002", "module": "Bảo mật",
        "name": "Không log password trong console",
        "goal": "Đảm bảo an toàn thông tin.",
        "pre": "Debug mode.",
        "input": "Đăng nhập.",
        "steps": "1. Bật log\n2. Đăng nhập",
        "expect": "Không có log in plain password ra stdout/stderr.",
        "type": "Bảo mật", "priority": "Cao", "status": "Chưa thực thi",
        "note": ""
    },
]

# ------------- TAO TAI LIEU ---------------

doc = Document()

# Thiet lap section landscape de bang rong
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
new_w, new_h = section.page_height, section.page_width
section.page_width = new_w
section.page_height = new_h
section.left_margin = Cm(1.2)
section.right_margin = Cm(1.2)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

# Font mac dinh cho style Normal
style = doc.styles['Normal']
style.font.name = "Times New Roman"
style.font.size = Pt(11)

# ================ TRANG BIA ================
for _ in range(3):
    doc.add_paragraph()

add_paragraph(doc, "TÀI LIỆU KIỂM THỬ PHẦN MỀM",
              bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER,
              color=RGBColor(0x1F, 0x4E, 0x79))
add_paragraph(doc, "(Software Test Cases Document)",
              italic=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
              color=RGBColor(0x59, 0x59, 0x59))

doc.add_paragraph()
doc.add_paragraph()

add_paragraph(doc, "ỨNG DỤNG VOCAB QUEST",
              bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER,
              color=RGBColor(0xC0, 0x50, 0x4D))
add_paragraph(doc, "Học từ vựng tiếng Anh qua trò chơi",
              italic=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info_tbl = doc.add_table(rows=4, cols=2)
info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
info_tbl.autofit = False
for r in info_tbl.rows:
    r.cells[0].width = Cm(6)
    r.cells[1].width = Cm(10)

rows_info = [
    ("Dự án", "VocabQuest - Ứng dụng học từ vựng (Flutter + Firebase)"),
    ("Nền tảng", "Android / iOS / Web (Flutter)"),
    ("Loại tài liệu", "Bộ Test Case kiểm thử toàn hệ thống"),
    ("Ngày lập", date.today().strftime("%d/%m/%Y")),
]
for idx, (k, v) in enumerate(rows_info):
    fill_cell(info_tbl.rows[idx].cells[0], k, bold=True, size=12, bg="DDEBF7")
    fill_cell(info_tbl.rows[idx].cells[1], v, size=12)

doc.add_page_break()

# ================ MUC LUC (thu cong don gian) ================
add_heading(doc, "MỤC LỤC", level=1, color=RGBColor(0x1F, 0x4E, 0x79),
            align=WD_ALIGN_PARAGRAPH.CENTER)

toc_items = [
    ("1.  Giới thiệu tài liệu", 3),
    ("2.  Phạm vi kiểm thử", 3),
    ("3.  Chiến lược và tiêu chí kiểm thử", 4),
    ("4.  Môi trường kiểm thử", 5),
    ("5.  Mô tả các cột trong bảng test case", 5),
    ("6.  Bảng tổng hợp test case theo module", 6),
    ("7.  CHI TIẾT TEST CASE", 7),
    ("8.  Tổng kết và khuyến nghị", "cuối"),
]
for t, p in toc_items:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    r = para.add_run(t)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    tab = para.add_run(f"\t{p}")
    tab.font.size = Pt(12)

doc.add_page_break()

# ================ 1. GIOI THIEU ================
add_heading(doc, "1. GIỚI THIỆU TÀI LIỆU", level=1,
            color=RGBColor(0x1F, 0x4E, 0x79))

add_paragraph(doc,
    "Tài liệu này mô tả bộ test case được thiết kế để kiểm thử toàn diện "
    "ứng dụng VocabQuest - một ứng dụng học từ vựng tiếng Anh thông qua các "
    "trò chơi tương tác (Quiz, Matching, Memory, Word Puzzle). Mục tiêu là "
    "đảm bảo các chức năng hoạt động đúng đặc tả, xử lý đúng các tình huống "
    "bất thường và mang lại trải nghiệm ổn định cho người dùng cuối.",
    size=12)

add_paragraph(doc,
    "Bộ test case được tổ chức theo module chức năng, mỗi test case bao gồm "
    "các cột thông tin chuẩn giống như các tài liệu kiểm thử thực tế trong "
    "doanh nghiệp, hỗ trợ dễ dàng cho việc thực thi, theo dõi và báo cáo.",
    size=12)

# ================ 2. PHAM VI ================
add_heading(doc, "2. PHẠM VI KIỂM THỬ", level=1,
            color=RGBColor(0x1F, 0x4E, 0x79))

add_paragraph(doc, "Các module/chức năng được kiểm thử:", bold=True, size=12)

scope = [
    "Xác thực: Đăng ký, Đăng nhập, Quên mật khẩu, Đăng xuất",
    "Màn hình Home: Header, stats, daily challenge, điều hướng",
    "Bốn trò chơi: Quiz, Matching, Memory, Word Puzzle",
    "Pack Selection, Level Map, Game Menu, Game Result",
    "Hồ sơ, Chỉnh sửa hồ sơ (đổi tên, đổi avatar)",
    "Yêu thích, Bảng xếp hạng, Lịch sử chơi game",
    "Cài đặt: theme, ngôn ngữ, âm thanh, thông báo",
    "Hệ thống Streak & Level, phần thưởng theo mốc",
    "Cửa hàng (Shop) - mua pack từ vựng",
    "Dịch vụ nền: Notification, Firestore, Local Storage",
    "Các khía cạnh phi chức năng: hiệu năng, giao diện, đa ngôn ngữ, bảo mật",
]
for s in scope:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(s)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

# ================ 3. CHIEN LUOC ================
add_heading(doc, "3. CHIẾN LƯỢC VÀ TIÊU CHÍ KIỂM THỬ", level=1,
            color=RGBColor(0x1F, 0x4E, 0x79))

add_paragraph(doc, "3.1. Các loại test case được thiết kế:", bold=True, size=12)
types = [
    ("Chức năng", "Kiểm tra luồng nghiệp vụ chính hoạt động đúng."),
    ("Xác thực", "Kiểm tra các ràng buộc dữ liệu đầu vào."),
    ("Giá trị biên", "Kiểm tra các giá trị cận trên/cận dưới của ràng buộc."),
    ("Âm tính", "Kiểm tra hành vi khi có lỗi hoặc dữ liệu bất hợp lệ."),
    ("Tích hợp", "Kiểm tra sự phối hợp giữa các module (Firestore, Auth, Provider)."),
    ("Giao diện", "Kiểm tra hiển thị, responsive, animation."),
    ("Hiệu năng", "Kiểm tra thời gian phản hồi, mượt khi thao tác nhiều."),
    ("Bảo mật", "Kiểm tra quyền truy cập và bảo vệ dữ liệu."),
]
tbl_types = doc.add_table(rows=1, cols=2)
tbl_types.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl_types.rows[0].cells
fill_cell(hdr[0], "Loại test", align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(hdr[1], "Mô tả", align=WD_ALIGN_PARAGRAPH.CENTER)
style_table_header(tbl_types.rows[0])
for t_name, t_desc in types:
    row = tbl_types.add_row().cells
    fill_cell(row[0], t_name, bold=True, size=11)
    fill_cell(row[1], t_desc, size=11)
    set_cell_borders(row[0]); set_cell_borders(row[1])
set_col_widths(tbl_types, [5, 18])

add_paragraph(doc, "3.2. Tiêu chí đánh giá Pass/Fail:", bold=True, size=12)
add_paragraph(doc,
    "• Pass: Kết quả thực tế khớp với 'Kết quả mong đợi' trên mọi bước.\n"
    "• Fail: Có ít nhất một bước cho kết quả khác với mong đợi hoặc phát sinh crash, exception.\n"
    "• Blocked: Không thể thực thi do phụ thuộc vào tính năng khác đang hỏng.\n"
    "• Chưa thực thi: Chưa có dữ liệu kiểm thử hoặc chưa tới lượt.",
    size=12)

# ================ 4. MOI TRUONG ================
add_heading(doc, "4. MÔI TRƯỜNG KIỂM THỬ", level=1,
            color=RGBColor(0x1F, 0x4E, 0x79))

env_tbl = doc.add_table(rows=1, cols=2)
env_hdr = env_tbl.rows[0].cells
fill_cell(env_hdr[0], "Thành phần", align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(env_hdr[1], "Cấu hình", align=WD_ALIGN_PARAGRAPH.CENTER)
style_table_header(env_tbl.rows[0])

env_rows = [
    ("Hệ điều hành thiết bị", "Android 10+, iOS 14+, Windows 11 (web)"),
    ("Thiết bị kiểm thử", "Máy ảo 360x640, 1080x1920; Tablet 10 inch; iPhone 13"),
    ("Framework", "Flutter SDK (stable), Dart 3.x"),
    ("Backend", "Firebase Authentication, Cloud Firestore"),
    ("Mạng", "Wi-Fi / 4G; mô phỏng offline và latency cao"),
    ("Dữ liệu", "Tài khoản test riêng, dataset từ vựng sơ cấp/trung cấp/nâng cao"),
]
for k, v in env_rows:
    row = env_tbl.add_row().cells
    fill_cell(row[0], k, bold=True, size=11)
    fill_cell(row[1], v, size=11)
    set_cell_borders(row[0]); set_cell_borders(row[1])
set_col_widths(env_tbl, [7, 16])

# ================ 5. MO TA COT ================
add_heading(doc, "5. MÔ TẢ CÁC CỘT TRONG BẢNG TEST CASE",
            level=1, color=RGBColor(0x1F, 0x4E, 0x79))

cols_desc = [
    ("ID", "Mã định danh duy nhất của test case (VD: TC_LOGIN_001)."),
    ("Module", "Chức năng/màn hình được kiểm thử."),
    ("Tên test case", "Tiêu đề ngắn gọn mô tả mục đích kịch bản."),
    ("Mục tiêu", "Điều cần xác minh bằng test case này."),
    ("Điều kiện tiên quyết", "Trạng thái hệ thống/dữ liệu cần có trước khi thực thi."),
    ("Dữ liệu đầu vào", "Giá trị cụ thể nhập vào hoặc các tham số thao tác."),
    ("Các bước thực hiện", "Danh sách bước thao tác theo thứ tự."),
    ("Kết quả mong đợi", "Kết quả đúng theo đặc tả."),
    ("Loại test", "Phân loại test (Chức năng, Âm tính, Giá trị biên, ...)."),
    ("Mức độ ưu tiên", "Cao / Trung bình / Thấp - dùng để ưu tiên thực thi."),
    ("Trạng thái", "Pass / Fail / Blocked / Chưa thực thi."),
    ("Ghi chú", "Các lưu ý, liên kết bug, phiên bản test, v.v."),
]
col_tbl = doc.add_table(rows=1, cols=2)
hdr = col_tbl.rows[0].cells
fill_cell(hdr[0], "Cột", align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(hdr[1], "Ý nghĩa", align=WD_ALIGN_PARAGRAPH.CENTER)
style_table_header(col_tbl.rows[0])
for k, v in cols_desc:
    row = col_tbl.add_row().cells
    fill_cell(row[0], k, bold=True, size=11)
    fill_cell(row[1], v, size=11)
    set_cell_borders(row[0]); set_cell_borders(row[1])
set_col_widths(col_tbl, [5, 18])

# ================ 6. TONG HOP ================
add_heading(doc, "6. BẢNG TỔNG HỢP TEST CASE THEO MODULE",
            level=1, color=RGBColor(0x1F, 0x4E, 0x79))

# dem theo module
from collections import Counter
module_counter = Counter(tc["module"] for tc in TESTCASES)

sum_tbl = doc.add_table(rows=1, cols=3)
hdr = sum_tbl.rows[0].cells
fill_cell(hdr[0], "STT", align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(hdr[1], "Module / Chức năng", align=WD_ALIGN_PARAGRAPH.CENTER)
fill_cell(hdr[2], "Số lượng test case", align=WD_ALIGN_PARAGRAPH.CENTER)
style_table_header(sum_tbl.rows[0])

total = 0
for i, (m, c) in enumerate(module_counter.items(), 1):
    row = sum_tbl.add_row().cells
    fill_cell(row[0], i, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    fill_cell(row[1], m, size=11)
    fill_cell(row[2], c, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=11)
    for c_ in row:
        set_cell_borders(c_)
    total += c

row = sum_tbl.add_row().cells
fill_cell(row[0], "", bg="FFF2CC")
fill_cell(row[1], "TỔNG CỘNG", bold=True, bg="FFF2CC", size=12)
fill_cell(row[2], total, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          bg="FFF2CC", size=12,
          color=RGBColor(0xC0, 0x00, 0x00))
for c_ in row:
    set_cell_borders(c_)
set_col_widths(sum_tbl, [2, 14, 6])

doc.add_page_break()

# ================ 7. CHI TIET TEST CASE ================
add_heading(doc, "7. CHI TIẾT TEST CASE", level=1,
            color=RGBColor(0x1F, 0x4E, 0x79))
add_paragraph(doc,
    "Bảng test case được trình bày theo dạng NGANG: mỗi dòng là một test case, "
    "các cột trải ngang theo thứ tự các trường thông tin. Bảng được nhóm theo module "
    "để dễ theo dõi và thực thi.",
    italic=True, size=11, color=RGBColor(0x59, 0x59, 0x59))

LABELS = [
    ("STT",                  "stt"),
    ("ID",                   "id"),
    ("Module",               "module"),
    ("Tên test case",        "name"),
    ("Mục tiêu",             "goal"),
    ("Điều kiện tiên quyết", "pre"),
    ("Dữ liệu đầu vào",      "input"),
    ("Các bước thực hiện",   "steps"),
    ("Kết quả mong đợi",     "expect"),
    ("Loại test",            "type"),
    ("Ưu tiên",              "priority"),
    ("Trạng thái",           "status"),
    ("Ghi chú",              "note"),
]

# Do rong tung cot (cm) - tong ~ 25cm vua khit trang A4 landscape
COL_WIDTHS = [1.0, 2.0, 2.0, 2.8, 2.8, 2.6, 2.4, 3.0, 3.0, 1.5, 1.2, 1.4, 1.8]

# Nhom test case theo module de de theo doi
from itertools import groupby
groups = []
for m, items in groupby(TESTCASES, key=lambda x: x["module"]):
    groups.append((m, list(items)))

stt_counter = 0
for module_name, items in groups:
    # Tieu de module
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(f"▸ Module: {module_name}  ({len(items)} test case)")
    r.bold = True
    r.font.size = Pt(13)
    r.font.name = "Times New Roman"
    r.font.color.rgb = RGBColor(0xC0, 0x50, 0x4D)

    # Tao bang voi header + moi dong la 1 test case
    t = doc.add_table(rows=1, cols=len(LABELS))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    # Header row
    for i, (label, _) in enumerate(LABELS):
        fill_cell(t.rows[0].cells[i], label,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_borders(t.rows[0].cells[i])
    style_table_header(t.rows[0])

    # Data rows
    for tc in items:
        stt_counter += 1
        row = t.add_row().cells
        for i, (_, key) in enumerate(LABELS):
            if key == "stt":
                value = stt_counter
            else:
                value = tc.get(key, "")

            cell = row[i]
            set_cell_borders(cell)

            if key == "stt":
                fill_cell(cell, value, bold=True, size=9,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            elif key == "id":
                fill_cell(cell, value, bold=True, size=9,
                          align=WD_ALIGN_PARAGRAPH.CENTER,
                          color=RGBColor(0x1F, 0x4E, 0x79))
            elif key == "priority":
                fill_cell(cell, value, bold=True, size=9,
                          align=WD_ALIGN_PARAGRAPH.CENTER,
                          color=severity_color(value))
            elif key == "status":
                fill_cell(cell, value, size=9,
                          align=WD_ALIGN_PARAGRAPH.CENTER,
                          bg=status_bg(value))
            elif key == "type":
                fill_cell(cell, value, size=9,
                          align=WD_ALIGN_PARAGRAPH.CENTER,
                          color=RGBColor(0x1F, 0x4E, 0x79))
            elif key == "steps":
                # Giu xuong dong cho cac buoc
                cell.text = ""
                first = True
                for line in str(value).split("\n"):
                    p = cell.paragraphs[0] if first else cell.add_paragraph()
                    pr = p.add_run(line)
                    pr.font.size = Pt(9)
                    pr.font.name = "Times New Roman"
                    first = False
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            else:
                fill_cell(cell, value, size=9)

    set_col_widths(t, COL_WIDTHS)

# ================ 8. TONG KET ================
doc.add_page_break()
add_heading(doc, "8. TỔNG KẾT VÀ KHUYẾN NGHỊ", level=1,
            color=RGBColor(0x1F, 0x4E, 0x79))

add_paragraph(doc,
    f"Tổng số test case đã thiết kế: {total} test case, trải rộng trên "
    f"{len(module_counter)} module/chức năng lớn của ứng dụng VocabQuest.",
    size=12)

add_paragraph(doc, "Khuyến nghị triển khai:", bold=True, size=12)
recs = [
    "Thực thi trước các test case mức Ưu tiên CAO thuộc luồng Đăng ký - Đăng nhập - Chơi game - Lưu kết quả.",
    "Tự động hoá các kịch bản ổn định (Auth, Streak, Level) bằng Flutter integration_test.",
    "Kiểm thử thủ công cho các kịch bản UI/animation (Memory flip, timer Quiz, hiệu ứng ghép từ).",
    "Dùng Firebase Emulator để thử nghiệm transaction mua pack, lưu điểm, cập nhật streak mà không ảnh hưởng dữ liệu thật.",
    "Trong mỗi chu kỳ phát hành, cập nhật cột 'Trạng thái' và 'Ghi chú' để theo dõi tiến độ kiểm thử.",
    "Liên kết mỗi test case fail với một ticket/issue (Jira, GitHub) để dễ truy vết.",
]
for s in recs:
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(s)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

doc.add_paragraph()
add_paragraph(doc, "--- Hết tài liệu ---", italic=True,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              color=RGBColor(0x80, 0x80, 0x80))

# Luu
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print("SAVED:", OUT_PATH)
print("TOTAL TEST CASES:", total)
print("MODULES:", len(module_counter))
