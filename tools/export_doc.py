# -*- coding: utf-8 -*-
"""Xuat tai lieu logic VocabQuest sang file .docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = r"C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_Logic_Game.docx"

doc = Document()

# ===== Styling helpers =====
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x4A, 0x44, 0xCC)


def h3(text):
    p = doc.add_heading(text, level=3)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    doc.add_paragraph()


# ===== Title page =====
title = doc.add_heading('VocabQuest', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Tài liệu phân tích logic game & hệ thống điểm')
run.italic = True
run.font.size = Pt(14)
doc.add_paragraph()

# ===== 1. Tong quan 4 games =====
h1('1. Tổng quan 4 mini-games')
make_table(
    ['Game', 'Type ID', 'Icon', 'Gradient', 'Mô tả ngắn'],
    [
        ['Nối từ', 'matching', 'puzzle-piece', 'Tím',
         'Ghép từ Anh ↔ nghĩa Việt (2 cột)'],
        ['Trắc nghiệm', 'quiz', 'circle-question', 'Cam',
         'Nghe/đọc từ, chọn 1 trong 4 nghĩa'],
        ['Xếp chữ', 'word_puzzle', 'spell-check', 'Hồng',
         'Sắp xếp chữ cái thành từ tiếng Anh'],
        ['Lật thẻ', 'memory', 'brain', 'Xanh dương',
         'Pexeso — lật thẻ tìm cặp từ-nghĩa'],
    ],
)

# ===== 2. Luat choi tung game =====
h1('2. Luật chơi chi tiết từng game')

h2('2.1. Matching (Nối từ)')
make_table(
    ['Thuộc tính', 'Giá trị'],
    [
        ['Layout', '2 cột (trái: English word, phải: Vietnamese meaning)'],
        ['Số cặp/round', '4 cặp'],
        ['Số rounds', 'Tối đa 3 (tự giảm nếu pack ít từ)'],
        ['Không lặp từ', '✅ 12 từ unique nếu pack đủ'],
        ['Timer', '60 giây cho toàn game'],
        ['Điểm/cặp đúng lần đầu', '+10'],
        ['Retry sau sai', 'Cho phép, nhưng KHÔNG cộng điểm'],
        ['Báo đỏ khi sai', 'Chỉ 2 ô vừa click (word trái + meaning phải)'],
        ['Pair "fail"', 'Tính theo meaning user click sai'],
        ['Max điểm/round', '40 (4 × 10)'],
        ['Max điểm/game', '120 (3 × 40)'],
    ],
)

h2('2.2. Quiz (Trắc nghiệm)')
make_table(
    ['Thuộc tính', 'Giá trị'],
    [
        ['Layout', '1 từ English hiện giữa, 4 đáp án nghĩa bên dưới'],
        ['Số câu', '10 câu'],
        ['Timer', '15s / câu (độc lập, hết giờ = sai)'],
        ['Flutter TTS', '✅ Phát âm từ tiếng Anh (en-US)'],
        ['Điểm/câu đúng', '10 + thời gian còn lại (0-15)'],
        ['Điểm/câu sai', '0'],
        ['Timeout', 'Tính là sai, tự chuyển câu sau 2s'],
        ['Retry', 'Không có — đã chọn không đổi được'],
        ['Max điểm/câu', '25 (10 + 15s còn lại)'],
        ['Max điểm/game', '≈ 250 (hiếm — cần đúng tức thì cả 10 câu)'],
    ],
)

h2('2.3. Word Puzzle (Xếp chữ)')
make_table(
    ['Thuộc tính', 'Giá trị'],
    [
        ['Layout', '1 từ mỗi lượt, chữ cái shuffle dưới, ô trống trên'],
        ['Số từ', '5 từ'],
        ['Timer', 'Không có'],
        ['Điểm/từ đúng', '+20'],
        ['Điểm/từ sai', '0, chuyển sau 1.8s'],
        ['Hint', 'Tiết lộ 1 chữ đúng, tốn 10 coin'],
        ['Check', 'Tự check khi đã điền đủ'],
        ['Max điểm/game', '100 (5 × 20)'],
    ],
)

h2('2.4. Memory (Lật thẻ)')
make_table(
    ['Thuộc tính', 'Giá trị'],
    [
        ['Layout', 'Grid 3 cột × 4 hàng = 12 thẻ (6 cặp word-meaning)'],
        ['Số rounds', 'Tối đa 3 (tự giảm nếu pack < 18 từ)'],
        ['Không lặp từ', '✅ 18 từ unique nếu pack đủ'],
        ['Timer', '90 giây toàn game'],
        ['Mặt úp', 'Tím + icon EN/VI + badge'],
        ['Mặt ngửa', 'Xanh dương (EN) / Hồng (VI) với text'],
        ['Match khi', 'Cùng vocabId + 1 là word, 1 là meaning'],
        ['Điểm/cặp match lần đầu', '+20'],
        ['Retry sau sai', 'Cho phép nhưng KHÔNG cộng điểm'],
        ['Không match', 'Úp lại sau 800ms'],
        ['Max điểm/game', '360 (3 × 6 × 20)'],
    ],
)

# ===== 3. Coin & XP =====
h1('3. Công thức Coin & XP từ kết quả game')
para(
    'File game_provider.dart — _calculateRewards() — áp dụng cho mọi game '
    'sau khi tính xong score.',
    italic=True,
)

h2('3.1. Base (bắt buộc)')
make_table(
    ['Loại', 'Công thức'],
    [
        ['Coin cơ bản', 'round(score × 0.5)'],
        ['XP cơ bản', '= score'],
    ],
)

h2('3.2. Bonus độ chính xác (accuracy = correct / total)')
make_table(
    ['Accuracy', 'Coin bonus', 'XP bonus'],
    [
        ['= 100%', '+50', '+30'],
        ['≥ 80%', '+20', '+15'],
        ['< 80%', '0', '0'],
    ],
)

h2('3.3. Bonus tốc độ')
make_table(
    ['Điều kiện', 'Bonus'],
    [
        ['Thời gian chơi < 30 giây', '+10 coin'],
    ],
)

h2('3.4. Nhân hệ số theo độ khó pack')
make_table(
    ['Level pack', 'Nhân XP'],
    [
        ['beginner', '×1.0'],
        ['intermediate', '×1.5'],
        ['advanced', '×2.0'],
    ],
)

h3('Ví dụ tính toán')
para(
    'Chơi Quiz intermediate được 150 điểm, 10/10 đúng, 25 giây:',
    bold=True,
)
p = doc.add_paragraph()
p.add_run('• Base: 75 coin, 150 XP\n')
p.add_run('• Accuracy 100%: +50 coin, +30 XP\n')
p.add_run('• Speed < 30s: +10 coin\n')
p.add_run('• Nhân intermediate: XP × 1.5 = (150+30) × 1.5 = 270 XP\n')
p.add_run('→ Tổng: 135 coin + 270 XP').bold = True

# ===== 4. Star rating =====
h1('4. Star rating (sao cuối game)')
para(
    'File game_result_model.dart — dựa trên accuracy.',
    italic=True,
)
make_table(
    ['Accuracy', 'Sao', 'UI'],
    [
        ['≥ 90%', '3 sao', '"Tuyệt vời!" + crown + confetti'],
        ['≥ 60%', '2 sao', '"Làm tốt lắm!" + medal + confetti'],
        ['≥ 30%', '1 sao', '"Cố lên nhé!"'],
        ['< 30%', '0 sao', '"Hãy thử lại!"'],
    ],
)
para(
    'Điều kiện mở khóa level: cần ≥ 2 sao để pass level và nhận Level Reward.',
    bold=True,
)

# ===== 5. Level Reward =====
h1('5. Level Reward (pass level lần đầu)')
para(
    'File rewards.dart — chỉ nhận 1 lần mỗi (gameType × packId × levelIndex).',
    italic=True,
)
make_table(
    ['Level', 'Coin', 'XP', 'Ghi chú'],
    [
        ['Level 1', '+30', '+20', ''],
        ['Level 2', '+60', '+40', ''],
        ['Level 3', '+150', '+100', '⭐ packCompleted = true'],
    ],
)
para(
    'Logic: user phải đạt ≥ 2 sao + progress level mới cao hơn progress đã '
    'có → ghi vào users/{uid}/progress[gameType|packId], cộng thưởng atomic '
    'qua Firestore transaction.',
)

# ===== 6. Streak =====
h1('6. Hệ thống Streak')
para(
    'File streak_calculator.dart — tính trong transaction saveGameResult.',
    italic=True,
)

h2('6.1. Luật cập nhật streak')
make_table(
    ['Tình huống', 'Streak mới', 'Trạng thái'],
    [
        ['Chưa từng chơi (lastPlayed = null)', '1', 'Bắt đầu'],
        ['Cùng ngày (đã chơi hôm nay)', 'giữ nguyên', 'Không đổi'],
        ['Hôm qua mới chơi', '+1', 'Tăng'],
        ['Bỏ ≥ 1 ngày', 'reset = 1', 'Mất streak cũ'],
    ],
)
para('longestStreak luôn = max(cũ, streak mới).')

h2('6.2. Milestone reward')
make_table(
    ['Streak', 'Bonus coin', 'Bonus XP'],
    [
        ['3 ngày', '+30', '+20'],
        ['7 ngày', '+100', '+75'],
        ['14 ngày', '+200', '+150'],
        ['30 ngày', '+500', '+400'],
        ['60 ngày', '+1000', '+800'],
        ['100 ngày', '+2500', '+2000'],
    ],
)
para(
    'Khi đạt milestone: Game Result hiển thị banner cam-hồng + confetti + '
    'lửa pulse animation. Bonus coin/XP tự động cộng vào user qua Firestore '
    'transaction.',
)

# ===== 7. User Level =====
h1('7. User Level (theo tổng XP)')
para(
    'File constants.dart — class LevelSystem.xpThresholds.',
    italic=True,
)
make_table(
    ['User Level', 'XP cần', 'XP để lên level sau'],
    [
        ['Lv 1', '0', '100'],
        ['Lv 2', '100', '150'],
        ['Lv 3', '250', '250'],
        ['Lv 4', '500', '500'],
        ['Lv 5', '1000', '1000'],
        ['Lv 6', '2000', '1500'],
        ['Lv 7', '3500', '2000'],
        ['Lv 8', '5500', '2500'],
        ['Lv 9', '8000', '4000'],
        ['Lv 10', '12000', 'cap'],
    ],
)

# ===== 8. Transaction flow =====
h1('8. Atomic transaction flow (khi chơi xong)')
para(
    'File firestore_service.dart — saveGameResult() làm tất cả trong 1 '
    'transaction:',
    italic=True,
)
code = doc.add_paragraph()
code.paragraph_format.left_indent = Cm(0.5)
code_run = code.add_run(
    'BEGIN TRANSACTION (users/{uid})\n'
    '  1. Đọc: streak, longestStreak, lastPlayedDate\n'
    '  2. Tính: newStreak, milestoneHit, bonusCoins, bonusXP\n'
    '  3. Tạo doc mới ở game_results/{autoId} với GameResultModel\n'
    '  4. UPDATE users/{uid}:\n'
    '     totalScore += result.score\n'
    '     totalCoins += result.coinsEarned + bonusCoins (milestone)\n'
    '     totalXP    += result.xpEarned + bonusXP (milestone)\n'
    '     streak      = newStreak\n'
    '     longestStreak = max(old, newStreak)\n'
    '     lastPlayedDate = now\n'
    'COMMIT\n'
    '→ Return StreakOutcome cho UI hiển thị celebration'
)
code_run.font.name = 'Consolas'
code_run.font.size = Pt(10)

para(
    'Sau transaction chính: nếu đủ ≥ 2 sao + packId/levelIndex → gọi '
    'recordLevelComplete() (transaction riêng) → cộng Level Reward nếu '
    'level mới cao hơn cũ.',
)

# ===== 9. Shop =====
h1('9. Shop — mua pack từ vựng')
make_table(
    ['Hành động', 'Coin cost', 'Logic'],
    [
        ['Mua pack', 'Biến thiên',
         'Transaction: kiểm đủ coin + chưa sở hữu → trừ coin + append '
         'ownedPacks. Free pack: beginner, intermediate, advanced'],
    ],
)

# ===== 10. Features =====
h1('10. Các tính năng hỗ trợ')
make_table(
    ['Tính năng', 'Lưu ở đâu', 'Ghi chú'],
    [
        ['Avatar upload', 'Firestore base64 (≤500KB)',
         'Lưu data URI trong users/{uid}.avatarUrl'],
        ['Favorites', 'SharedPreferences local',
         'Offline, list gameType IDs'],
        ['Nhắc nhở hàng ngày',
         'Local notification + timezone (Asia/HCM)',
         'Exact alarm, reschedule khi đổi giờ/bật lại app'],
        ['Ngôn ngữ', 'SharedPreferences',
         'VI mặc định, switch VI↔EN reactive qua context.t(key)'],
        ['Dark mode', 'SharedPreferences', 'Theme switch runtime'],
    ],
)

# ===== Save =====
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
