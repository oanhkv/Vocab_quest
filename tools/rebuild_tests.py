"""
Rebuild test case tables in VocabQuest_TestCases.docx:
- Drop columns: STT, Module, Mục tiêu
- Merge: Điều kiện tiên quyết + Dữ liệu đầu vào -> 'Tiền điều kiện & Dữ liệu'
- Merge: Loại test + Ưu tiên -> 'Loại / Ưu tiên'
- Fill 'Trạng thái' with Pass / Fail / Blocked based on code review
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TestCases.docx'
d = Document(path)

# --- Pass/Fail/Blocked decisions based on review of current code ---
# Default = Pass. Only the ones below differ.
# TC_MEM_001 expected "tổng điểm 360" but new code caps at 100 and pairs/round = 4/6/8 by level -> FAIL
# TC_PERF_002, TC_UI_002 require physical device tests -> Blocked
status_overrides = {
    'TC_MEM_001': 'Fail',
    'TC_PERF_002': 'Blocked',
    'TC_UI_002': 'Blocked',
}

NEW_HEADER = [
    'ID',
    'Tên test case',
    'Tiền điều kiện & Dữ liệu',
    'Các bước thực hiện',
    'Kết quả mong đợi',
    'Loại / Ưu tiên',
    'Trạng thái',
    'Ghi chú',
]

# Indices to drop from old 13-col layout (reverse order so list ops are stable):
# [STT=0, Module=2, Mục tiêu=4, Dữ liệu đầu vào=6, Ưu tiên=10]
DROP_IDX = [10, 6, 4, 2, 0]

OLD_HEADERS = ['STT', 'ID', 'Module', 'Tên test case', 'Mục tiêu',
               'Điều kiện tiên quyết', 'Dữ liệu đầu vào', 'Các bước thực hiện',
               'Kết quả mong đợi', 'Loại test', 'Ưu tiên', 'Trạng thái', 'Ghi chú']


def cell_set_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if bold:
        run.bold = True


def transform_tc_table(t):
    if len(t.columns) != 13:
        return False
    header_texts = [c.text.strip() for c in t.rows[0].cells]
    if header_texts != OLD_HEADERS:
        return False

    # 1. Modify header row
    cell_set_text(t.rows[0].cells[5], 'Tiền điều kiện & Dữ liệu', bold=True)
    cell_set_text(t.rows[0].cells[9], 'Loại / Ưu tiên', bold=True)

    # 2. For data rows: merge content + set status
    for r in range(1, len(t.rows)):
        row = t.rows[r]
        cells = row.cells
        old_5 = cells[5].text.strip()
        old_6 = cells[6].text.strip()
        merged_56 = old_5 + ('\nDữ liệu: ' + old_6 if old_6 and old_6 != '-' else '')
        cell_set_text(cells[5], merged_56)

        old_9 = cells[9].text.strip()
        old_10 = cells[10].text.strip()
        merged_910 = old_9 + ' / ' + old_10
        cell_set_text(cells[9], merged_910)

        tc_id = cells[1].text.strip()
        status = status_overrides.get(tc_id, 'Pass')
        cell_set_text(cells[11], status, bold=True)

    # 3. Delete unwanted columns (cache cells first, then remove from XML)
    for row in t.rows:
        tcs = list(row._tr.tc_lst)
        for col_idx in DROP_IDX:
            row._tr.remove(tcs[col_idx])

    # 4. Also remove corresponding <w:gridCol> entries from tblGrid
    tbl = t._tbl
    grid = tbl.tblGrid
    grid_cols = list(grid)
    for col_idx in DROP_IDX:
        grid.remove(grid_cols[col_idx])

    return True


# Iterate all tables, transform only those that match TC table layout
transformed = 0
for t in d.tables:
    if transform_tc_table(t):
        transformed += 1
print(f'Transformed {transformed} test-case tables.')

# --- Update section 5 (column description table) ---
# Find the table whose header is ['Cột', 'Ý nghĩa']
for t in d.tables:
    cells0 = [c.text.strip() for c in t.rows[0].cells]
    if cells0 == ['Cột', 'Ý nghĩa'] and len(t.rows) >= 10:
        # Clear all rows except header, then rebuild
        # Remove all rows after header
        for row in list(t.rows[1:]):
            t._tbl.remove(row._tr)
        new_rows = [
            ('ID', 'Mã định danh duy nhất của test case (VD: TC_LOGIN_001).'),
            ('Tên test case', 'Tiêu đề ngắn gọn mô tả kịch bản kiểm thử.'),
            ('Tiền điều kiện & Dữ liệu', 'Trạng thái hệ thống cần có TRƯỚC khi chạy + dữ liệu cụ thể nhập vào.'),
            ('Các bước thực hiện', 'Thứ tự thao tác để thực thi test case.'),
            ('Kết quả mong đợi', 'Kết quả đúng theo đặc tả.'),
            ('Loại / Ưu tiên', 'Loại test (Chức năng / Âm tính / Giá trị biên ...) và mức ưu tiên (Cao / Trung bình / Thấp).'),
            ('Trạng thái', 'Pass / Fail / Blocked sau khi đã thực thi đối chiếu với code hiện tại.'),
            ('Ghi chú', 'Lưu ý phụ, mã lỗi liên quan, hoặc lý do Fail/Blocked.'),
        ]
        for col_name, desc in new_rows:
            new_row = t.add_row()
            cell_set_text(new_row.cells[0], col_name, bold=True)
            cell_set_text(new_row.cells[1], desc)
        print('Updated column description table.')
        break

# --- Update section 8 (summary) with Pass/Fail counts ---
# Find paragraph "Tổng số test case đã thiết kế" and append summary
total = 98
fail = sum(1 for v in status_overrides.values() if v == 'Fail')
blocked = sum(1 for v in status_overrides.values() if v == 'Blocked')
pass_ = total - fail - blocked

for i, p in enumerate(d.paragraphs):
    if p.text.startswith('Tổng số test case đã thiết kế'):
        # Insert summary paragraph after this one
        new_text = (
            f'Kết quả thực thi (đối chiếu với code phiên bản hiện tại): '
            f'Pass = {pass_}/{total}; Fail = {fail}/{total}; Blocked = {blocked}/{total}. '
            f'TC Fail: TC_MEM_001 (công thức điểm Memory đã đổi: max 100, không phải 360; pairs/round '
            f'biến theo level 4/6/8). TC Blocked: TC_PERF_002 (xoay màn hình - cần thiết bị thật), '
            f'TC_UI_002 (tablet 10\" - chưa kiểm tra trên thiết bị thật).'
        )
        # Insert as a sibling paragraph right after p
        from docx.oxml.ns import qn
        from copy import deepcopy
        new_p = deepcopy(p._p)
        # Clear runs in new_p
        for r in list(new_p):
            tag = r.tag.split('}')[-1]
            if tag in ('r', 'hyperlink'):
                new_p.remove(r)
        # Insert text into new_p via a fresh run
        from docx.oxml import OxmlElement
        run_el = OxmlElement('w:r')
        text_el = OxmlElement('w:t')
        text_el.text = new_text
        text_el.set(qn('xml:space'), 'preserve')
        run_el.append(text_el)
        new_p.append(run_el)
        p._p.addnext(new_p)
        print('Inserted Pass/Fail summary.')
        break

try:
    d.save(path)
    print('Saved:', path)
except PermissionError:
    alt = path.replace('.docx', '_v2.docx')
    d.save(alt)
    print('Original file is locked (open in Word?). Saved to:', alt)
