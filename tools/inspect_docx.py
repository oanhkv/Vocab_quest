import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_Logic_Game.docx'
d = Document(path)
print('=== STYLES USED ===')
seen = set()
for p in d.paragraphs:
    name = p.style.name
    if name not in seen:
        seen.add(name)
        print(' -', name)

print('\n=== PARAGRAPHS ===')
for i, p in enumerate(d.paragraphs):
    print(f'[{i}] ({p.style.name}) {p.text[:120]}')
