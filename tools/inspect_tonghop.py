import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TongHop.docx'
d = Document(path)

print('=== STYLES ===')
seen = set()
for p in d.paragraphs:
    if p.style.name not in seen:
        seen.add(p.style.name)
        print(' -', p.style.name)

print('\n=== HEADINGS / structure ===')
for i, p in enumerate(d.paragraphs):
    if 'Heading' in p.style.name or p.style.name == 'Title':
        print(f'P[{i}] ({p.style.name}) {p.text[:140]}')

print('\n=== ALL TABLES (header rows only) ===')
for ti, t in enumerate(d.tables):
    head = [c.text.strip()[:40] for c in t.rows[0].cells]
    print(f'TABLE[{ti}] rows={len(t.rows)} cols={len(t.columns)} head={head}')
