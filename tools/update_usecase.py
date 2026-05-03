"""
Cập nhật PHẦN IV trong VocabQuest_TongHop_v2.docx để khớp với phiên bản PlantUML mới
(rút gọn còn 4 actor + 13 UC chính + 4 mini-game con; đường thẳng linetype ortho).
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TongHop_v2.docx'
d = Document(path)


def cell_set_text(cell, text, bold=False, mono=False, size=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if bold:
        run.bold = True
    if mono:
        run.font.name = 'Consolas'
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Consolas')
        rFonts.set(qn('w:hAnsi'), 'Consolas')
    if size:
        run.font.size = Pt(size)


def rebuild_table(t, header, rows):
    """Clear table, set header to `header`, then add `rows`."""
    # Remove all rows
    for row in list(t.rows):
        t._tbl.remove(row._tr)
    # Add header
    new_row = t.add_row()
    while len(new_row.cells) < len(header):
        new_row.add_cell()
    for i, h in enumerate(header):
        cell_set_text(new_row.cells[i], h, bold=True)
    # Add data rows
    for row_data in rows:
        new_row = t.add_row()
        for i, val in enumerate(row_data):
            cell_set_text(new_row.cells[i], val)


# ===== 1. UPDATE ACTOR TABLE =====
actor_header = ['STT', 'Tên Actor', 'Loại', 'Mô tả', 'Use case liên quan']
actor_rows = [
    ['A1', 'Khách (Guest)', 'Primary',
     'Người dùng chưa đăng nhập, chỉ truy cập màn hình Login/Register.',
     'UC-01, UC-02'],
    ['A2', 'Học viên (Learner)', 'Primary',
     'Người dùng đã đăng nhập. Là actor trung tâm — chơi 4 mini-game, '
     'mở level, mua pack, xem leaderboard, cá nhân hóa app. KẾ THỪA từ Khách.',
     'UC-03 → UC-12'],
    ['A3', 'Firebase Authentication', 'Secondary',
     'Dịch vụ xác thực ngoài (email/password). Xử lý đăng ký, đăng nhập, đăng xuất.',
     'UC-01, UC-02, UC-03'],
    ['A4', 'Cloud Firestore', 'Secondary',
     'Cơ sở dữ liệu NoSQL ngoài app. Lưu user, game_results, progress, '
     'ownedPacks, leaderboard, streak.',
     'UC-05, UC-06, UC-10, UC-11, UC-13'],
]

# ===== 2. UPDATE USE CASE TABLE =====
uc_header = ['Mã', 'Tên Use Case', 'Actor chính', 'Mô tả ngắn']
uc_rows = [
    ['UC-01', 'Đăng ký tài khoản', 'Khách',
     'Tạo tài khoản mới (email + mật khẩu); khởi tạo document user trên Firestore.'],
    ['UC-02', 'Đăng nhập', 'Khách',
     'Xác thực bằng email + mật khẩu, vào màn hình chính.'],
    ['UC-03', 'Đăng xuất', 'Học viên',
     'Xoá session FirebaseAuth, quay về màn Login.'],
    ['UC-04', 'Chọn pack & level', 'Học viên',
     'Duyệt danh sách pack đã sở hữu, chọn level đã mở khóa để chơi.'],
    ['UC-05', 'Chơi mini-game', 'Học viên',
     'Use case TỔNG QUÁT — được hiện thực bởi 4 game con bên dưới.'],
    ['UC-05a', 'Chơi Quiz', '(con của UC-05)',
     'Trả lời 10 câu trắc nghiệm 4 đáp án; mỗi câu 15 giây.'],
    ['UC-05b', 'Chơi Matching', '(con của UC-05)',
     'Nối 4 cặp word ↔ meaning trong 60 giây, tối đa 3 round.'],
    ['UC-05c', 'Chơi Memory', '(con của UC-05)',
     'Lật thẻ tìm cặp; pairs/round 4-6-8 theo level; 90 giây cả ván.'],
    ['UC-05d', 'Chơi Word Puzzle', '(con của UC-05)',
     'Sắp xếp chữ cái thành từ tiếng Anh, 5 từ mỗi ván; có hint -10 coin.'],
    ['UC-06', 'Xem kết quả', 'Học viên',
     'Hiển thị điểm, sao, coin/XP, streak bonus sau mỗi ván chơi.'],
    ['UC-07', 'Mở khóa level mới', '—',
     'Khi đạt sao đủ ngưỡng (≥2 sao thường, =3 sao cho Memory) — mở rộng UC-06.'],
    ['UC-08', 'Xem hồ sơ', 'Học viên',
     'Hiển thị tên, email, avatar, stats theo từng loại game.'],
    ['UC-09', 'Sửa hồ sơ', '(extend UC-08)',
     'Đổi tên hiển thị, đổi avatar (≤ 500KB, lưu base64 trên Firestore).'],
    ['UC-10', 'Mua pack từ vựng', 'Học viên',
     'Trả coin để sở hữu pack mới. Transaction kiểm coin + append ownedPacks.'],
    ['UC-11', 'Xem leaderboard', 'Học viên',
     'Top 100 theo Score / XP, hiển thị thứ hạng cá nhân.'],
    ['UC-12', 'Cài đặt app', 'Học viên',
     'Đổi theme (Dark/Light), ngôn ngữ (VI/EN), bật/tắt âm thanh.'],
    ['UC-13', 'Cập nhật streak', 'Hệ thống',
     'Use case ngầm: tăng/giữ/reset streak theo lastPlayedDate. Bắt buộc trong UC-05.'],
]

# Find and rebuild tables
for t in d.tables:
    if not t.rows:
        continue
    h = [c.text.strip() for c in t.rows[0].cells]
    if h == ['STT', 'Tên Actor', 'Loại', 'Mô tả', 'Use case liên quan']:
        rebuild_table(t, actor_header, actor_rows)
        print('Updated Actor table.')
    elif h == ['Mã', 'Tên Use Case', 'Actor chính', 'Mô tả ngắn']:
        rebuild_table(t, uc_header, uc_rows)
        print('Updated Use Case table.')

# ===== 3. UPDATE PLANTUML #1 CODE BLOCK =====
new_uml1 = '''@startuml VocabQuest_UseCase
skinparam linetype ortho
skinparam packageStyle rectangle
skinparam shadowing false
skinparam usecase {
  BackgroundColor #FFF8E7
  BorderColor #444444
}
skinparam actor {
  BackgroundColor #E3F2FD
  BorderColor #1565C0
}
left to right direction
title Biểu đồ Use Case tổng quát — VocabQuest

\' ==== Actors ====
actor "Khách"     as Guest
actor "Học viên"  as Learner
Learner --|> Guest : kế thừa

actor "Firebase Auth"   as FAuth   <<system>>
actor "Cloud Firestore" as FStore  <<system>>

\' ==== System boundary ====
rectangle "Hệ thống VocabQuest" {
  usecase "UC-01\\nĐăng ký"            as UC01
  usecase "UC-02\\nĐăng nhập"          as UC02
  usecase "UC-03\\nĐăng xuất"          as UC03
  usecase "UC-04\\nChọn pack & level"  as UC04
  usecase "UC-05\\nChơi mini-game"     as UC05
  usecase "UC-06\\nXem kết quả"        as UC06
  usecase "UC-07\\nMở khóa level"      as UC07

  usecase "UC-05a\\nQuiz"        as UC05a
  usecase "UC-05b\\nMatching"    as UC05b
  usecase "UC-05c\\nMemory"      as UC05c
  usecase "UC-05d\\nWord Puzzle" as UC05d

  usecase "UC-08\\nXem hồ sơ"        as UC08
  usecase "UC-09\\nSửa hồ sơ"        as UC09
  usecase "UC-10\\nMua pack"         as UC10
  usecase "UC-11\\nXem leaderboard"  as UC11
  usecase "UC-12\\nCài đặt app"      as UC12
  usecase "UC-13\\nCập nhật streak"  as UC13
}

\' ==== Associations: Guest ====
Guest -- UC01
Guest -- UC02

\' ==== Associations: Learner ====
Learner -- UC03
Learner -- UC04
Learner -- UC05
Learner -- UC08
Learner -- UC10
Learner -- UC11
Learner -- UC12

\' ==== Secondary actors ====
UC01 -- FAuth
UC02 -- FAuth
UC03 -- FAuth
UC05 -- FStore
UC06 -- FStore
UC10 -- FStore
UC11 -- FStore
UC13 -- FStore

\' ==== Generalization: 4 game kế thừa "Chơi mini-game" ====
UC05a --|> UC05
UC05b --|> UC05
UC05c --|> UC05
UC05d --|> UC05

\' ==== Include (bắt buộc) ====
UC05 ..> UC06 : <<include>>
UC05 ..> UC13 : <<include>>
UC04 ..> UC02 : <<include>>

\' ==== Extend (tùy chọn) ====
UC07 ..> UC06 : <<extend>>
UC10 ..> UC04 : <<extend>>
UC09 ..> UC08 : <<extend>>

@enduml
'''

# Find paragraph with @startuml VocabQuest_UseCase and replace its text
replaced_uml1 = False
for p in d.paragraphs:
    if '@startuml VocabQuest_UseCase' in p.text and 'VocabQuest_Relations' not in p.text:
        # Clear existing runs
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        # Add new run with monospace font
        run = p.add_run(new_uml1)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Consolas')
        rFonts.set(qn('w:hAnsi'), 'Consolas')
        replaced_uml1 = True
        print('Replaced PlantUML #1 code block.')
        break
if not replaced_uml1:
    print('WARNING: PlantUML #1 paragraph NOT found.')

# ===== 4. UPDATE GIẢI THÍCH 7.x =====
# Map: heading text (or starting text of paragraph) → new text
explanation_updates = {
    '7.1. Phân vùng Actor':
        'Biểu đồ rút gọn còn 4 actor: Khách và Học viên (primary, bên trái) + Firebase Auth '
        'và Cloud Firestore (secondary, bên phải). Học viên kế thừa từ Khách bằng quan hệ '
        'generalization, nên không cần vẽ lại các association mà Khách đã có. Việc tách '
        'Firebase thành 2 actor (Auth vs Firestore) phản ánh đúng kiến trúc: dịch vụ xác '
        'thực và lưu trữ là hai hệ thống ngoài độc lập.',
    '7.2. Các quan hệ kế thừa':
        'Có hai quan hệ generalization trong biểu đồ. Thứ nhất, Học viên --|> Khách: bất kỳ '
        'chức năng nào Khách làm được (UC-01 đăng ký, UC-02 đăng nhập) thì Học viên cũng '
        'làm được. Thứ hai và quan trọng hơn — UC-05a, UC-05b, UC-05c, UC-05d --|> UC-05: '
        'bốn mini-game (Quiz, Matching, Memory, Word Puzzle) là các hiện thực cụ thể của use '
        'case tổng quát "Chơi mini-game". Cách vẽ này giúp giữ biểu đồ gọn mà vẫn thể hiện '
        'đủ 4 game đặc trưng của app.',
    '7.3. Các quan hệ <<include>>':
        'UC-05 (Chơi mini-game) <<include>> UC-06 (Xem kết quả) và UC-13 (Cập nhật streak): '
        'mỗi lần kết thúc 1 ván chơi, hai use case này LUÔN được kích hoạt — không có lựa '
        'chọn bỏ qua (saveGameResult chạy trong 1 transaction Firestore). UC-04 (Chọn pack '
        '& level) <<include>> UC-02 (Đăng nhập): không thể chọn pack nếu chưa có session.',
    '7.4. Các quan hệ <<extend>>':
        'UC-07 (Mở khóa level) <<extend>> UC-06: chỉ chạy KHI sao đạt ngưỡng (≥ 2 cho 3 game '
        'thường, = 3 cho Memory). UC-10 (Mua pack) <<extend>> UC-04: trong màn chọn pack, '
        'chỉ pack chưa sở hữu mới có nút Mua. UC-09 (Sửa hồ sơ) <<extend>> UC-08 (Xem hồ sơ): '
        'từ màn xem hồ sơ, nhấn icon edit mới mở luồng chỉnh sửa.',
    '7.5. UC-19 thuộc về ai?':
        'UC-13 (Cập nhật streak — đổi mã từ UC-19 ở bản trước) không có actor người dùng nối '
        'trực tiếp vì nó chạy NGẦM trong transaction lưu kết quả game (xem PHẦN I — mục 8). '
        'Nó được kích hoạt qua quan hệ <<include>> từ UC-05, nên trong sơ đồ chỉ thể hiện '
        'liên kết tới Firestore (actor phụ) mà không có association với Học viên.',
}

# Strategy: walk through paragraphs; when encountering Heading 2 with key text,
# replace the next Normal paragraph's text with new content.
paragraphs = d.paragraphs
for i, p in enumerate(paragraphs):
    if 'Heading' in p.style.name and p.text.strip() in explanation_updates:
        new_text = explanation_updates[p.text.strip()]
        # Find next Normal paragraph
        for j in range(i + 1, min(i + 5, len(paragraphs))):
            np = paragraphs[j]
            if np.style.name == 'Normal' and np.text.strip():
                # Replace
                for r in list(np.runs):
                    r._element.getparent().remove(r._element)
                np.add_run(new_text)
                print(f'Updated explanation: {p.text.strip()[:50]}')
                break

# Also rename heading 7.5 since it still says "UC-19"
for p in d.paragraphs:
    if p.text.strip() == '7.5. UC-19 thuộc về ai?':
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.add_run('7.5. UC-13 (Cập nhật streak) — actor ngầm')
        print('Renamed heading 7.5')
        break

# Save
try:
    d.save(path)
    print('Saved:', path)
except PermissionError:
    alt = path.replace('_v2.docx', '_v3.docx')
    d.save(alt)
    print('Locked. Saved to:', alt)
