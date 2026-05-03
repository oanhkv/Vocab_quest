"""
Append PHẦN IV — Biểu đồ Use Case to VocabQuest_TongHop.docx:
- Bảng Actor
- Bảng Use Case chính
- Code PlantUML cho 2 biểu đồ
- Giải thích biểu đồ và các loại quan hệ
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TongHop.docx'
d = Document(path)


def add_table(rows, header_bold=True):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            if r_idx == 0 and header_bold:
                run.bold = True
    return t


def add_code_block(code):
    """Add a paragraph styled as monospaced code block (single paragraph)."""
    p = d.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    # Set East Asian font as well
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    return p


# ===================== PHẦN IV =====================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN IV — BIỂU ĐỒ USE CASE TỔNG QUÁT (UML)', style='Title')

# ---------- 1. Giới thiệu ----------
d.add_paragraph('1. Giới thiệu', style='Heading 1')
d.add_paragraph(
    'Phần này mô tả biểu đồ Use Case tổng quát của ứng dụng VocabQuest theo chuẩn UML 2.x. '
    'Biểu đồ thể hiện các tác nhân (actor) tương tác với hệ thống và các tình huống sử dụng '
    '(use case) chính. Code PlantUML kèm theo có thể paste vào https://www.plantuml.com/plantuml '
    'hoặc plugin PlantUML trong VS Code/IntelliJ để render.',
    style='Normal',
)

# ---------- 2. Bảng Actor ----------
d.add_paragraph('2. Bảng tác nhân (Actor)', style='Heading 1')
d.add_paragraph(
    'Tác nhân là thực thể NGOÀI hệ thống có tương tác với app. Trong VocabQuest có 2 nhóm: '
    'tác nhân chính (Primary — người dùng trực tiếp) và tác nhân phụ (Secondary — hệ thống bên ngoài).',
    style='Normal',
)
add_table([
    ['STT', 'Tên Actor', 'Loại', 'Mô tả', 'Use case liên quan'],
    ['A1', 'Khách (Guest)', 'Primary',
     'Người dùng chưa đăng nhập, chỉ truy cập được màn hình Login/Register.',
     'UC-01, UC-02, UC-03'],
    ['A2', 'Học viên (Learner)', 'Primary',
     'Người dùng đã đăng ký và đăng nhập thành công. Là actor trung tâm — chơi game, '
     'xem điểm, mua pack, cá nhân hóa app. KẾ THỪA từ Khách.',
     'UC-04 → UC-20'],
    ['A3', 'Firebase Authentication', 'Secondary',
     'Dịch vụ xác thực ngoài (email/password). Xử lý đăng ký, đăng nhập, gửi mail reset.',
     'UC-01, UC-02, UC-03, UC-04'],
    ['A4', 'Cloud Firestore', 'Secondary',
     'Cơ sở dữ liệu NoSQL ngoài app. Lưu user, game_results, leaderboard, progress, ownedPacks.',
     'UC-06, UC-07, UC-09, UC-10, UC-13, UC-14, UC-19'],
    ['A5', 'Local Notification (Android/iOS OS)', 'Secondary',
     'Hệ điều hành phát thông báo nhắc học hàng ngày dựa trên lịch app đặt.',
     'UC-17'],
    ['A6', 'TTS Engine (Flutter TTS)', 'Secondary',
     'Dịch vụ chuyển văn bản thành giọng nói tiếng Anh để phát âm từ vựng trong Quiz.',
     'UC-18'],
])

# ---------- 3. Bảng Use Case ----------
d.add_paragraph('3. Bảng Use Case chính', style='Heading 1')
add_table([
    ['Mã', 'Tên Use Case', 'Actor chính', 'Mô tả ngắn'],
    ['UC-01', 'Đăng ký tài khoản', 'Khách',
     'Tạo tài khoản mới với email + mật khẩu, nhận tài liệu user khởi tạo trên Firestore.'],
    ['UC-02', 'Đăng nhập', 'Khách',
     'Xác thực bằng email + mật khẩu, vào màn hình Home.'],
    ['UC-03', 'Quên mật khẩu', 'Khách',
     'Yêu cầu Firebase gửi email reset mật khẩu.'],
    ['UC-04', 'Đăng xuất', 'Học viên',
     'Xoá session, quay về màn Login.'],
    ['UC-05', 'Xem trang chủ', 'Học viên',
     'Hiển thị header (tên, level, streak, coin), daily challenge, các tile điều hướng.'],
    ['UC-06', 'Chơi mini-game', 'Học viên',
     'Chọn 1 trong 4 game (Quiz / Matching / Memory / Word Puzzle), trả lời/thao tác để tích điểm.'],
    ['UC-07', 'Xem kết quả game', 'Học viên',
     'Xem điểm, sao, coin/XP nhận được sau khi hoàn thành 1 lượt chơi.'],
    ['UC-08', 'Chọn pack & level', 'Học viên',
     'Duyệt pack, xem level map, chọn level đã mở khóa để chơi.'],
    ['UC-09', 'Mở khóa level mới', 'Học viên',
     'Khi đạt ≥ 2 sao (Memory: 3 sao) thì level kế tiếp tự mở.'],
    ['UC-10', 'Mua pack từ vựng', 'Học viên',
     'Trả coin để sở hữu pack mới. Transaction kiểm coin + append ownedPacks.'],
    ['UC-11', 'Xem hồ sơ', 'Học viên',
     'Hiển thị tên, email, avatar, stats theo từng loại game.'],
    ['UC-12', 'Chỉnh sửa hồ sơ', 'Học viên',
     'Đổi tên hiển thị, đổi avatar (≤ 500KB, lưu base64).'],
    ['UC-13', 'Xem bảng xếp hạng', 'Học viên',
     'Top 100 theo Score / XP, hiển thị thứ hạng cá nhân.'],
    ['UC-14', 'Xem lịch sử chơi', 'Học viên',
     'Danh sách 50 game gần nhất với điểm, sao, thời lượng.'],
    ['UC-15', 'Quản lý yêu thích', 'Học viên',
     'Thêm / bỏ favorite cho pack hoặc game.'],
    ['UC-16', 'Cài đặt app', 'Học viên',
     'Đổi theme (Dark/Light), ngôn ngữ (VI/EN), bật/tắt âm thanh.'],
    ['UC-17', 'Đặt nhắc nhở học hàng ngày', 'Học viên',
     'Đặt giờ nhắc → app schedule local notification qua OS.'],
    ['UC-18', 'Phát âm từ vựng', 'Học viên',
     'Trong Quiz, nhấn icon loa → TTS phát âm từ tiếng Anh đang hiển thị.'],
    ['UC-19', 'Cập nhật streak', 'Hệ thống',
     'Sau mỗi lần lưu game result: tăng / giữ / reset streak theo lastPlayedDate.'],
    ['UC-20', 'Nhận thưởng milestone', 'Học viên',
     'Khi streak chạm mốc 3/7/14/30/60/100 ngày → cộng coin/XP bonus.'],
])

# ---------- 4. Sơ đồ Use Case tổng quát ----------
d.add_paragraph('4. Biểu đồ Use Case tổng quát (PlantUML)', style='Heading 1')
d.add_paragraph(
    'Code PlantUML dưới đây dùng các quan hệ chuẩn UML: association (nét thẳng), '
    'generalization (mũi tên rỗng), <<include>> và <<extend>> (mũi tên đứt). '
    'Paste vào trình render PlantUML để xuất hình ảnh.',
    style='Normal',
)

main_uml = """@startuml VocabQuest_UseCase
left to right direction
skinparam packageStyle rectangle
skinparam usecase {
  BackgroundColor #FFF8E7
  BorderColor #444444
}
skinparam actor {
  BackgroundColor #E3F2FD
  BorderColor #1565C0
}
title Biểu đồ Use Case tổng quát — VocabQuest

' ==== Actors ====
actor "Khách\\n(Guest)"            as Guest
actor "Học viên\\n(Learner)"       as Learner
Learner --|> Guest : kế thừa

actor "Firebase Auth"              as FAuth   <<system>>
actor "Cloud Firestore"            as FStore  <<system>>
actor "OS Notification"            as OSNotif <<system>>
actor "TTS Engine"                 as TTS     <<system>>

' ==== System boundary ====
rectangle "Hệ thống VocabQuest" {
  usecase "UC-01\\nĐăng ký"                 as UC01
  usecase "UC-02\\nĐăng nhập"               as UC02
  usecase "UC-03\\nQuên mật khẩu"           as UC03
  usecase "UC-04\\nĐăng xuất"               as UC04
  usecase "UC-05\\nXem trang chủ"           as UC05
  usecase "UC-06\\nChơi mini-game"          as UC06
  usecase "UC-07\\nXem kết quả"             as UC07
  usecase "UC-08\\nChọn pack & level"       as UC08
  usecase "UC-09\\nMở khóa level"           as UC09
  usecase "UC-10\\nMua pack"                as UC10
  usecase "UC-11\\nXem hồ sơ"               as UC11
  usecase "UC-12\\nChỉnh sửa hồ sơ"         as UC12
  usecase "UC-13\\nXem leaderboard"         as UC13
  usecase "UC-14\\nXem lịch sử"             as UC14
  usecase "UC-15\\nQuản lý yêu thích"       as UC15
  usecase "UC-16\\nCài đặt app"             as UC16
  usecase "UC-17\\nĐặt nhắc nhở"            as UC17
  usecase "UC-18\\nPhát âm từ vựng"         as UC18
  usecase "UC-19\\nCập nhật streak"         as UC19
  usecase "UC-20\\nNhận thưởng milestone"   as UC20
}

' ==== Associations: Guest ====
Guest -- UC01
Guest -- UC02
Guest -- UC03

' ==== Associations: Learner ====
Learner -- UC04
Learner -- UC05
Learner -- UC06
Learner -- UC07
Learner -- UC08
Learner -- UC10
Learner -- UC11
Learner -- UC12
Learner -- UC13
Learner -- UC14
Learner -- UC15
Learner -- UC16
Learner -- UC17

' ==== Secondary actors (right side) ====
UC01 -- FAuth
UC02 -- FAuth
UC03 -- FAuth
UC04 -- FAuth

UC06 -- FStore
UC07 -- FStore
UC10 -- FStore
UC13 -- FStore
UC14 -- FStore
UC19 -- FStore

UC17 -- OSNotif
UC18 -- TTS

' ==== Relationships (include / extend) ====
UC06 ..> UC07 : <<include>>
UC06 ..> UC19 : <<include>>
UC10 ..> UC02 : <<include>>
UC07 <.. UC09 : <<extend>>
UC19 <.. UC20 : <<extend>>
UC06 <.. UC18 : <<extend>>
UC08 <.. UC10 : <<extend>>

@enduml
"""
add_code_block(main_uml)

# ---------- 5. Sơ đồ minh họa các loại đường ----------
d.add_paragraph('5. Biểu đồ minh họa các loại quan hệ trong Use Case Diagram', style='Heading 1')
d.add_paragraph(
    'UML Use Case có 4 loại đường nối chính. Biểu đồ riêng dưới đây minh họa từng loại để bạn '
    'dễ tham chiếu khi đọc biểu đồ chính ở mục 4.',
    style='Normal',
)

relations_uml = """@startuml VocabQuest_Relations
left to right direction
title Các loại quan hệ trong Use Case Diagram (UML)

actor "Actor A"  as A
actor "Actor B"  as B

usecase "UC X"   as X
usecase "UC Y"   as Y
usecase "UC Z"   as Z
usecase "UC W"   as W
usecase "UC V"   as V

' 1. Association — nét thẳng giữa Actor và Use Case
A -- X : 1. Association (nét thẳng)

' 2. Generalization — mũi tên rỗng (Actor B kế thừa Actor A)
B --|> A : 2. Generalization (kế thừa)

' 3. Include — mũi tên đứt với <<include>> (UC Y bắt buộc gọi UC Z)
Y ..> Z : 3. <<include>>\\nbắt buộc

' 4. Extend — mũi tên đứt với <<extend>> (UC W mở rộng UC V — tùy chọn)
W ..> V : 4. <<extend>>\\ntùy chọn

A -- Y
B -- W
@enduml
"""
add_code_block(relations_uml)

# ---------- 6. Bảng giải thích các loại đường ----------
d.add_paragraph('6. Giải thích các loại đường nối', style='Heading 1')
add_table([
    ['Loại quan hệ', 'Ký hiệu PlantUML', 'Hình ảnh', 'Ý nghĩa', 'Ví dụ trong VocabQuest'],
    ['Association', 'A -- UC',
     'Nét thẳng (—)',
     'Actor có tương tác trực tiếp với Use Case. Là quan hệ phổ biến nhất trong sơ đồ Use Case.',
     'Học viên — Chơi mini-game'],
    ['Generalization', 'B --|> A',
     'Nét thẳng + mũi tên rỗng (—▷)',
     'Kế thừa: Actor B (hoặc UC con) thừa hưởng mọi quan hệ của Actor A (UC cha).',
     'Học viên kế thừa Khách: làm được mọi thứ Khách làm + thêm chức năng riêng'],
    ['Include', 'X ..> Y : <<include>>',
     'Nét đứt + mũi tên + stereotype <<include>>',
     'Bắt buộc: Use Case X LUÔN gọi Use Case Y trong quá trình thực thi. Y không thể bỏ qua.',
     'Chơi mini-game <<include>> Cập nhật streak (mọi lần chơi xong đều update streak)'],
    ['Extend', 'X ..> Y : <<extend>>',
     'Nét đứt + mũi tên + stereotype <<extend>>',
     'Tùy chọn: Use Case X mở rộng Y CHỈ KHI điều kiện thỏa mãn. Y vẫn chạy được nếu không có X.',
     'Nhận thưởng milestone <<extend>> Cập nhật streak (chỉ kích hoạt khi streak chạm mốc 3/7/14...)'],
])

# ---------- 7. Giải thích biểu đồ chính ----------
d.add_paragraph('7. Giải thích biểu đồ Use Case chính', style='Heading 1')

d.add_paragraph('7.1. Phân vùng Actor', style='Heading 2')
d.add_paragraph(
    'Bên TRÁI biểu đồ là 2 actor chính (Khách và Học viên), nối với các use case họ trực tiếp '
    'thao tác. Bên PHẢI là 4 actor phụ (hệ thống ngoài) — chúng nhận yêu cầu từ use case nội bộ '
    'và trả về kết quả. Cách bố trí 2 phía giúp nhìn rõ ranh giới giữa "người dùng → app" và '
    '"app → dịch vụ ngoài".',
    style='Normal',
)

d.add_paragraph('7.2. Các quan hệ kế thừa', style='Heading 2')
d.add_paragraph(
    'Học viên --|> Khách: bất kỳ chức năng nào Khách làm được (đăng ký, đăng nhập, quên mật '
    'khẩu) thì Học viên cũng làm được mà không cần vẽ lại đường nối. Đây là cách giảm nhiễu '
    'biểu đồ khi 2 actor có giao thoa chức năng.',
    style='Normal',
)

d.add_paragraph('7.3. Các quan hệ <<include>>', style='Heading 2')
d.add_paragraph(
    'UC-06 (Chơi mini-game) <<include>> UC-07 (Xem kết quả) và UC-19 (Cập nhật streak): khi '
    'chơi xong 1 game, cả hai use case này LUÔN được kích hoạt — không có lựa chọn bỏ qua. '
    'UC-10 (Mua pack) <<include>> UC-02 (Đăng nhập): không thể mua nếu chưa đăng nhập.',
    style='Normal',
)

d.add_paragraph('7.4. Các quan hệ <<extend>>', style='Heading 2')
d.add_paragraph(
    'UC-09 (Mở khóa level) <<extend>> UC-07: chỉ chạy KHI sao đạt ngưỡng (≥2 cho 3 game đầu, '
    '=3 cho Memory). UC-20 (Nhận thưởng milestone) <<extend>> UC-19: chỉ chạy KHI streak '
    'chạm mốc. UC-18 (Phát âm) <<extend>> UC-06: chỉ kích hoạt khi user nhấn nút loa trong '
    'Quiz. UC-10 (Mua pack) <<extend>> UC-08: trong màn chọn pack, có pack chưa sở hữu thì '
    'mới có nút Mua.',
    style='Normal',
)

d.add_paragraph('7.5. UC-19 thuộc về ai?', style='Heading 2')
d.add_paragraph(
    'UC-19 (Cập nhật streak) không có actor người dùng nối trực tiếp vì nó chạy NGẦM trong '
    'transaction lưu kết quả game (xem PHẦN I — mục 8). Nó được kích hoạt qua quan hệ '
    '<<include>> từ UC-06, nên trong sơ đồ chỉ thể hiện liên kết tới Firestore (actor phụ) '
    'mà không có association trực tiếp với Học viên.',
    style='Normal',
)

# ---------- 8. Hướng dẫn render ----------
d.add_paragraph('8. Hướng dẫn render PlantUML', style='Heading 1')
d.add_paragraph(
    '1. Online: vào https://www.plantuml.com/plantuml/uml/ → paste code → nhấn "Submit" → '
    'tải ảnh PNG/SVG. 2. VS Code: cài extension "PlantUML" của jebbs → mở file .puml → Alt+D '
    'để preview. 3. IntelliJ/Android Studio: cài plugin "PlantUML Integration" → tạo file '
    '.puml → preview pane bên phải. 4. Xuất ảnh chất lượng cao: chọn định dạng SVG khi render '
    'để giữ nét rõ khi phóng to trong Word.',
    style='Normal',
)

# Save
try:
    d.save(path)
    print('Saved:', path)
except PermissionError:
    alt = path.replace('.docx', '_v2.docx')
    d.save(alt)
    print('Original locked. Saved to:', alt)
