# -*- coding: utf-8 -*-
"""Xuat danh sach thu vien su dung trong VocabQuest sang .docx."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = r"C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_ThuVien.docx"

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x4A, 0x44, 0xCC)


def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic


def make_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            t.rows[r_idx].cells[c_idx].text = str(val)
    doc.add_paragraph()


# ===== Cover =====
title = doc.add_heading('VocabQuest', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Danh sách thư viện sử dụng trong project')
r.italic = True
r.font.size = Pt(14)
doc.add_paragraph()

# ===== 1. Firebase =====
h1('1. Firebase (Backend & Cloud)')
make_table(
    ['Package', 'Version', 'Mục đích', 'Dùng ở đâu'],
    [
        ['firebase_core', '^2.24.2',
         'Khởi tạo Firebase, bắt buộc để dùng các Firebase SDK khác',
         'main.dart — Firebase.initializeApp()'],
        ['firebase_auth', '^4.15.3',
         'Xác thực người dùng: đăng ký, đăng nhập, đăng xuất, reset mật khẩu',
         'auth_service.dart'],
        ['cloud_firestore', '^4.13.6',
         'Database NoSQL realtime — lưu users, game_results, leaderboard; '
         'dùng transaction atomic cho streak/reward',
         'firestore_service.dart'],
        ['firebase_storage', '^11.5.6',
         'Cloud storage (ảnh). Hiện KHÔNG dùng thực tế — avatar đã chuyển '
         'base64 vào Firestore vì cần Blaze plan',
         'storage_service.dart (dự phòng)'],
    ],
)

# ===== 2. State & Local Storage =====
h1('2. State Management & Local Storage')
make_table(
    ['Package', 'Version', 'Mục đích', 'Dùng ở đâu'],
    [
        ['provider', '^6.1.1',
         'Quản lý state reactive, rebuild widget khi state đổi',
         'UserProvider, SettingsProvider, GameProvider, FavoritesProvider'],
        ['shared_preferences', '^2.2.2',
         'Lưu key-value đơn giản trên device (offline) — settings, '
         'favorites, first-time flag',
         'local_storage.dart'],
    ],
)

# ===== 3. Icons =====
h1('3. Icons')
make_table(
    ['Package', 'Version', 'Mục đích', 'Dùng ở đâu'],
    [
        ['cupertino_icons', '^1.0.6',
         'Icon kiểu iOS (Apple design)',
         'Ít dùng, mặc định Flutter có sẵn'],
        ['font_awesome_flutter', '^10.6.0',
         'Bộ icon Font Awesome — trophy, coins, gamepad, puzzle-piece...',
         'Profile stats, game icons'],
        ['lucide_icons', '^0.257.0',
         'Icon Lucide (thiết kế hiện đại, đẹp) — dùng xuyên suốt UI',
         'Home, Settings, Game menu, nút back...'],
    ],
)

# ===== 4. UI & Animation =====
h1('4. UI & Animation')
make_table(
    ['Package', 'Version', 'Mục đích', 'Dùng ở đâu'],
    [
        ['google_fonts', '^6.1.0',
         'Font Google miễn phí (Poppins) cho app đẹp hơn mặc định',
         'theme.dart'],
        ['flutter_svg', '^2.0.9',
         'Render file SVG',
         'Dự phòng khi cần icon/illustration vector'],
        ['lottie', '^2.7.0',
         'Phát animation Lottie (JSON dạng After Effects)',
         'Dự phòng — chưa dùng animation cụ thể'],
        ['cached_network_image', '^3.3.0',
         'Load ảnh từ URL + cache trên máy (không tải lại)',
         'UserAvatar cho URL http (không dùng cho base64)'],
        ['percent_indicator', '^4.2.3',
         'Progress bar tròn/dọc đẹp',
         'Profile XP bar, Pack tiến độ, Daily challenge'],
        ['confetti', '^0.7.0',
         'Hiệu ứng ăn mừng (pháo hoa, bông nhiều màu)',
         'GameResultScreen khi ≥2 sao hoặc milestone streak'],
        ['flutter_animate', '^4.3.0',
         'Chain animation dễ dùng: .animate().fadeIn().slideY()',
         'Gần như tất cả screens — hero banner, cards, list items'],
        ['shimmer', '^3.0.0',
         'Hiệu ứng loading skeleton / glow sweep',
         'Featured game card (shimmer sweep mỗi 3s)'],
        ['animated_text_kit', '^4.2.2',
         'Text có animation (typewriter, wave...)',
         'Dự phòng, chưa dùng nhiều'],
    ],
)

# ===== 5. Audio & Speech =====
h1('5. Audio & Speech')
make_table(
    ['Package', 'Version', 'Mục đích', 'Dùng ở đâu'],
    [
        ['audioplayers', '^5.2.1',
         'Phát file MP3 (nhạc nền, SFX đúng/sai, click)',
         'audio_service.dart — playClick, playCorrect, playWrong, bgMusic'],
        ['flutter_tts', '^3.8.5',
         'Text-to-speech — đọc từ tiếng Anh chuẩn (en-US)',
         'QuizGame — nút phát âm từ hiện tại'],
    ],
)

# ===== 6. Notification =====
h1('6. Notification')
make_table(
    ['Package', 'Version', 'Mục đích', 'Dùng ở đâu'],
    [
        ['flutter_local_notifications', '^17.2.4',
         'Hiển thị thông báo local (không cần server) — schedule, channel, '
         'exact alarm',
         'notification_service.dart — nhắc nhở học hàng ngày'],
        ['timezone', '^0.9.4',
         'Quản lý múi giờ chính xác (Asia/Ho_Chi_Minh) cho scheduled '
         'notifications',
         'notification_service.dart — tz.setLocalLocation()'],
    ],
)

# ===== 7. Media & Input =====
h1('7. Media & Input')
make_table(
    ['Package', 'Version', 'Mục đích', 'Dùng ở đâu'],
    [
        ['image_picker', '^1.0.7',
         'Chọn ảnh từ camera hoặc thư viện thiết bị',
         'EditProfileScreen — pick avatar, resize 400×400 quality 75'],
    ],
)

# ===== 8. Utilities =====
h1('8. Utilities')
make_table(
    ['Package', 'Version', 'Mục đích', 'Dùng ở đâu'],
    [
        ['intl', '^0.19.0',
         'Format ngày tháng, số, locale',
         'Format lịch sử chơi, thời gian trong leaderboard'],
        ['uuid', '^4.2.1',
         'Sinh ID duy nhất (v4)',
         'Dự phòng — thực tế Firestore auto-gen ID'],
    ],
)

# ===== 9. Dev dependencies =====
h1('9. Dev dependencies (chỉ dùng khi develop)')
make_table(
    ['Package', 'Version', 'Mục đích'],
    [
        ['flutter_test', 'sdk', 'Framework test unit/widget'],
        ['flutter_lints', '^3.0.0',
         'Rule linter — warn code không theo convention'],
        ['flutter_launcher_icons', '^0.14.4',
         'Tool auto generate icon app (Android mipmap) từ 1 ảnh PNG'],
    ],
)

# ===== Summary =====
h1('Tổng kết nhóm')
make_table(
    ['Nhóm', 'Số package', 'Mục tiêu'],
    [
        ['Firebase', '4', 'Auth + Database + Storage'],
        ['State & Storage', '2', 'Quản lý state reactive + lưu local'],
        ['Icons', '3', '3 bộ icon đa dạng'],
        ['UI & Animation', '9',
         'Tạo UI đẹp, mượt, có feedback trực quan'],
        ['Audio & Speech', '2', 'Hiệu ứng âm thanh + phát âm từ'],
        ['Notification', '2', 'Nhắc nhở học hàng ngày'],
        ['Media Input', '1', 'Chọn ảnh avatar'],
        ['Utilities', '2', 'Helper format / ID'],
        ['Dev only', '3', 'Test, lint, build tool'],
        ['TỔNG', '28', ''],
    ],
)

# ===== Core deps =====
h1('Thư viện core không thể thiếu')
para('Nếu cắt gọn app, 5 nhóm thư viện sau là bắt buộc:')
p = doc.add_paragraph()
p.style = 'List Number'
for text in [
    'firebase_core + firebase_auth + cloud_firestore — toàn bộ backend',
    'provider — state management',
    'shared_preferences — settings offline',
    'flutter_local_notifications + timezone — nhắc nhở (feature cốt lõi)',
    'image_picker — đổi avatar',
]:
    doc.add_paragraph(text, style='List Number')

para(
    'Các thư viện UI/animation có thể thay thế bằng Flutter built-in '
    'nhưng UX sẽ kém mượt hơn.',
    italic=True,
)

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
