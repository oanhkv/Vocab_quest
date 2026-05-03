import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from copy import deepcopy

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_Logic_Game.docx'
d = Document(path)

# Helper: add table with header row formatted similar to existing tables
def add_table(rows):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(val)
            if r_idx == 0:
                run.bold = True
    return t

# ===== Section 11: Memory Game logic chi tiết =====
d.add_paragraph('', style='Normal')
d.add_paragraph('11. Memory Game — Phân tích logic chi tiết', style='Heading 1')
d.add_paragraph(
    'File lib/screens/games/memory_game.dart — bản cập nhật mới nhất. '
    'Phần này bổ sung & hiệu chỉnh cho mục 2.4 ở trên (số cặp/round, cơ chế điểm, tính sao).',
    style='Normal',
)

# 11.1 Quy trình hoạt động
d.add_paragraph('11.1. Quy trình hoạt động', style='Heading 2')
d.add_paragraph(
    'Khởi tạo (initState → _loadInitial): gọi GameProvider.startGame(); nạp pool từ widget.words '
    'hoặc JsonService.loadVocab(level); _dedupeWords khử trùng lặp theo id (fallback word|meaning lowercase); '
    'shuffle pool; _totalRounds = pool.length ~/ pairsPerRound, clamp [1, 3] để đảm bảo các round không lặp từ; '
    '_setupRound() → _startTimer() (đếm ngược 90s).',
    style='Normal',
)
d.add_paragraph(
    'Setup round (_setupRound): lấy pairsPerRound từ vựng đầu tiên trong _remainingPool, xoá khỏi pool. '
    'Mỗi từ tạo 2 thẻ — 1 thẻ word (EN), 1 thẻ meaning (VI), cùng vocabId. Trộn thẻ rồi reset _firstPick / '
    '_secondPick / _failedPairIds.',
    style='Normal',
)
d.add_paragraph(
    'Tap thẻ (_onTapCard → _checkMatch): bỏ qua tap khi !_canTap, đã lật, hoặc đã match. Lưu thẻ thứ 1, '
    'tap thẻ thứ 2 thì khoá _canTap = false và check match. Match khi vocabId trùng AND khác loại '
    '(first.isWord != second.isWord). Đúng → cả 2 isMatched = true, mở lại _canTap, playCorrect. Sai → '
    'thêm vocabId vào _failedPairIds, playWrong, sau 800ms lật úp lại. Khi mọi thẻ đã match → _nextRound() '
    '(delay 900ms); hết round → _endGame().',
    style='Normal',
)
d.add_paragraph(
    'Kết thúc (_endGame): huỷ timer, build MemoryStats, tính memoryStars. GameProvider.finishGame(...) trả '
    'outcome (result + streak). Cập nhật user (score/coins/XP/streak). Mở level mới CHỈ khi memoryStars >= 3 '
    '→ gọi recordLevelComplete. Điều hướng sang GameResultScreen với overrideStars: memoryStars.',
    style='Normal',
)

# 11.2 Luật chơi
d.add_paragraph('11.2. Luật chơi', style='Heading 2')
add_table([
    ['Thuộc tính', 'Giá trị'],
    ['Thời gian', '90 giây cho cả ván (chung cho mọi round)'],
    ['Số cặp/round', 'Level 1 = 4 cặp, Level 2 = 6 cặp, Level 3 = 8 cặp (default 6)'],
    ['Số round', 'pool.length / pairsPerRound, clamp 1–3'],
    ['Layout grid', '4 cặp → 4 cột × 2 hàng; 6 cặp → 3 × 4; 8 cặp → 4 × 4'],
    ['Thẻ', 'Mỗi từ vựng = 2 thẻ (EN word + VI meaning) cùng vocabId'],
    ['Điều kiện match', 'Cùng vocabId AND khác loại (1 word + 1 meaning)'],
    ['Thẻ sai', 'Úp lại sau 800ms, vocabId được đánh dấu "đã từng fail"'],
    ['Thoát giữa chừng', 'Mất tiến trình (dialog xác nhận)'],
])

# 11.3 Cơ chế tính điểm
d.add_paragraph('11.3. Cơ chế tính điểm (max 100)', style='Heading 2')
d.add_paragraph(
    'Công thức (_score):  raw = firstAttemptMatches + 0.5 × afterFailMatches;  '
    'score = round(raw / totalPairs × 100), cap ở 100. '
    'Trong đó totalPairs = totalRounds × pairsPerRound.',
    style='Normal',
)
add_table([
    ['Trường hợp', 'Điểm'],
    ['Cặp ghép đúng ngay lần đầu', '100 / totalPairs điểm'],
    ['Cặp ghép đúng sau khi đã từng sai (vocabId nằm trong _failedPairIds)', 'Nửa điểm so với lần đầu'],
    ['Tổng điểm tối đa', '100 (cap)'],
])
d.add_paragraph(
    'Sau mỗi match, _syncProviderScore chỉ gửi DELTA sang GameProvider để tránh đếm trùng. '
    'GameProvider.addCorrect() chỉ tăng cho lần match đầu; addWrong() cộng cho mỗi lần lật sai cặp.',
    style='Normal',
)

# 11.4 Tính sao
d.add_paragraph('11.4. Tính sao (chỉ xét khi đã ghép xong toàn bộ)', style='Heading 2')
add_table([
    ['Điều kiện', 'Sao'],
    ['Chưa ghép xong (hết giờ)', '0 sao'],
    ['Ghép xong, timeLeft >= 50s', '3 sao — điều kiện DUY NHẤT mở level mới'],
    ['Ghép xong, timeLeft >= 30s', '2 sao'],
    ['Ghép xong, timeLeft >= 1s', '1 sao'],
])
d.add_paragraph(
    'Lưu ý: Memory game KHÔNG dùng star rating chuẩn theo accuracy như các game khác — nó dùng overrideStars '
    'truyền vào GameResultScreen. Sao đo TỐC ĐỘ (timeLeft), điểm đo ĐỘ CHÍNH XÁC (số lần sai). '
    'Người chơi có thể đạt 3 sao dù lật sai nhiều, miễn vẫn xong toàn bộ trong thời gian quy định.',
    style='Normal',
)

# 11.5 Hằng số quan trọng
d.add_paragraph('11.5. Hằng số quan trọng trong code', style='Heading 2')
add_table([
    ['Hằng số', 'Giá trị', 'Ý nghĩa'],
    ['_totalTime', '90', 'Tổng thời gian (giây) cho cả ván'],
    ['_maxScore', '100', 'Điểm tối đa sau scale'],
    ['_starThreshold3', '50', 'timeLeft tối thiểu để đạt 3 sao'],
    ['_starThreshold2', '30', 'timeLeft tối thiểu để đạt 2 sao'],
    ['_starThreshold1', '1', 'timeLeft tối thiểu để đạt 1 sao'],
    ['flip-back delay', '800ms', 'Delay úp thẻ lại khi sai'],
    ['next-round delay', '900ms', 'Delay chuyển round'],
    ['flip animation', '400ms', 'Thời gian animation lật thẻ 3D (rotateY)'],
])

# 11.6 Khác biệt với mục 2.4
d.add_paragraph('11.6. Hiệu chỉnh so với mục 2.4 cũ', style='Heading 2')
add_table([
    ['Thuộc tính', 'Mục 2.4 (cũ)', 'Code hiện tại (đúng)'],
    ['Số cặp/round', 'Cố định 6', 'Theo level: 4 / 6 / 8'],
    ['Layout grid', 'Chỉ 3 × 4', '4×2 (L1) / 3×4 (L2) / 4×4 (L3)'],
    ['Điểm/cặp', '+20 cố định', '100/totalPairs (scale theo tổng cặp)'],
    ['Cơ chế retry', 'Không cộng điểm', 'Cộng nửa điểm (không phải 0)'],
    ['Max điểm/game', '360', '100 (cap)'],
    ['Tính sao', 'Theo accuracy (chuẩn chung)', 'Theo timeLeft (override riêng)'],
    ['Mở level mới', '>= 2 sao', 'Yêu cầu 3 sao'],
])

out_path = path
d.save(out_path)
print('Saved:', out_path)
