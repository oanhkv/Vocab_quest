# -*- coding: utf-8 -*-
"""Xuat 2 bieu do phan cap chuc nang (BPC) cua VocabQuest."""

import os
import matplotlib
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

matplotlib.rcParams['font.family'] = ['Segoe UI', 'Arial', 'DejaVu Sans']
matplotlib.rcParams['axes.unicode_minus'] = False

OUT_DIR = r"C:\Users\Kieu Anh\Desktop\CD2"
IMG1 = os.path.join(OUT_DIR, '_bpc_user.png')
IMG2 = os.path.join(OUT_DIR, '_bpc_game.png')
DOCX = os.path.join(OUT_DIR, 'VocabQuest_BPC.docx')

# ========= Hierarchy data =========

DIAGRAM_1 = {
    'root': 'VOCABQUEST — Phân hệ\nTài khoản & Cá nhân hóa',
    'modules': [
        {
            'code': '1',
            'title': 'Xác thực',
            'children': [
                '1.1 Đăng ký',
                '1.2 Đăng nhập',
                '1.3 Đăng xuất',
                '1.4 Quên mật khẩu',
            ],
        },
        {
            'code': '2',
            'title': 'Hồ sơ cá nhân',
            'children': [
                '2.1 Xem hồ sơ',
                '2.2 Sửa tên hiển thị',
                '2.3 Upload avatar',
                '2.4 Thống kê cá nhân',
            ],
        },
        {
            'code': '3',
            'title': 'Cài đặt ứng dụng',
            'children': [
                '3.1 Giao diện tối / sáng',
                '3.2 Âm thanh & nhạc nền',
                '3.3 Thông báo & giờ nhắc',
                '3.4 Ngôn ngữ VI / EN',
                '3.5 Đánh giá & chia sẻ',
            ],
        },
        {
            'code': '4',
            'title': 'Yêu thích',
            'children': [
                '4.1 Thêm game yêu thích',
                '4.2 Danh sách đã lưu',
                '4.3 Gỡ khỏi yêu thích',
            ],
        },
    ],
}

DIAGRAM_2 = {
    'root': 'VOCABQUEST — Phân hệ\nChơi game & Học tập',
    'modules': [
        {
            'code': '5',
            'title': 'Kho từ vựng',
            'children': [
                '5.1 Xem danh sách pack',
                '5.2 Mua pack bằng coin',
                '5.3 Xem chi tiết pack',
            ],
        },
        {
            'code': '6',
            'title': 'Chọn game',
            'children': [
                '6.1 Menu mini game',
                '6.2 Chọn pack từ vựng',
                '6.3 Chọn level (map)',
            ],
        },
        {
            'code': '7',
            'title': 'Mini Games',
            'children': [
                '7.1 Nối từ (Matching)',
                '7.2 Trắc nghiệm (Quiz)',
                '7.3 Xếp chữ (Puzzle)',
                '7.4 Lật thẻ (Memory)',
            ],
        },
        {
            'code': '8',
            'title': 'Kết quả & Phần thưởng',
            'children': [
                '8.1 Chấm sao 1-3',
                '8.2 Tính Coin & XP',
                '8.3 Thưởng mở level',
                '8.4 Milestone streak',
                '8.5 Animation ăn mừng',
            ],
        },
        {
            'code': '9',
            'title': 'Thống kê & Thành tích',
            'children': [
                '9.1 Lịch sử chơi',
                '9.2 Bảng xếp hạng',
                '9.3 User Level (XP)',
                '9.4 Streak hiện tại & kỷ lục',
            ],
        },
    ],
}


# ========= Drawing helpers =========

def draw_box(ax, cx, cy, text, *, width, height,
             face, edge, text_color, bold=False):
    """Ve 1 box co bo goc tron."""
    box = FancyBboxPatch(
        (cx - width / 2, cy - height / 2),
        width, height,
        boxstyle='round,pad=0.02,rounding_size=0.08',
        linewidth=1.5,
        edgecolor=edge,
        facecolor=face,
        zorder=2,
    )
    ax.add_patch(box)
    ax.text(
        cx, cy, text,
        ha='center', va='center',
        fontsize=10 if not bold else 12,
        fontweight='bold' if bold else 'normal',
        color=text_color,
        zorder=3,
    )


def draw_line(ax, x1, y1, x2, y2, color='#9D97FF'):
    ax.plot([x1, x2], [y1, y2], color=color, lw=1.2, zorder=1)


def render_diagram(diagram, out_path):
    modules = diagram['modules']
    n = len(modules)
    # Tinh width can thiet theo so module + so children max
    max_child = max(len(m['children']) for m in modules)
    slot_w = 3.6  # do rong 1 slot module
    fig_w = max(slot_w * n + 1.5, 14)
    fig_h = 3 + max_child * 0.85 + 1.5
    fig, ax = plt.subplots(figsize=(fig_w, fig_h))
    ax.set_xlim(0, fig_w)
    ax.set_ylim(0, fig_h)
    ax.set_aspect('equal')
    ax.axis('off')

    # Root
    root_x = fig_w / 2
    root_y = fig_h - 0.8
    draw_box(
        ax, root_x, root_y, diagram['root'],
        width=5.2, height=1.1,
        face='#6C63FF', edge='#4A44CC', text_color='white', bold=True,
    )

    # Horizontal backbone line (neu >= 2 module)
    mod_y = root_y - 1.6
    # Compute module x positions
    total_w = slot_w * n
    start_x = (fig_w - total_w) / 2 + slot_w / 2
    mod_xs = [start_x + i * slot_w for i in range(n)]

    # Vertical drop from root
    drop_y = root_y - 0.55 - 0.4  # below root
    draw_line(ax, root_x, root_y - 0.55, root_x, drop_y, color='#6C63FF')
    # Horizontal connector at drop_y across all modules
    draw_line(ax, mod_xs[0], drop_y, mod_xs[-1], drop_y, color='#6C63FF')
    for mx in mod_xs:
        draw_line(ax, mx, drop_y, mx, mod_y + 0.4, color='#6C63FF')

    # Modules
    for m, mx in zip(modules, mod_xs):
        draw_box(
            ax, mx, mod_y, f"{m['code']}. {m['title']}",
            width=slot_w * 0.88, height=0.75,
            face='#9D97FF', edge='#6C63FF', text_color='white', bold=True,
        )

        # Children dọc xuống
        child_y = mod_y - 1.15
        for child in m['children']:
            # line tu module xuong box con
            draw_line(ax, mx, mod_y - 0.4, mx, child_y + 0.3, color='#BDB5FF')
            draw_box(
                ax, mx, child_y, child,
                width=slot_w * 0.88, height=0.6,
                face='white', edge='#6C63FF', text_color='#2D2D3F',
            )
            child_y -= 0.78

    plt.tight_layout()
    plt.savefig(out_path, dpi=160, bbox_inches='tight', facecolor='white')
    plt.close()


# ========= Render both =========

render_diagram(DIAGRAM_1, IMG1)
render_diagram(DIAGRAM_2, IMG2)
print('Rendered diagrams:', IMG1, IMG2)

# ========= Build docx =========

doc = Document()

# Title
title = doc.add_heading('VocabQuest', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Biểu đồ phân cấp chức năng (BPC)')
run.italic = True
run.font.size = Pt(14)

intro = doc.add_paragraph(
    'Hệ thống VocabQuest được chia thành 2 phân hệ chính để dễ nhìn và '
    'quản lý: (1) Tài khoản & Cá nhân hóa và (2) Chơi game & Học tập. '
    'Hai phân hệ liên kết với nhau qua node gốc VOCABQUEST ở cấp 0.'
)

# Diagram 1
h = doc.add_heading('Biểu đồ 1 — Phân hệ Tài khoản & Cá nhân hóa', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
doc.add_paragraph(
    'Phân hệ quản lý người dùng: xác thực tài khoản, thông tin cá nhân, '
    'cài đặt ứng dụng và danh sách game yêu thích.'
)
doc.add_picture(IMG1, width=Cm(17))
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.add_run('Hình 1: BPC phân hệ Tài khoản & Cá nhân hóa').italic = True

doc.add_page_break()

# Diagram 2
h = doc.add_heading('Biểu đồ 2 — Phân hệ Chơi game & Học tập', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
doc.add_paragraph(
    'Phân hệ cốt lõi: quản lý kho từ vựng, luồng chọn-chơi 4 mini game, '
    'cơ chế phần thưởng và thống kê thành tích.'
)
doc.add_picture(IMG2, width=Cm(17))
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.add_run('Hình 2: BPC phân hệ Chơi game & Học tập').italic = True

doc.add_page_break()

# Liên kết giữa 2 phân hệ
h = doc.add_heading('Liên kết giữa 2 phân hệ', level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
doc.add_paragraph(
    'Hai phân hệ không độc lập mà trao đổi dữ liệu qua tài khoản người '
    'dùng (Firestore users/{uid}):'
)

link_rows = [
    ('Xác thực (1)', '→', 'Mọi phân hệ',
     'Mở khóa toàn bộ chức năng sau khi đăng nhập'),
    ('Hồ sơ (2)', '←→', 'Kết quả & Thưởng (8)',
     'Điểm, coin, XP, streak, level user cập nhật ngược vào hồ sơ'),
    ('Cài đặt (3)', '→', 'Mini Games (7)',
     'Âm thanh, ngôn ngữ, thông báo ảnh hưởng trải nghiệm game'),
    ('Yêu thích (4)', '→', 'Chọn game (6)',
     'User có thể chơi nhanh từ danh sách yêu thích'),
    ('Kho từ vựng (5)', '→', 'Mini Games (7)',
     'Pack đã mua/sở hữu cung cấp dữ liệu từ vựng cho 4 game'),
    ('Kết quả (8)', '→', 'Thống kê (9)',
     'Mọi kết quả ghi vào lịch sử, leaderboard và cập nhật streak'),
]

table = doc.add_table(rows=1 + len(link_rows), cols=4)
table.style = 'Light Grid Accent 1'
headers = ['Nguồn', 'Chiều', 'Đích', 'Ý nghĩa']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for r_idx, row in enumerate(link_rows, start=1):
    for c_idx, val in enumerate(row):
        table.rows[r_idx].cells[c_idx].text = val

doc.add_paragraph()
note = doc.add_paragraph()
note.add_run('Ghi chú kỹ thuật: ').bold = True
note.add_run(
    'Tất cả dữ liệu người dùng (profile, tiến độ, streak, coin, XP) được '
    'lưu đồng bộ trên Firestore collection users. Cài đặt cục bộ (giao '
    'diện, âm thanh, ngôn ngữ, giờ nhắc, yêu thích) lưu trong '
    'SharedPreferences để hoạt động offline.'
)

doc.save(DOCX)
print('Saved:', DOCX)
