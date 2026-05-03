import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TongHop_v2.docx'
d = Document(path)

# Check Actor & UC tables
for t in d.tables:
    h = [c.text.strip() for c in t.rows[0].cells]
    if h == ['STT', 'Tên Actor', 'Loại', 'Mô tả', 'Use case liên quan']:
        print(f'=== ACTOR TABLE ({len(t.rows)-1} actor) ===')
        for row in t.rows:
            print('  |', ' | '.join(c.text.strip()[:35] for c in row.cells))
    elif h == ['Mã', 'Tên Use Case', 'Actor chính', 'Mô tả ngắn']:
        print(f'\n=== USE CASE TABLE ({len(t.rows)-1} UC) ===')
        for row in t.rows:
            print('  |', ' | '.join(c.text.strip()[:35] for c in row.cells))

# Check PlantUML #1 first lines
print('\n=== PLANTUML #1 (first 6 lines) ===')
for p in d.paragraphs:
    if '@startuml VocabQuest_UseCase' in p.text:
        for line in p.text.splitlines()[:6]:
            print('  ', line)
        break

# Check headings 7.x
print('\n=== HEADINGS 7.x ===')
for p in d.paragraphs:
    if p.style.name == 'Heading 2' and p.text.strip().startswith('7.'):
        print(' -', p.text.strip())
