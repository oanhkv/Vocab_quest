import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TestCases_v2.docx'
d = Document(path)

# Show first TC table
for i, t in enumerate(d.tables):
    if len(t.columns) == 8:
        print(f'=== TABLE[{i}] cols={len(t.columns)} rows={len(t.rows)} ===')
        for r_idx, row in enumerate(t.rows[:3]):
            cells = [c.text.replace('\n', ' | ')[:60] for c in row.cells]
            print(f'  r{r_idx}: {cells}')
        print()

# Count statuses across all TC tables
pass_n = fail_n = blocked_n = 0
fail_ids, blocked_ids = [], []
for t in d.tables:
    if len(t.columns) != 8:
        continue
    h = [c.text.strip() for c in t.rows[0].cells]
    if h[0] != 'ID':
        continue
    for row in t.rows[1:]:
        tc_id = row.cells[0].text.strip()
        status = row.cells[6].text.strip()
        if status == 'Pass': pass_n += 1
        elif status == 'Fail': fail_n += 1; fail_ids.append(tc_id)
        elif status == 'Blocked': blocked_n += 1; blocked_ids.append(tc_id)
print(f'Pass={pass_n}, Fail={fail_n} {fail_ids}, Blocked={blocked_n} {blocked_ids}')
