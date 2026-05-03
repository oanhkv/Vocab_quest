"""
Bổ sung các chương còn thiếu để VocabQuest_TongHop_v2.docx trở thành báo cáo phân tích
hệ thống đầy đủ. Append các PHẦN V → XV vào cuối file.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TongHop_v2.docx'
d = Document(path)


def add_table(rows, header_bold=True):
    t = d.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            if r_idx == 0 and header_bold:
                run.bold = True
    return t


def add_code_block(code):
    p = d.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    return p


def H1(t): d.add_paragraph(t, style='Heading 1')
def H2(t): d.add_paragraph(t, style='Heading 2')
def H3(t): d.add_paragraph(t, style='Heading 3')
def P(t):  d.add_paragraph(t, style='Normal')


# =============================================================
# PHẦN V — TỔNG QUAN ĐỀ TÀI
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN V — TỔNG QUAN ĐỀ TÀI', style='Title')

H1('1. Giới thiệu')
P('VocabQuest là ứng dụng học từ vựng tiếng Anh trên thiết bị di động được phát triển '
  'bằng Flutter (frontend) và Firebase (backend). Ứng dụng kết hợp 4 mini-game tương tác '
  '(Quiz, Matching, Memory, Word Puzzle) với hệ thống điểm thưởng – streak – cấp độ để '
  'biến việc học từ vựng thành trải nghiệm chơi game (gamification). Hỗ trợ đa nền tảng: '
  'Android, iOS, Web, Windows, macOS, Linux.')

H1('2. Lý do chọn đề tài')
P('Học từ vựng là khâu nền tảng nhưng dễ gây nhàm chán nhất khi học ngoại ngữ. Các app '
  'flashcard truyền thống thường thiếu yếu tố thúc đẩy duy trì thói quen. VocabQuest '
  'giải quyết vấn đề này bằng cách: (1) đa dạng hoá hoạt động tương tác qua 4 game khác '
  'nhau cho cùng một bộ từ; (2) hệ thống streak và milestone reward để duy trì thói quen '
  'hàng ngày; (3) cơ chế mở khóa level + leaderboard tạo động lực cạnh tranh.')

H1('3. Mục tiêu đề tài')
add_table([
    ['Mục tiêu', 'Mô tả'],
    ['Về sản phẩm',
     'Xây dựng hoàn chỉnh ứng dụng Flutter chạy đa nền tảng, đăng ký/đăng nhập '
     'Firebase, có 4 mini-game hoạt động độc lập, lưu trữ và đồng bộ dữ liệu '
     'realtime qua Firestore.'],
    ['Về kỹ thuật',
     'Áp dụng Provider pattern cho state management, transaction Firestore atomic '
     'cho lưu kết quả game, thiết kế Security Rules an toàn ở backend.'],
    ['Về trải nghiệm',
     'UI/UX tối ưu cho cả tablet và phone, hỗ trợ đa ngôn ngữ (VI/EN), '
     'dark/light mode, animation mượt 60fps.'],
])

H1('4. Phạm vi đề tài')
add_table([
    ['Phạm vi', 'Bao gồm', 'Không bao gồm'],
    ['Chức năng',
     'Auth (đăng ký/đăng nhập/quên mật khẩu), 4 mini-game, leaderboard, profile, '
     'streak, level system, shop, settings, notification.',
     'Chat, social network, voice recognition, AI sinh câu hỏi.'],
    ['Nền tảng',
     'Android (chính), iOS, Web (test).',
     'Windows/macOS/Linux desktop (chỉ build, không test sâu).'],
    ['Backend',
     'Firebase Auth, Cloud Firestore, Firebase Storage (avatar — đang dùng base64).',
     'Cloud Functions, BigQuery, ML Kit, custom backend server.'],
])

H1('5. Đối tượng người dùng mục tiêu')
add_table([
    ['Phân nhóm', 'Đặc điểm', 'Nhu cầu chính'],
    ['Học sinh – Sinh viên (15–22)',
     'Học tiếng Anh chương trình phổ thông, đại học; thường xuyên dùng smartphone.',
     'Học từ vựng theo chủ đề SGK, ôn thi chứng chỉ A1–B2.'],
    ['Người đi làm (23–35)',
     'Cải thiện tiếng Anh giao tiếp công việc, ôn TOEIC/IELTS.',
     'Học theo gói chủ đề công việc, học ngắn 5–10 phút/ngày.'],
    ['Người mới bắt đầu',
     'Chưa có nền tảng, sợ "học chính thống".',
     'Pack beginner miễn phí, giao diện thân thiện, không áp lực.'],
])


# =============================================================
# PHẦN VI — PHÂN TÍCH YÊU CẦU
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN VI — PHÂN TÍCH YÊU CẦU HỆ THỐNG', style='Title')

H1('1. Yêu cầu chức năng (Functional Requirements)')
P('Tổng hợp các yêu cầu chức năng theo nhóm. Mỗi yêu cầu gắn với 1+ Use Case ở PHẦN IV.')
add_table([
    ['Mã', 'Nhóm', 'Yêu cầu', 'UC liên quan'],
    ['FR-01', 'Tài khoản', 'Người dùng đăng ký bằng email/mật khẩu, được tặng 100 coin khởi điểm.', 'UC-01'],
    ['FR-02', 'Tài khoản', 'Đăng nhập bằng email/mật khẩu; có chức năng quên mật khẩu qua email.', 'UC-02'],
    ['FR-03', 'Tài khoản', 'Đăng xuất, xoá session local + Firebase.', 'UC-03'],
    ['FR-04', 'Chơi game', 'Hiển thị 4 mini-game: Quiz, Matching, Memory, Word Puzzle; mỗi game có 3 độ khó.', 'UC-05a→d'],
    ['FR-05', 'Chơi game', 'Chấm điểm 0–250 (Quiz), 0–120 (Matching), 0–100 (Memory/Puzzle); chấm sao 0–3.', 'UC-06'],
    ['FR-06', 'Chơi game', 'Mở level kế tiếp khi đạt sao đủ ngưỡng (≥ 2 hoặc 3 với Memory).', 'UC-07'],
    ['FR-07', 'Tiến độ', 'Cập nhật streak, longestStreak; thưởng milestone 3/7/14/30/60/100 ngày.', 'UC-13'],
    ['FR-08', 'Tiến độ', 'Tính level user (1–10) theo tổng XP tích luỹ.', '—'],
    ['FR-09', 'Cá nhân hoá', 'Xem hồ sơ; sửa tên hiển thị, đổi avatar (≤ 500 KB, lưu base64).', 'UC-08, UC-09'],
    ['FR-10', 'Cá nhân hoá', 'Bật/tắt dark mode, đổi ngôn ngữ VI/EN, bật/tắt âm thanh, đặt giờ nhắc.', 'UC-12'],
    ['FR-11', 'Cộng đồng', 'Hiển thị bảng xếp hạng top 100 theo Score và XP.', 'UC-11'],
    ['FR-12', 'Mua sắm', 'Mua pack từ vựng bằng coin (transaction atomic kiểm coin + ownedPacks).', 'UC-10'],
    ['FR-13', 'Lịch sử', 'Hiển thị 50 game gần nhất với điểm/sao/thời lượng.', '—'],
])

H1('2. Yêu cầu phi chức năng (Non-functional Requirements)')
add_table([
    ['Mã', 'Loại', 'Yêu cầu cụ thể'],
    ['NFR-01', 'Hiệu năng',
     'Khởi động lạnh ≤ 3 giây trên thiết bị phổ thông. Animation game đạt ≥ 60 fps. '
     'Lệnh đọc/ghi Firestore ≤ 1 giây trong điều kiện mạng 4G.'],
    ['NFR-02', 'Bảo mật',
     'Mật khẩu chỉ qua Firebase Auth (không lưu local plaintext). Firestore Security '
     'Rules cấm user A truy cập dữ liệu user B. Không log password ra console.'],
    ['NFR-03', 'Tin cậy',
     'Lưu kết quả game phải atomic (1 transaction): coin/XP/streak/result hoặc lưu '
     'đầy đủ hoặc không lưu gì. Không có trạng thái "lưng chừng".'],
    ['NFR-04', 'Khả dụng',
     'Hỗ trợ offline cho favorites và settings (SharedPreferences). Khi mất mạng vẫn '
     'cho người dùng chơi local; sync lên Firestore khi có lại mạng.'],
    ['NFR-05', 'Tương thích',
     'Android 7.0+ (API 24+), iOS 13+, web Chrome/Edge/Safari mới nhất.'],
    ['NFR-06', 'Mở rộng',
     'Cấu trúc Provider tách services, models, screens. Thêm game mới chỉ cần '
     'tạo screen mới + thêm enum gameType.'],
    ['NFR-07', 'Bảo trì',
     'Mã nguồn theo convention Dart, có docstring tiếng Việt cho service & game logic. '
     'Tách config (constants.dart, design_tokens.dart, theme.dart) khỏi business logic.'],
    ['NFR-08', 'Đa ngôn ngữ',
     'Toàn bộ UI text qua AppLocalizations. Hot-switch VI/EN không cần restart.'],
])


# =============================================================
# PHẦN VII — KIẾN TRÚC HỆ THỐNG
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN VII — KIẾN TRÚC HỆ THỐNG', style='Title')

H1('1. Kiến trúc tổng thể')
P('Hệ thống VocabQuest áp dụng kiến trúc client-server hai tầng (2-tier) với Firebase '
  'làm Backend-as-a-Service (BaaS) — không cần server riêng. Client Flutter chạy đa nền '
  'tảng, giao tiếp trực tiếp với Firebase Auth (xác thực) và Cloud Firestore (lưu trữ).')

P('Sơ đồ kiến trúc tổng thể (PlantUML):')
arch_uml = '''@startuml VocabQuest_Architecture
skinparam linetype ortho
skinparam shadowing false
skinparam component {
  BackgroundColor #FFF8E7
  BorderColor #444444
}
skinparam database {
  BackgroundColor #E8F5E9
  BorderColor #2E7D32
}
skinparam cloud {
  BackgroundColor #E3F2FD
  BorderColor #1565C0
}
title Kiến trúc tổng thể — VocabQuest

actor "Người dùng" as User

package "CLIENT (Flutter App)" {
  component "UI Layer\\n(Screens + Widgets)"          as UI
  component "State Layer\\n(Providers)"               as STATE
  component "Service Layer\\n(Auth/Firestore/Audio)"  as SVC
  component "Data Models"                             as MODEL
  database "Local Storage\\n(SharedPreferences)"      as LOCAL
}

cloud "FIREBASE (BaaS)" {
  component "Firebase Auth"      as FAUTH
  database  "Cloud Firestore"    as FSTORE
  component "Firebase Storage"   as FSTORAGE
}

User --> UI
UI --> STATE
STATE --> SVC
SVC --> MODEL
SVC --> LOCAL
SVC --> FAUTH
SVC --> FSTORE
SVC --> FSTORAGE
@enduml
'''
add_code_block(arch_uml)

H1('2. Tầng Client — Flutter')
add_table([
    ['Layer', 'Vai trò', 'Thư mục / file đại diện'],
    ['UI Layer (Presentation)',
     'Render giao diện, nhận input người dùng, đẩy event lên State Layer.',
     'lib/screens/, lib/widgets/'],
    ['State Layer (Provider)',
     'Quản lý state chung (UserProvider, GameProvider, FavoritesProvider, '
     'SettingsProvider). Notify listeners khi state đổi.',
     'lib/providers/'],
    ['Service Layer',
     'Đóng gói gọi API Firebase, TTS, Audio, Notification. Trả về Future/Stream.',
     'lib/services/'],
    ['Data Models',
     'POJO class, factory fromFirestore() / toMap() để map JSON ↔ Object.',
     'lib/models/'],
    ['Local Storage',
     'Lưu setting (theme, ngôn ngữ, sound), favorites, isFirstTime.',
     'shared_preferences'],
])

H1('3. Tầng Backend — Firebase BaaS')
add_table([
    ['Dịch vụ', 'Vai trò', 'Lưu ý'],
    ['Firebase Authentication',
     'Quản lý tài khoản email/password; cấp UID dùng làm khóa user document.',
     'Không lưu password ở Firestore.'],
    ['Cloud Firestore',
     'NoSQL document database; lưu users, game_results, vocabularies, '
     'leaderboard. Realtime sync qua snapshot stream.',
     'Áp dụng Security Rules — xem PHẦN VIII.'],
    ['Firebase Storage',
     '(Tùy chọn) Lưu avatar người dùng kích cỡ lớn. Hiện app dùng base64 trong '
     'Firestore để đơn giản hoá deploy.',
     'Có thể bật khi cần.'],
])

H1('4. Pattern Provider — State Management')
P('App áp dụng MVVM-style với package `provider`: View (Screen) → ViewModel (Provider) → '
  'Model. Mỗi Provider extends ChangeNotifier; gọi notifyListeners() khi state đổi để '
  'rebuild các Consumer widget tương ứng.')


# =============================================================
# PHẦN VIII — THIẾT KẾ CSDL FIREBASE FIRESTORE
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN VIII — THIẾT KẾ CSDL FIREBASE FIRESTORE', style='Title')

H1('1. Tổng quan Cloud Firestore')
P('Cloud Firestore là cơ sở dữ liệu NoSQL, document-oriented, lưu dữ liệu dạng JSON-like '
  'theo cấu trúc collection → document → (subcollection / fields). Khác với SQL truyền '
  'thống, Firestore không có schema cứng và không hỗ trợ JOIN — quan hệ phải được mô '
  'hình hoá bằng cách: (1) embed (nhúng) dữ liệu con vào document cha, hoặc (2) reference '
  '(lưu ID document liên quan và query thêm).')

P('VocabQuest dùng kết hợp cả hai: dữ liệu tiến độ (progress, ownedPacks, streak) được '
  'EMBED trong document users; còn lịch sử game (game_results) là collection riêng có '
  'reference tới userId — vì nó tăng nhanh, không nên embed mảng quá lớn vào users.')

H1('2. Sơ đồ tổng quan các Collection')

erd_uml = '''@startuml VocabQuest_Firestore_Schema
skinparam linetype ortho
skinparam shadowing false
skinparam class {
  BackgroundColor #FFF8E7
  BorderColor #444444
}
title Sơ đồ Collections — Cloud Firestore VocabQuest

class "users/{uid}" as users {
  + uid : string (= doc.id)
  + email : string
  + displayName : string
  + avatarUrl : string (base64 data URI)
  --
  + totalScore : int
  + totalCoins : int
  + totalXP : int
  + level : int (1-10)
  --
  + streak : int
  + longestStreak : int
  + lastPlayedDate : timestamp
  --
  + hearts : int (0-5)
  + ownedPacks : array<string>
  + progress : map<string, int>
  + createdAt : timestamp
}

class "game_results/{auto_id}" as results {
  + id : string (auto)
  + userId : string  (FK → users.uid)
  + userDisplayName : string
  + gameType : string
  + level : string
  + score : int
  + correctAnswers : int
  + totalQuestions : int
  + timeTakenSeconds : int
  + coinsEarned : int
  + xpEarned : int
  + playedAt : timestamp
}

class "user_settings/{uid}" as settings {
  + theme : string
  + language : string
  + soundEnabled : bool
  + notificationEnabled : bool
  + reminderHour : int
  + reminderMinute : int
}

class "vocabularies/{vocabId}" as vocabs {
  + id : string
  + word : string
  + meaning : string
  + level : string
  + topic : string
  + audioUrl : string
}

class "achievements/{achievementId}" as achv {
  + id : string
  + name : string
  + description : string
  + iconUrl : string
  + xpReward : int
  + condition : string
}

class "user_achievements/{auto_id}" as user_achv {
  + userId : string  (FK)
  + achievementId : string (FK)
  + unlockedAt : timestamp
}

users "1" -- "*" results : 1 user có nhiều game_result
users "1" -- "1" settings : 1-1
users "1" -- "*" user_achv : 1-n
achv  "1" -- "*" user_achv : 1-n

@enduml
'''
add_code_block(erd_uml)

H1('3. Chi tiết từng Collection')

H2('3.1. Collection users')
P('Document key = UID Firebase Auth (đảm bảo 1-1 với tài khoản). Embed các trường tiến '
  'độ thay vì tạo subcollection — vì dữ liệu nhỏ, đọc thường xuyên.')
add_table([
    ['Trường', 'Kiểu', 'Mặc định', 'Ghi chú'],
    ['uid', 'string', '(= doc.id)', 'Khóa chính, trùng UID Auth'],
    ['email', 'string', '""', 'Email đăng ký'],
    ['displayName', 'string', '""', 'Tên hiển thị (sửa được)'],
    ['avatarUrl', 'string', '""', 'Data URI base64 (≤ 500 KB)'],
    ['totalScore', 'int', '0', 'Tổng điểm tích luỹ qua mọi game'],
    ['totalCoins', 'int', '100', 'Coin hiện có (khởi tạo +100)'],
    ['totalXP', 'int', '0', 'Tổng XP — quyết định user level'],
    ['level', 'int', '1', 'User level 1–10 (tính từ totalXP)'],
    ['streak', 'int', '0', 'Số ngày liên tiếp học hiện tại'],
    ['longestStreak', 'int', '0', 'Streak cao nhất từng đạt'],
    ['hearts', 'int', '5', 'Tim hiện tại (cơ chế dự phòng)'],
    ['ownedPacks', 'array<string>', '["beginner","intermediate","advanced"]',
     '3 pack mặc định miễn phí'],
    ['progress', 'map<string,int>', '{}',
     'Key dạng "gameType|packId" → level cao nhất đã pass (1–3)'],
    ['lastPlayedDate', 'timestamp', 'null', 'Để tính streak'],
    ['createdAt', 'timestamp', 'now()', 'Thời điểm đăng ký'],
])

H2('3.2. Collection game_results')
P('Mỗi document = 1 lần chơi xong. Dùng auto-id. Index theo userId + playedAt để query '
  'lịch sử nhanh.')
add_table([
    ['Trường', 'Kiểu', 'Ghi chú'],
    ['userId', 'string', 'FK → users.uid; Security Rule chỉ cho owner ghi'],
    ['userDisplayName', 'string', 'Snapshot tên user lúc chơi (denormalize cho leaderboard)'],
    ['gameType', 'string', 'matching | quiz | memory | word_puzzle'],
    ['level', 'string', 'beginner | intermediate | advanced'],
    ['score', 'int', 'Điểm số ván chơi'],
    ['correctAnswers', 'int', 'Số câu/ô đúng'],
    ['totalQuestions', 'int', 'Tổng câu/ô'],
    ['timeTakenSeconds', 'int', 'Thời lượng chơi'],
    ['coinsEarned', 'int', 'Coin thưởng (đã cộng vào users.totalCoins)'],
    ['xpEarned', 'int', 'XP thưởng (đã cộng vào users.totalXP)'],
    ['playedAt', 'timestamp', 'Khi nào lưu — sort lịch sử'],
])

H2('3.3. Collection user_settings')
P('Tách riêng khỏi users vì settings có thể đồng bộ sang nhiều thiết bị độc lập với '
  'dữ liệu chơi. App hiện chủ yếu lưu local SharedPreferences, collection này dự phòng '
  'cho sync đa thiết bị.')

H2('3.4. Collection vocabularies')
P('Dữ liệu master từ vựng. Chỉ admin write — user read-only (Security Rules). Hiện '
  'app load từ assets JSON để giảm cost Firestore, collection này dự phòng khi cần '
  'cập nhật từ vựng động.')

H2('3.5. Collection achievements & user_achievements')
P('achievements là master list các thành tựu (huy hiệu, mốc). user_achievements ghi '
  'lại user nào đã mở khóa achievement nào (quan hệ many-to-many).')

H1('4. Quan hệ giữa các Collection')
add_table([
    ['Quan hệ', 'Loại', 'Cách hiện thực'],
    ['users ↔ game_results', '1 — n',
     'game_results.userId tham chiếu users.uid. Query: where("userId", "==", uid).'],
    ['users ↔ user_settings', '1 — 1',
     'Cùng key = UID. Đọc song song với users khi cần.'],
    ['users ↔ achievements', 'n — n',
     'Qua bảng trung gian user_achievements.'],
    ['users ↔ vocabularies', 'n — n (gián tiếp)',
     'users.ownedPacks chứa packId; vocabularies thuộc 1 pack qua trường level/topic. '
     'App lọc bằng ownedPacks ở client.'],
])

H1('5. Security Rules — bảo vệ dữ liệu')
P('Toàn bộ truy cập Firestore phải qua rules; không có client nào bypass được. Quy tắc '
  'thực tế trong file firestore.rules:')

rules_block = '''rules_version = '2';
service cloud.firestore {
  match /databases/{database}/documents {

    function isAuthenticated() { return request.auth != null; }
    function isOwner(userId)   { return isAuthenticated()
                                       && request.auth.uid == userId; }

    // users — đọc: ai đã đăng nhập; ghi: chỉ chủ document
    match /users/{userId} {
      allow read: if isAuthenticated();
      allow create, update: if isOwner(userId);
      allow delete: if false;       // không cho xoá
    }

    // game_results — đọc: bất kỳ ai login (cho leaderboard);
    //                ghi: chỉ tạo, userId == auth.uid; cấm sửa/xoá
    match /game_results/{resultId} {
      allow read: if isAuthenticated();
      allow create: if isAuthenticated()
                       && request.resource.data.userId == request.auth.uid;
      allow update, delete: if false;
    }

    match /user_settings/{userId} {
      allow read, write: if isOwner(userId);
    }

    match /vocabularies/{vocabId} {
      allow read: if isAuthenticated();
      allow write: if false;       // chỉ admin (qua console)
    }

    // Mặc định cuối: cấm tất cả
    match /{document=**} {
      allow read, write: if false;
    }
  }
}'''
add_code_block(rules_block)

H1('6. Index & Query patterns')
add_table([
    ['Use case', 'Query', 'Index cần'],
    ['Lịch sử user',
     'where("userId", "==", uid).orderBy("playedAt", DESC).limit(50)',
     'Composite (userId ASC, playedAt DESC)'],
    ['Leaderboard theo score',
     'orderBy("totalScore", DESC).limit(100)',
     'Single field totalScore (auto)'],
    ['Leaderboard theo XP',
     'orderBy("totalXP", DESC).limit(100)',
     'Single field totalXP (auto)'],
    ['Lấy user',
     'doc(uid).get()/.snapshots()',
     'Không cần — lookup theo doc-id'],
])

H1('7. Atomic Transactions trong app')
P('Hai transaction quan trọng (xem chi tiết trong PHẦN I — mục 8):')
add_table([
    ['Transaction', 'Mục đích', 'Đảm bảo'],
    ['saveGameResult',
     'Lưu game_result + cộng coin/XP/score + cập nhật streak/milestone trong 1 batch.',
     'Hoặc lưu đầy đủ hoặc rollback toàn bộ — không có trạng thái lưng chừng.'],
    ['purchasePack',
     'Trừ coin + append packId vào ownedPacks.',
     'Không cho mua trùng / thiếu coin; tránh race condition khi mua liên tiếp.'],
    ['updateLevelProgress',
     'Cập nhật progress[gameType|packId] khi user pass level mới + cộng level reward.',
     'Chỉ ghi nếu level mới > level hiện tại (idempotent).'],
])


# =============================================================
# PHẦN IX — BIỂU ĐỒ HOẠT ĐỘNG (Activity Diagram)
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN IX — BIỂU ĐỒ HOẠT ĐỘNG (Activity Diagram)', style='Title')

H1('1. Activity 1 — Đăng ký tài khoản')
act1 = '''@startuml ActDangKy
skinparam shadowing false
title Activity — Đăng ký tài khoản

start
:Mở màn Register;
:Nhập tên, email, mật khẩu, xác nhận;
:Tích checkbox điều khoản;
if (Form hợp lệ?) then (Không)
  :Hiển thị lỗi validator;
  stop
endif
:Gọi FirebaseAuth.createUserWithEmailAndPassword;
if (Email đã tồn tại?) then (Có)
  :Hiển thị "Email đã được đăng ký";
  stop
endif
:Tạo document users/{uid} với defaults
 (totalCoins=100, ownedPacks=[3 pack free]);
:Đăng nhập tự động;
:Điều hướng sang Home;
stop
@enduml
'''
add_code_block(act1)

H1('2. Activity 2 — Chơi 1 ván game (luồng chính)')
act2 = '''@startuml ActChoiGame
skinparam shadowing false
title Activity — Chơi 1 ván Mini-game

start
:Chọn pack & level từ Pack Selection;
:Mở Game Screen tương ứng;
:GameProvider.startGame() — bắt đầu đếm thời gian;
repeat
  :Người dùng tương tác (chọn đáp án / lật thẻ / nối từ);
  if (Đúng?) then (Có)
    :Cộng điểm, addCorrect();
  else (Không)
    :addWrong();
  endif
repeat while (Còn lượt và còn thời gian?) is (Có)
->Không;
:GameProvider.finishGame()
 → tính score, accuracy, sao;
:[Transaction Firestore] saveGameResult():
  - Lưu game_result document
  - Cộng coin/XP/totalScore vào users
  - Tính streak, milestone bonus
  - Cập nhật lastPlayedDate;
if (Sao >= ngưỡng (>=2 hoặc =3 cho Memory)?) then (Có)
  :updateLevelProgress() — mở level mới
   + Level Reward;
else (Không)
endif
:Mở GameResultScreen
 (điểm, sao, coin/XP nhận, streak banner nếu milestone);
stop
@enduml
'''
add_code_block(act2)

H1('3. Activity 3 — Mua pack')
act3 = '''@startuml ActMuaPack
skinparam shadowing false
title Activity — Mua pack từ vựng

start
:Mở Shop;
:Chọn pack muốn mua;
:Nhấn nút "Mua" (giá X coin);
:[Transaction] purchasePack():
  - Đọc users/{uid};
if (totalCoins >= X?) then (Không)
  :Throw "Không đủ xu";
  :Toast lỗi;
  stop
endif
if (packId đã trong ownedPacks?) then (Có)
  :Throw "Đã sở hữu";
  :Toast lỗi;
  stop
endif
:Trừ X coin, append packId;
:Commit transaction;
:Cập nhật UserProvider local;
:Toast "Mua thành công";
stop
@enduml
'''
add_code_block(act3)


# =============================================================
# PHẦN X — BIỂU ĐỒ TUẦN TỰ (Sequence Diagram)
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN X — BIỂU ĐỒ TUẦN TỰ (Sequence Diagram)', style='Title')

H1('1. Sequence — Lưu kết quả game (saveGameResult)')
seq1 = '''@startuml SeqSaveResult
skinparam shadowing false
title Sequence — Lưu kết quả game (atomic)

actor User
participant "GameScreen" as UI
participant "GameProvider" as GP
participant "FirestoreService" as FS
database "Firestore" as DB

User -> UI : Hoàn thành ván chơi
UI -> GP : finishGame(userId, gameType, level)
GP -> GP : _calculateRewards() (coin/XP/bonus)
GP -> FS : saveGameResult(userId, result)
FS -> DB : runTransaction(BEGIN)
DB -> FS : snapshot users/{uid}
FS -> FS : tính newStreak, milestone từ lastPlayedDate
FS -> DB : update users (totalCoins/XP/score, streak, lastPlayedDate)
FS -> DB : create game_results/{auto_id}
DB -> FS : COMMIT
FS --> GP : GameOutcome(result, streakOutcome)
GP --> UI : outcome
UI -> User : Mở GameResultScreen
@enduml
'''
add_code_block(seq1)

H1('2. Sequence — Mua pack')
seq2 = '''@startuml SeqMuaPack
skinparam shadowing false
title Sequence — Mua pack từ vựng

actor User
participant "ShopScreen" as UI
participant "UserProvider" as UP
participant "FirestoreService" as FS
database "Firestore" as DB

User -> UI : Nhấn "Mua" pack X (giá P)
UI -> UP : purchasePack(uid, packId, P)
UP -> FS : purchasePack(uid, packId, P)
FS -> DB : runTransaction(BEGIN)
DB -> FS : snapshot users/{uid}
alt totalCoins < P
  FS --> UP : throw "Không đủ xu"
  UP --> UI : error
  UI -> User : Toast lỗi
else đã sở hữu
  FS --> UP : throw "Đã sở hữu"
  UP --> UI : error
  UI -> User : Toast lỗi
else OK
  FS -> DB : update users.totalCoins -= P, ownedPacks += [packId]
  DB -> FS : COMMIT
  FS --> UP : new UserModel
  UP -> UP : notifyListeners()
  UP --> UI : success
  UI -> User : Toast "Mua thành công"
end
@enduml
'''
add_code_block(seq2)


# =============================================================
# PHẦN XI — BIỂU ĐỒ LỚP (Class Diagram)
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN XI — BIỂU ĐỒ LỚP (Class Diagram)', style='Title')

P('Biểu đồ lớp tóm tắt các thành phần chính: Models (data), Providers (state), '
  'Services (Firebase + utility). Đã giản lược nhiều thuộc tính phụ để giữ rõ.')

class_uml = '''@startuml VocabQuest_ClassDiagram
skinparam linetype ortho
skinparam shadowing false
skinparam class {
  BackgroundColor<<model>>    #FFF8E7
  BackgroundColor<<provider>> #E3F2FD
  BackgroundColor<<service>>  #E8F5E9
}
title Class Diagram — VocabQuest (rút gọn)

' ===== MODELS =====
class UserModel <<model>> {
  + uid : String
  + email : String
  + displayName : String
  + totalScore : int
  + totalCoins : int
  + totalXP : int
  + level : int
  + streak : int
  + ownedPacks : List<String>
  + progress : Map<String,int>
  + fromFirestore(doc) : UserModel
  + toMap() : Map
  + copyWith(...) : UserModel
}

class GameResultModel <<model>> {
  + userId : String
  + gameType : String
  + level : String
  + score : int
  + correctAnswers : int
  + timeTakenSeconds : int
  + coinsEarned : int
  + xpEarned : int
  + accuracy : double
  + stars : int
}

class VocabModel <<model>> {
  + id : String
  + word : String
  + meaning : String
  + level : String
}

' ===== PROVIDERS =====
class UserProvider <<provider>> {
  - _user : UserModel?
  + load(uid) : void
  + updateLocalUser(addScore, addCoins, addXP) : void
  + recordLevelComplete(...) : Future<LevelReward?>
  + applyStreakOutcome(outcome) : void
  + notifyListeners()
}

class GameProvider <<provider>> {
  - _score : int
  - _correct : int
  - _wrong : int
  - _startTime : DateTime
  + startGame() : void
  + addScore(d) : void
  + addCorrect() : void
  + addWrong() : void
  + finishGame(...) : Future<GameOutcome>
}

class FavoritesProvider <<provider>> {
  - _items : Set<String>
  + toggle(id) : void
  + contains(id) : bool
}

class SettingsProvider <<provider>> {
  + theme : ThemeMode
  + language : String
  + soundEnabled : bool
  + notificationEnabled : bool
}

' ===== SERVICES =====
class FirestoreService <<service>> {
  + getUser(uid) : Future<UserModel?>
  + streamUser(uid) : Stream
  + saveGameResult(...) : Future
  + purchasePack(uid, packId, price) : Future<UserModel>
  + updateLevelProgress(...) : Future<bool>
  + streamLeaderboard(...) : Stream
}

class AuthService <<service>> {
  + signUp(email, password) : Future
  + signIn(email, password) : Future
  + signOut() : Future
  + resetPassword(email) : Future
}

class AudioService <<service>> {
  + playClick() : void
  + playCorrect() : void
  + playWrong() : void
}

class JsonService <<service>> {
  + loadVocab(level) : Future<List<VocabModel>>
}

' ===== Quan hệ =====
UserProvider     --> UserModel
GameProvider     --> GameResultModel
FirestoreService --> UserModel
FirestoreService --> GameResultModel
JsonService      --> VocabModel
UserProvider     --> FirestoreService : uses
GameProvider     --> FirestoreService : uses
UserProvider     --> AuthService      : uses

@enduml
'''
add_code_block(class_uml)


# =============================================================
# PHẦN XII — THIẾT KẾ GIAO DIỆN
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN XII — THIẾT KẾ GIAO DIỆN', style='Title')

H1('1. Sơ đồ điều hướng (Navigation Graph)')
nav_uml = '''@startuml NavGraph
skinparam linetype ortho
skinparam shadowing false
skinparam state {
  BackgroundColor #FFF8E7
  BorderColor #444444
}
title Sơ đồ điều hướng giữa các màn hình

[*] --> Splash
Splash --> Login : chưa login
Splash --> Home  : đã login

state Login {
  Login --> Register
  Login --> ForgotPassword
}
Register --> Home
ForgotPassword --> Login

state Home {
  Home --> GameMenu        : tile "Chơi game"
  Home --> Leaderboard     : tile "Xếp hạng"
  Home --> Profile         : icon avatar
  Home --> Shop            : tile "Cửa hàng"
  Home --> Favorites
  Home --> History
  Home --> Settings
}

GameMenu --> PackSelection
PackSelection --> LevelMap
LevelMap --> GameScreen : Quiz/Match/Mem/Puzzle
GameScreen --> GameResult
GameResult --> Home
GameResult --> GameScreen : Chơi lại

Profile --> EditProfile
Settings --> Login : Đăng xuất
@enduml
'''
add_code_block(nav_uml)

H1('2. Bảng các màn hình chính')
add_table([
    ['Mã màn', 'Tên', 'Mục đích', 'Widget chính'],
    ['SC-01', 'Splash', 'Hiển thị logo, kiểm tra trạng thái auth, điều hướng tiếp.',
     'Animated logo, auto-redirect'],
    ['SC-02', 'Login', 'Đăng nhập email/mật khẩu.',
     'TextField, ElevatedButton, link Register/Forgot'],
    ['SC-03', 'Register', 'Tạo tài khoản mới.',
     'Form 4 trường + checkbox điều khoản'],
    ['SC-04', 'Forgot Password', 'Gửi email reset mật khẩu.', 'TextField email'],
    ['SC-05', 'Home', 'Hub trung tâm: stats, daily challenge, tile điều hướng.',
     'Header user, streak chip, daily progress, grid tiles'],
    ['SC-06', 'Game Menu', 'Chọn 1 trong 4 mini-game.',
     'Card 4 game với gradient riêng'],
    ['SC-07', 'Pack Selection', 'Chọn pack từ vựng.',
     'List/Grid pack với owned/locked badge'],
    ['SC-08', 'Level Map', 'Chọn level (1–3) trong pack.',
     'Path UI với star indicator'],
    ['SC-09', 'Quiz Game', 'Trả lời 10 câu trắc nghiệm.',
     'Question card + 4 option, timer 15s, TTS button'],
    ['SC-10', 'Matching Game', 'Nối từ ↔ nghĩa, 2 cột.',
     'Two-column tap-to-match'],
    ['SC-11', 'Memory Game', 'Lật thẻ tìm cặp.',
     'Grid 3D flip cards, badge EN/VI'],
    ['SC-12', 'Word Puzzle', 'Sắp xếp chữ thành từ.',
     'Letter bank + slot row, hint button'],
    ['SC-13', 'Game Result', 'Tổng kết ván chơi.',
     'Stars animation, coin/XP, streak banner, replay'],
    ['SC-14', 'Profile', 'Thông tin user, stats game.',
     'Avatar, name, stats grid'],
    ['SC-15', 'Edit Profile', 'Sửa tên + đổi avatar.',
     'TextField, image picker'],
    ['SC-16', 'Leaderboard', 'Top 100 score/XP.',
     'Tab Score/XP, list user, highlight self'],
    ['SC-17', 'Favorites', 'Danh sách đã yêu thích.', 'Empty state + list'],
    ['SC-18', 'History', '50 game gần nhất.', 'List with date/score/stars'],
    ['SC-19', 'Shop', 'Mua pack từ vựng.', 'Pack card + giá + nút mua'],
    ['SC-20', 'Settings', 'Cấu hình app.',
     'Switches: theme, lang, sound, noti, time picker'],
])


# =============================================================
# PHẦN XIII — KIỂM THỬ (PLACEHOLDER)
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN XIII — KIỂM THỬ HỆ THỐNG', style='Title')
P('[ CHỖ DÀNH RIÊNG ] — Phần này được tách thành tài liệu riêng "VocabQuest_TestCases_v2.docx" '
  'gồm 98 test case theo 27 module, đã đối chiếu với code và đánh trạng thái '
  'Pass / Fail / Blocked. Bạn copy nội dung từ file đó và dán vào đây để gộp vào báo cáo '
  'tổng hợp. Cấu trúc gợi ý của Phần này:')
add_table([
    ['Mục', 'Nội dung sẽ chèn'],
    ['1. Giới thiệu kiểm thử', 'Mục đích, phạm vi kiểm thử (từ file Test Cases mục 1–2)'],
    ['2. Chiến lược & tiêu chí', 'Loại test, tiêu chí Pass/Fail (mục 3 file Test Cases)'],
    ['3. Môi trường', 'Thiết bị, framework, mạng, dữ liệu (mục 4)'],
    ['4. Mô tả 8 cột test case', 'Ý nghĩa các cột bảng test case (mục 5)'],
    ['5. Tổng hợp test case theo module', '27 module – 98 TC (mục 6)'],
    ['6. Chi tiết 98 test case', 'Bảng test case ngang theo module (mục 7)'],
    ['7. Kết quả thực thi', 'Pass = 95, Fail = 1, Blocked = 2 (cập nhật mới nhất)'],
    ['8. Tổng kết kiểm thử', 'Khuyến nghị triển khai (mục 8)'],
])


# =============================================================
# PHẦN XIV — TRIỂN KHAI & VẬN HÀNH
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN XIV — TRIỂN KHAI & VẬN HÀNH', style='Title')

H1('1. Cấu hình môi trường phát triển')
add_table([
    ['Thành phần', 'Phiên bản / Yêu cầu'],
    ['Flutter SDK', 'Stable channel (3.x)'],
    ['Dart', '3.x (đi kèm Flutter)'],
    ['Android Studio', 'Hedgehog 2023+ (build Android, emulator)'],
    ['Xcode', '15+ (chỉ cần khi build iOS)'],
    ['Firebase CLI', 'flutterfire_cli để cấu hình firebase_options.dart'],
    ['JDK', '17 (gradle 8.x yêu cầu)'],
])

H1('2. Quy trình build & deploy')
add_table([
    ['Nền tảng', 'Lệnh build', 'Output', 'Phát hành'],
    ['Android (debug)', 'flutter build apk --debug', 'app-debug.apk', 'Test nội bộ'],
    ['Android (release)', 'flutter build appbundle', 'app-release.aab', 'Google Play'],
    ['iOS', 'flutter build ipa', '*.ipa', 'TestFlight → App Store'],
    ['Web', 'flutter build web --release', 'build/web/', 'Firebase Hosting'],
])

H1('3. Cấu hình Firebase')
P('1. Tạo project Firebase qua console. 2. Bật Authentication (Email/Password). '
  '3. Tạo Firestore database (mode native). 4. Deploy rules: firebase deploy --only firestore:rules. '
  '5. Cài firebase_options.dart bằng flutterfire configure. 6. Tải google-services.json '
  '(Android) và GoogleService-Info.plist (iOS) — đã include trong repo.')

H1('4. Backup & Monitoring')
add_table([
    ['Hạng mục', 'Cách thực hiện'],
    ['Backup Firestore',
     'gcloud firestore export gs://<bucket>/backup-<date> — hàng tuần.'],
    ['Crash report',
     'Tích hợp Firebase Crashlytics (chưa bật, để chương sau).'],
    ['Analytics',
     'Firebase Analytics (track event đăng ký, chơi game, mua pack).'],
    ['Cost monitoring',
     'Xem Firebase usage tab; alert khi gần ngưỡng quota free.'],
])


# =============================================================
# PHẦN XV — KẾT LUẬN & HƯỚNG PHÁT TRIỂN
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHẦN XV — KẾT LUẬN & HƯỚNG PHÁT TRIỂN', style='Title')

H1('1. Kết quả đạt được')
add_table([
    ['Khía cạnh', 'Đạt được'],
    ['Chức năng', '4 mini-game hoạt động ổn định, hệ thống điểm/sao/streak/level đầy đủ, '
                  'shop mua pack, leaderboard top 100, profile + edit, đa ngôn ngữ VI/EN, '
                  'dark/light mode.'],
    ['Kỹ thuật', 'Provider state management, Firestore transaction atomic, Security '
                 'Rules chặt chẽ, animation 3D flip mượt 60fps, hỗ trợ 6 nền tảng.'],
    ['Kiểm thử', '95/98 test case Pass theo bộ 27 module — tỷ lệ 96.9%.'],
    ['Tài liệu', 'Báo cáo phân tích đầy đủ 15 phần, kèm 9 biểu đồ UML PlantUML.'],
])

H1('2. Hạn chế hiện tại')
add_table([
    ['Hạn chế', 'Mô tả'],
    ['Avatar lưu base64',
     'Tăng size document users (~500KB). Nên migrate sang Firebase Storage.'],
    ['Chưa có Crashlytics',
     'Khó truy vết crash thực tế từ user.'],
    ['Vocab static từ JSON',
     'Khó cập nhật từ vựng mà không phát hành app mới.'],
    ['Chưa có quên-mật-khẩu UI hoàn chỉnh',
     'Hiện chỉ gửi email reset, không có flow xác nhận trong app.'],
    ['Test trên thiết bị thật còn ít',
     '2 test case Blocked do thiếu thiết bị xoay màn hình + tablet thật.'],
])

H1('3. Hướng phát triển')
add_table([
    ['Tính năng', 'Mô tả & Ưu tiên'],
    ['AI sinh câu hỏi',
     'Dùng LLM tạo quiz đa dạng theo chủ đề người dùng — Cao.'],
    ['Voice recognition',
     'Bài tập phát âm đối chiếu với speech-to-text — Cao.'],
    ['Cloud Functions',
     'Tự động cleanup game_results cũ, tính achievement server-side — Trung bình.'],
    ['Bảng xếp hạng theo nhóm bạn',
     'Mời bạn bè tạo group leaderboard — Trung bình.'],
    ['In-app purchase',
     'Mua coin bằng tiền thật qua Google Play Billing / Apple StoreKit — Thấp.'],
    ['PWA Web đầy đủ',
     'Tối ưu Web build, deploy Firebase Hosting + tối ưu offline-first — Thấp.'],
    ['Migrate avatar sang Storage',
     'Giảm kích cỡ document users, tăng tốc đọc — Cao.'],
    ['Crashlytics + Analytics',
     'Theo dõi crash + hành vi user thực tế — Cao.'],
])


# =============================================================
# MỤC LỤC ĐỀ XUẤT (cuối file để tham chiếu)
# =============================================================
d.add_paragraph('', style='Normal')
d.add_paragraph('PHỤ LỤC — CẤU TRÚC BÁO CÁO ĐỀ XUẤT', style='Title')
P('Khi gộp các phần lại thành báo cáo cuối cùng, bố cục logic được khuyến nghị như sau '
  '(dùng cho Mục lục ở đầu file):')
add_table([
    ['Chương', 'Tên', 'Nguồn nội dung'],
    ['1', 'Tổng quan đề tài', 'PHẦN V'],
    ['2', 'Phân tích yêu cầu', 'PHẦN VI'],
    ['3', 'Kiến trúc hệ thống', 'PHẦN VII'],
    ['4', 'Biểu đồ phân cấp chức năng (BPC)', 'PHẦN II (đã có)'],
    ['5', 'Biểu đồ Use Case', 'PHẦN IV (đã có)'],
    ['6', 'Biểu đồ hoạt động (Activity)', 'PHẦN IX'],
    ['7', 'Biểu đồ tuần tự (Sequence)', 'PHẦN X'],
    ['8', 'Biểu đồ lớp (Class)', 'PHẦN XI'],
    ['9', 'Thiết kế CSDL Firebase Firestore', 'PHẦN VIII'],
    ['10', 'Logic nghiệp vụ & Hệ thống điểm', 'PHẦN I (đã có)'],
    ['11', 'Thiết kế giao diện', 'PHẦN XII'],
    ['12', 'Thư viện sử dụng', 'PHẦN III (đã có)'],
    ['13', 'Kiểm thử hệ thống', 'PHẦN XIII (chèn từ file Test Cases)'],
    ['14', 'Triển khai & Vận hành', 'PHẦN XIV'],
    ['15', 'Kết luận & Hướng phát triển', 'PHẦN XV'],
])

# ========== Save ==========
try:
    d.save(path)
    print('Saved to:', path)
except PermissionError:
    alt = path.replace('_v2.docx', '_v3.docx')
    d.save(alt)
    print('File locked. Saved to:', alt)
