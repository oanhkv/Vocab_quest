import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

path = r'C:\Users\Kieu Anh\Desktop\CD2\VocabQuest_TestCases.docx'
d = Document(path)

body = d.element.body
i_para = 0
i_table = 0
for child in body.iterchildren():
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        p = d.paragraphs[i_para]; i_para += 1
        print(f'P[{i_para-1}] ({p.style.name}) {p.text[:160]}')
    elif tag == 'tbl':
        t = d.tables[i_table]; i_table += 1
        print(f'TABLE[{i_table-1}] rows={len(t.rows)} cols={len(t.columns)}')
        for r_idx, row in enumerate(t.rows):
            cells = [c.text.replace("\n"," | ")[:80] for c in row.cells]
            print(f'  r{r_idx}: {cells}')
